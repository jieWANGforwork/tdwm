#!/usr/bin/env python3
"""Build the locked Results TD V2-EMA-SG new-score report.

The report is downstream of two independent archival validators.  It accepts
only the complete 96-cell EMA epoch sweep, the complete 24-cell fixed-
checkpoint comparison, the accepted V2-EMA-SG training/original-score archive,
and the immutable V0/V1 Results TD document.  Any missing cell, changed commit,
selection mismatch, malformed success rate, or loss-curve discrepancy aborts
before a report artifact is written.

The generated DOCX appends a landscape section to the versioned V0/V1 report;
it never modifies the base document in place.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import io
import json
import math
import os
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from types import ModuleType
from typing import Any, Iterable, Mapping, Sequence

REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_ARTIFACT_ROOT = (
    REPOSITORY_ROOT
    / "reports/artifacts/actor_free_td_lewm_v2_ema_sg_new_scores_cube_seed3072"
)
DEFAULT_ORIGINAL_SUMMARY = DEFAULT_ARTIFACT_ROOT / "original_scores/summary.json"
DEFAULT_TRAINING_LOSS_CSV = (
    DEFAULT_ARTIFACT_ROOT / "original_scores/training_loss_curves.csv"
)
DEFAULT_NEW_SCORE_ROOT = DEFAULT_ARTIFACT_ROOT / "new_scores"
DEFAULT_NEW_SUMMARY = DEFAULT_NEW_SCORE_ROOT / "summary.json"
DEFAULT_NEW_RESULTS_CSV = DEFAULT_NEW_SCORE_ROOT / "results.csv"
DEFAULT_FIXED_RESULTS_CSV = DEFAULT_NEW_SCORE_ROOT / "fixed_checkpoint_results.csv"
DEFAULT_BASE_DOCUMENT = (
    REPOSITORY_ROOT / "reports/results_td_actor_free_td_lewm_v0_v1_cube_seed3072.docx"
)
DEFAULT_MARKDOWN_OUTPUT = (
    REPOSITORY_ROOT / "reports/actor_free_td_lewm_v2_ema_sg_new_scores_cube_seed3072.md"
)
DEFAULT_DOCX_OUTPUT = (
    REPOSITORY_ROOT
    / "reports/results_td_actor_free_td_lewm_v2_ema_sg_new_scores_cube_seed3072.docx"
)
DEFAULT_CHART_DIR = DEFAULT_ARTIFACT_ROOT / "figures"

SCHEMA_VERSION = 1
TRAINING_COMMIT = "18cd574d522515f20f4103509b1e660b2fc89ea6"
EVALUATION_COMMIT = "5456f3d18116812d078d4ec2e85ba1f83d89c7c7"
SHARED_EPISODE_SELECTION_SHA256 = (
    "e46ea81cce2e6a9a5df05ba04893b4181cbd8979340111a012c30f1efa2d7ee7"
)
FIXED_LAUNCHER_RANKS_SHA256 = (
    "88c204770f33c0b0220057d45b187766e3cfc54912e3f5ca49f2aa93d16437e9"
)
ACTION_NORMALIZATION_SHA256 = (
    "57f4d3c252e1805f4af1f614d20d1d1a064fa0d1d463ed5eb8ecf9dfc2b1a723"
)
# The reconciler schema retains these field-oriented names.  They are aliases,
# not two episode sets: all 120 cells share the episode-selection file above;
# 88c... is the fixed launcher's canonical valid-row-ranks digest.
EMA_SELECTION_SHA256 = SHARED_EPISODE_SELECTION_SHA256
FIXED_SELECTION_SHA256 = FIXED_LAUNCHER_RANKS_SHA256
BASE_DOCUMENT_SHA256 = (
    "e98ffb1c51f2874fbac407c4e3bbc1528c9872bc6e470ee9b7aaa67e23415b67"
)
VARIANTS = ("c", "d", "f", "g1", "g2", "g3")
VARIANT_LABELS = {
    "c": "C Goal-Projected TD",
    "d": "D Goal-Value Weighted TD",
    "f": "F Same-Future Advantage",
    "g1": "G1 Neighbor Action Advantage",
    "g2": "G2 Prefix-Mean Advantage",
    "g3": "G3 Prefix-Marginal Advantage",
}
EMA_EPOCHS = tuple(range(3, 11))
TRAINING_EPOCHS = tuple(range(1, 11))
SCORE_MODES = ("f_plus_g_first", "g_only_f_rollout_mean")
ORIGINAL_SCORE_MODES = ("f_only", "g_only", "f_plus_g")
FIXED_VERSIONS = ("v0", "v1", "v2")
EPISODES = 50
TRAINING_SEED = 3072
PLANNING_SEED = 42
GOAL_OFFSET = 50
STEPS_PER_EPOCH = 12_796
FIRST_ACTION_WEIGHT = 0.25
PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"
SHA256_LENGTH = 64

NEW_RESULTS_FIELDS = (
    "epoch",
    "variant",
    "method",
    "score_mode",
    "g_first_weight",
    "success_count",
    "success_rate",
    "success_rate_percent",
    "checkpoint_epoch",
    "checkpoint_global_step",
    "checkpoint_sha256",
    "training_manifest_sha256",
    "evaluation_commit",
    "source_scope",
    "source_status",
    "cell_id",
)
FIXED_RESULTS_FIELDS = (
    "version",
    "variant",
    "method",
    "score_mode",
    "g_first_weight",
    "success_count",
    "success_rate",
    "success_rate_percent",
    "checkpoint_sha256",
    "selection_sha256",
    "episode_selection_file_sha256",
    "action_normalization_sha256",
    "evaluation_commit",
    "job_id",
    "source_launcher_manifest",
)
TRAINING_LOSS_FIELDS = (
    "variant",
    "method",
    "display_name",
    "epoch",
    "train_total_loss",
    "train_prediction_loss",
    "train_prediction_online_reference_mse",
    "train_online_ema_latent_drift",
    "train_base_hybrid_td_loss",
    "train_method_hybrid_td_loss",
    "validation_total_loss",
    "validation_prediction_loss",
    "validation_prediction_online_reference_mse",
    "validation_online_ema_latent_drift",
    "validation_base_hybrid_td_loss",
    "validation_method_hybrid_td_loss",
)


class ResultsTDV2EMAError(ValueError):
    """The requested report cannot be built from the supplied evidence."""


@dataclass(frozen=True)
class ValidatedReportInputs:
    original_summary: Mapping[str, Any]
    training_loss_rows: tuple[Mapping[str, Any], ...]
    new_summary: Mapping[str, Any]
    new_rows: tuple[Mapping[str, Any], ...]
    fixed_rows: tuple[Mapping[str, Any], ...]
    base_document: Path


def _load_v1_builder() -> ModuleType:
    """Load the existing table/font/landscape helpers without a package clash."""

    module_name = "_tdwm_results_td_v1_builder_helpers"
    existing = sys.modules.get(module_name)
    if existing is not None:
        return existing
    source = REPOSITORY_ROOT / "scripts/build_results_td_v1.py"
    spec = importlib.util.spec_from_file_location(module_name, source)
    if spec is None or spec.loader is None:
        raise ImportError(f"Cannot load Results TD V1 helpers from {source}.")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


_v1 = _load_v1_builder()


def _mapping(value: Any, *, context: str) -> Mapping[str, Any]:
    if not isinstance(value, Mapping):
        raise ResultsTDV2EMAError(f"{context} must be an object.")
    return value


def _sequence(value: Any, *, context: str) -> Sequence[Any]:
    if not isinstance(value, list):
        raise ResultsTDV2EMAError(f"{context} must be a list.")
    return value


def _exact_int(value: Any, expected: int, *, context: str) -> int:
    if isinstance(value, bool) or not isinstance(value, int) or value != expected:
        raise ResultsTDV2EMAError(f"{context} must equal {expected}, found {value!r}.")
    return value


def _parse_int(value: Any, *, context: str) -> int:
    if isinstance(value, bool):
        raise ResultsTDV2EMAError(f"{context} must be an integer.")
    try:
        parsed = int(value)
    except (TypeError, ValueError) as exc:
        raise ResultsTDV2EMAError(f"{context} must be an integer.") from exc
    if str(value).strip() not in {str(parsed), f"+{parsed}"}:
        raise ResultsTDV2EMAError(f"{context} is not an exact integer: {value!r}.")
    return parsed


def _finite_float(value: Any, *, context: str) -> float:
    if isinstance(value, bool):
        raise ResultsTDV2EMAError(f"{context} must be numeric.")
    try:
        parsed = float(value)
    except (TypeError, ValueError) as exc:
        raise ResultsTDV2EMAError(f"{context} must be numeric.") from exc
    if not math.isfinite(parsed):
        raise ResultsTDV2EMAError(f"{context} must be finite.")
    return parsed


def _text(value: Any, *, context: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ResultsTDV2EMAError(f"{context} must be non-empty text.")
    return value.strip()


def _sha256(value: Any, *, context: str) -> str:
    text = _text(value, context=context)
    if len(text) != SHA256_LENGTH or any(
        character not in "0123456789abcdef" for character in text
    ):
        raise ResultsTDV2EMAError(f"{context} must be lowercase SHA-256.")
    return text


def _read_json(path: Path, *, context: str) -> Mapping[str, Any]:
    if not path.is_file():
        raise FileNotFoundError(f"{context} does not exist: {path}")
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ResultsTDV2EMAError(f"Cannot parse {context} at {path}: {exc}") from exc
    return _mapping(value, context=context)


def _read_csv(
    path: Path, *, fields: Sequence[str], context: str
) -> tuple[dict[str, str], ...]:
    if not path.is_file():
        raise FileNotFoundError(f"{context} does not exist: {path}")
    try:
        with path.open(newline="", encoding="utf-8") as stream:
            reader = csv.DictReader(stream)
            if tuple(reader.fieldnames or ()) != tuple(fields):
                raise ResultsTDV2EMAError(
                    f"{context} columns must equal {list(fields)!r}, found "
                    f"{list(reader.fieldnames or ())!r}."
                )
            return tuple(dict(row) for row in reader)
    except OSError as exc:
        raise ResultsTDV2EMAError(f"Cannot read {context} at {path}: {exc}") from exc


def _validate_rate(
    successes: Any,
    rate: Any,
    percent: Any,
    *,
    context: str,
) -> tuple[int, float, float]:
    count = _parse_int(successes, context=f"{context}.success_count")
    if not 0 <= count <= EPISODES:
        raise ResultsTDV2EMAError(f"{context}.success_count is outside [0, 50].")
    parsed_rate = _finite_float(rate, context=f"{context}.success_rate")
    parsed_percent = _finite_float(percent, context=f"{context}.success_rate_percent")
    expected_rate = count / EPISODES
    expected_percent = 100.0 * expected_rate
    if not math.isclose(parsed_rate, expected_rate, rel_tol=0.0, abs_tol=1e-12):
        raise ResultsTDV2EMAError(f"{context}.success_rate disagrees with {count}/50.")
    if not math.isclose(parsed_percent, expected_percent, rel_tol=0.0, abs_tol=1e-12):
        raise ResultsTDV2EMAError(
            f"{context}.success_rate_percent disagrees with {count}/50."
        )
    return count, parsed_rate, parsed_percent


def _validate_original_summary(path: Path) -> Mapping[str, Any]:
    summary = _read_json(path, context="original_scores summary")
    _exact_int(summary.get("schema_version"), 1, context="original.schema_version")
    study = _mapping(summary.get("study"), context="original.study")
    expected_study: Mapping[str, Any] = {
        "environment": "cube",
        "method_family": "actor_free_td_lewm_v2_ema_sg",
        "implementation_version": "v2_ema_sg",
        "training_revision": TRAINING_COMMIT,
        "training_count": 6,
        "evaluation_count": 18,
        "episodes_per_evaluation": EPISODES,
        "training_seed": TRAINING_SEED,
        "planning_seed": PLANNING_SEED,
        "goal_offset": GOAL_OFFSET,
        "score_modes": list(ORIGINAL_SCORE_MODES),
        "single_training_seed": True,
        "single_planning_selection": True,
    }
    for key, expected in expected_study.items():
        if study.get(key) != expected:
            raise ResultsTDV2EMAError(
                f"original.study.{key} must equal {expected!r}, "
                f"found {study.get(key)!r}."
            )
    local = _mapping(
        study.get("local_prediction_contract"),
        context="original.study.local_prediction_contract",
    )
    if local != {
        "loss": "ema_target_lewm_one_step_mse",
        "target": "ema_world_model_next_latent",
        "target_gradient": "stop_gradient",
    }:
        raise ResultsTDV2EMAError(
            "The original-score local prediction contract changed."
        )

    architecture = _mapping(
        summary.get("architecture"), context="original.architecture"
    )
    expected_architecture = {
        "state_dim": 192,
        "raw_action_dim": 25,
        "action_embedding_dim": 192,
        "task_dim": 192,
        "output_dim": 192,
        "predictor_parameter_count": 379_072,
        "world_model_parameter_count": 18_034_628,
        "single_shared_action_encoder": True,
        "online_lewm_trainable": True,
        "target_lewm": "ema_full_world_model",
        "actor": "none",
        "reward": "none",
    }
    for key, expected in expected_architecture.items():
        if architecture.get(key) != expected:
            raise ResultsTDV2EMAError(
                f"original.architecture.{key} must equal {expected!r}."
            )

    acceptance = _mapping(
        summary.get("training_acceptance"), context="original.training_acceptance"
    )
    if acceptance.get("status") != "PASS" or acceptance.get("warnings") != []:
        raise ResultsTDV2EMAError("V2-EMA-SG training acceptance must be a clean PASS.")
    _sha256(acceptance.get("sha256"), context="original.training_acceptance.sha256")

    selection = _mapping(summary.get("selection"), context="original.selection")
    _exact_int(
        selection.get("episode_count"),
        EPISODES,
        context="original.selection.episode_count",
    )
    if selection.get("episode_selection_json_sha256") != EMA_SELECTION_SHA256:
        raise ResultsTDV2EMAError(
            "Original scores do not use the locked shared episode selection."
        )

    methods = _mapping(summary.get("methods"), context="original.methods")
    if set(methods) != set(VARIANTS):
        raise ResultsTDV2EMAError(
            "Original scores must contain exactly C/D/F/G1/G2/G3."
        )
    for variant in VARIANTS:
        method = _mapping(methods[variant], context=f"original.methods.{variant}")
        if method.get("method") != f"actor_free_td_lewm_v2_ema_sg_{variant}":
            raise ResultsTDV2EMAError(
                f"Original method identity changed for {variant}."
            )
        for key in ("display_name", "network", "training_loss", "special_mechanism"):
            _text(method.get(key), context=f"original.methods.{variant}.{key}")
        training = _mapping(
            method.get("training"), context=f"original.{variant}.training"
        )
        _exact_int(
            training.get("seed"),
            TRAINING_SEED,
            context=f"original.{variant}.training.seed",
        )
        _exact_int(
            training.get("epochs"), 10, context=f"original.{variant}.training.epochs"
        )
        _exact_int(
            training.get("global_step"),
            127_960,
            context=f"original.{variant}.training.global_step",
        )
        curve = _sequence(
            training.get("loss_curve"),
            context=f"original.{variant}.training.loss_curve",
        )
        if len(curve) != 10 or {
            row.get("epoch") for row in curve if isinstance(row, Mapping)
        } != set(TRAINING_EPOCHS):
            raise ResultsTDV2EMAError(
                f"Original {variant} loss curve must contain epochs 1-10."
            )
        evaluations = _mapping(
            method.get("evaluations"), context=f"original.{variant}.evaluations"
        )
        if set(evaluations) != set(ORIGINAL_SCORE_MODES):
            raise ResultsTDV2EMAError(f"Original {variant} evaluations are incomplete.")
        checkpoint_hashes: set[str] = set()
        for mode in ORIGINAL_SCORE_MODES:
            run = _mapping(evaluations[mode], context=f"original.{variant}.{mode}")
            successes = run.get("success_count")
            rate = run.get("success_rate")
            count = _parse_int(
                successes, context=f"original.{variant}.{mode}.success_count"
            )
            if not 0 <= count <= EPISODES:
                raise ResultsTDV2EMAError(
                    f"Original {variant}/{mode} count is invalid."
                )
            if not math.isclose(
                _finite_float(rate, context=f"original.{variant}.{mode}.success_rate"),
                count / EPISODES,
                rel_tol=0.0,
                abs_tol=1e-12,
            ):
                raise ResultsTDV2EMAError(f"Original {variant}/{mode} rate is invalid.")
            checkpoint_hashes.add(
                _sha256(
                    run.get("checkpoint_sha256"),
                    context=f"original.{variant}.{mode}.checkpoint_sha256",
                )
            )
        if len(checkpoint_hashes) != 1:
            raise ResultsTDV2EMAError(
                f"Original {variant} modes use different checkpoints."
            )
    return summary


def _validate_training_losses(
    path: Path,
    *,
    original_summary: Mapping[str, Any],
) -> tuple[Mapping[str, Any], ...]:
    raw_rows = _read_csv(path, fields=TRAINING_LOSS_FIELDS, context="training loss CSV")
    if len(raw_rows) != len(VARIANTS) * len(TRAINING_EPOCHS):
        raise ResultsTDV2EMAError("Training loss CSV must contain exactly 60 rows.")
    methods = _mapping(original_summary["methods"], context="original.methods")
    numeric_fields = TRAINING_LOSS_FIELDS[4:]
    summary_field_map = {
        "train_total_loss": "train_loss",
        "train_prediction_loss": "train_prediction_loss",
        "train_prediction_online_reference_mse": "train_prediction_online_reference_mse",
        "train_online_ema_latent_drift": "train_online_ema_latent_drift",
        "train_base_hybrid_td_loss": "train_base_hybrid_td",
        "train_method_hybrid_td_loss": "train_method_hybrid_td",
        "validation_total_loss": "validation_loss",
        "validation_prediction_loss": "validation_prediction_loss",
        "validation_prediction_online_reference_mse": "validation_prediction_online_reference_mse",
        "validation_online_ema_latent_drift": "validation_online_ema_latent_drift",
        "validation_base_hybrid_td_loss": "validation_base_hybrid_td",
        "validation_method_hybrid_td_loss": "validation_method_hybrid_td",
    }
    normalized: list[Mapping[str, Any]] = []
    identities: set[tuple[str, int]] = set()
    for index, raw in enumerate(raw_rows):
        context = f"training_loss[{index}]"
        variant = raw["variant"]
        if variant not in VARIANTS:
            raise ResultsTDV2EMAError(f"{context}.variant is invalid.")
        epoch = _parse_int(raw["epoch"], context=f"{context}.epoch")
        identity = (variant, epoch)
        if identity in identities or epoch not in TRAINING_EPOCHS:
            raise ResultsTDV2EMAError(f"{context} duplicates or has an invalid epoch.")
        identities.add(identity)
        method = _mapping(methods[variant], context=f"original.methods.{variant}")
        if (
            raw["method"] != method["method"]
            or raw["display_name"] != method["display_name"]
        ):
            raise ResultsTDV2EMAError(
                f"{context} method metadata differs from summary."
            )
        values = {
            field: _finite_float(raw[field], context=f"{context}.{field}")
            for field in numeric_fields
        }
        summary_curve = {
            int(row["epoch"]): _mapping(row, context=f"original.{variant}.curve")
            for row in method["training"]["loss_curve"]
        }
        summary_row = summary_curve[epoch]
        for csv_field, summary_field in summary_field_map.items():
            expected = _finite_float(
                summary_row.get(summary_field),
                context=f"original.{variant}.E{epoch}.{summary_field}",
            )
            if not math.isclose(
                values[csv_field], expected, rel_tol=5e-10, abs_tol=1e-8
            ):
                raise ResultsTDV2EMAError(
                    f"{context}.{csv_field} differs from original summary."
                )
        normalized.append(
            {
                "variant": variant,
                "method": raw["method"],
                "display_name": raw["display_name"],
                "epoch": epoch,
                **values,
            }
        )
    expected = {(variant, epoch) for variant in VARIANTS for epoch in TRAINING_EPOCHS}
    if identities != expected:
        raise ResultsTDV2EMAError("Training loss CSV grid is incomplete.")
    return tuple(normalized)


def _expected_fixed_grid() -> set[tuple[str, str, str]]:
    return {
        (version, variant, "f_plus_g_first")
        for version in FIXED_VERSIONS
        for variant in VARIANTS
    } | {("v2", variant, "g_only_f_rollout_mean") for variant in VARIANTS}


def _validate_new_summary(path: Path) -> Mapping[str, Any]:
    summary = _read_json(path, context="new-score summary")
    expected_header = {
        "schema_version": SCHEMA_VERSION,
        "source": "actor_free_td_lewm_v2_ema_new_score_summary",
        "cell_count": 96,
        "selection_sha256": EMA_SELECTION_SHA256,
        "training_commit": TRAINING_COMMIT,
        "evaluation_commit": EVALUATION_COMMIT,
        "fixed_checkpoint_cell_count": 24,
        "fixed_checkpoint_selection_sha256": FIXED_SELECTION_SHA256,
    }
    for key, expected in expected_header.items():
        if summary.get(key) != expected:
            raise ResultsTDV2EMAError(
                f"new_summary.{key} must equal {expected!r}, found {summary.get(key)!r}."
            )
    by_epoch = _mapping(
        summary.get("results_by_epoch"), context="new_summary.results_by_epoch"
    )
    if set(by_epoch) != {str(epoch) for epoch in EMA_EPOCHS}:
        raise ResultsTDV2EMAError("New-score summary must contain epochs 3-10 exactly.")
    for epoch in EMA_EPOCHS:
        methods = _mapping(by_epoch[str(epoch)], context=f"new_summary.E{epoch}")
        if set(methods) != set(VARIANTS):
            raise ResultsTDV2EMAError(f"New-score E{epoch} variant grid is incomplete.")
        for variant in VARIANTS:
            modes = _mapping(
                methods[variant], context=f"new_summary.E{epoch}.{variant}"
            )
            if set(modes) != set(SCORE_MODES):
                raise ResultsTDV2EMAError(
                    f"New-score E{epoch}/{variant} modes are incomplete."
                )
            for mode in SCORE_MODES:
                run = _mapping(
                    modes[mode], context=f"new_summary.E{epoch}.{variant}.{mode}"
                )
                if set(run) != {
                    "success_count",
                    "success_rate",
                    "success_rate_percent",
                }:
                    raise ResultsTDV2EMAError(
                        f"New-score E{epoch}/{variant}/{mode} has unexpected fields."
                    )
                _validate_rate(
                    run["success_count"],
                    run["success_rate"],
                    run["success_rate_percent"],
                    context=f"new_summary.E{epoch}.{variant}.{mode}",
                )

    best = _mapping(
        summary.get("best_by_epoch_and_score_mode"),
        context="new_summary.best_by_epoch_and_score_mode",
    )
    if set(best) != {str(epoch) for epoch in EMA_EPOCHS}:
        raise ResultsTDV2EMAError("Best-by-epoch summary keys are incomplete.")
    for epoch in EMA_EPOCHS:
        epoch_best = _mapping(best[str(epoch)], context=f"new_summary.best.E{epoch}")
        if set(epoch_best) != set(SCORE_MODES):
            raise ResultsTDV2EMAError(f"Best-by-epoch E{epoch} modes are incomplete.")
        for mode in SCORE_MODES:
            candidates = [
                (
                    variant,
                    int(by_epoch[str(epoch)][variant][mode]["success_count"]),
                )
                for variant in VARIANTS
            ]
            top = max(count for _, count in candidates)
            expected = {
                (variant, count, 100.0 * count / EPISODES)
                for variant, count in candidates
                if count == top
            }
            recorded: set[tuple[str, int, float]] = set()
            for index, value in enumerate(
                _sequence(epoch_best[mode], context=f"new_summary.best.E{epoch}.{mode}")
            ):
                row = _mapping(
                    value, context=f"new_summary.best.E{epoch}.{mode}[{index}]"
                )
                variant = row.get("variant")
                if variant not in VARIANTS:
                    raise ResultsTDV2EMAError(
                        "Best-by-epoch contains an invalid variant."
                    )
                count = _parse_int(
                    row.get("success_count"), context="best.success_count"
                )
                percent = _finite_float(
                    row.get("success_rate_percent"), context="best.success_rate_percent"
                )
                recorded.add((str(variant), count, percent))
            if recorded != expected:
                raise ResultsTDV2EMAError(f"Best-by-epoch E{epoch}/{mode} is stale.")

    fixed = _sequence(
        summary.get("fixed_checkpoint_results"),
        context="new_summary.fixed_checkpoint_results",
    )
    if len(fixed) != 24:
        raise ResultsTDV2EMAError("New-score summary must contain 24 fixed results.")
    fixed_grid: set[tuple[str, str, str]] = set()
    for index, value in enumerate(fixed):
        row = _mapping(value, context=f"new_summary.fixed[{index}]")
        if set(row) != {
            "version",
            "variant",
            "score_mode",
            "success_count",
            "success_rate",
            "success_rate_percent",
        }:
            raise ResultsTDV2EMAError("Fixed summary row fields changed.")
        identity = (str(row["version"]), str(row["variant"]), str(row["score_mode"]))
        if identity in fixed_grid:
            raise ResultsTDV2EMAError(f"Duplicate fixed summary cell: {identity}.")
        fixed_grid.add(identity)
        _validate_rate(
            row["success_count"],
            row["success_rate"],
            row["success_rate_percent"],
            context=f"new_summary.fixed[{index}]",
        )
    if fixed_grid != _expected_fixed_grid():
        raise ResultsTDV2EMAError("Fixed summary grid is not the exact 24 cells.")
    return summary


def _expected_cell_id(epoch: int, variant: str, mode: str) -> str:
    suffix = mode + ("_alpha_0p25" if mode == "f_plus_g_first" else "")
    return f"v2_ema_e{epoch:02d}_{variant}_{suffix}"


def _validate_new_results(
    path: Path,
    *,
    summary: Mapping[str, Any],
) -> tuple[Mapping[str, Any], ...]:
    raw_rows = _read_csv(
        path, fields=NEW_RESULTS_FIELDS, context="new-score results CSV"
    )
    if len(raw_rows) != 96:
        raise ResultsTDV2EMAError("New-score results CSV must contain exactly 96 rows.")
    identities: set[tuple[int, str, str]] = set()
    cell_ids: set[str] = set()
    normalized: list[Mapping[str, Any]] = []
    by_epoch = summary["results_by_epoch"]
    for index, raw in enumerate(raw_rows):
        context = f"new_results[{index}]"
        epoch = _parse_int(raw["epoch"], context=f"{context}.epoch")
        variant = raw["variant"]
        mode = raw["score_mode"]
        identity = (epoch, variant, mode)
        if (
            epoch not in EMA_EPOCHS
            or variant not in VARIANTS
            or mode not in SCORE_MODES
            or identity in identities
        ):
            raise ResultsTDV2EMAError(
                f"{context} has an invalid or duplicate identity."
            )
        identities.add(identity)
        if raw["method"] != f"actor_free_td_lewm_v2_ema_sg_{variant}":
            raise ResultsTDV2EMAError(f"{context}.method is inconsistent.")
        expected_weight = FIRST_ACTION_WEIGHT if mode == "f_plus_g_first" else None
        if expected_weight is None:
            if raw["g_first_weight"] != "":
                raise ResultsTDV2EMAError(f"{context} must not record g_first_weight.")
        elif not math.isclose(
            _finite_float(raw["g_first_weight"], context=f"{context}.g_first_weight"),
            expected_weight,
            rel_tol=0.0,
            abs_tol=1e-12,
        ):
            raise ResultsTDV2EMAError(f"{context}.g_first_weight must be 0.25.")
        count, rate, percent = _validate_rate(
            raw["success_count"],
            raw["success_rate"],
            raw["success_rate_percent"],
            context=context,
        )
        _exact_int(
            _parse_int(raw["checkpoint_epoch"], context=f"{context}.checkpoint_epoch"),
            epoch,
            context=f"{context}.checkpoint_epoch",
        )
        _exact_int(
            _parse_int(
                raw["checkpoint_global_step"],
                context=f"{context}.checkpoint_global_step",
            ),
            epoch * STEPS_PER_EPOCH,
            context=f"{context}.checkpoint_global_step",
        )
        checkpoint_sha = _sha256(
            raw["checkpoint_sha256"], context=f"{context}.checkpoint_sha256"
        )
        training_manifest_sha = _sha256(
            raw["training_manifest_sha256"],
            context=f"{context}.training_manifest_sha256",
        )
        if raw["evaluation_commit"] != EVALUATION_COMMIT:
            raise ResultsTDV2EMAError(f"{context}.evaluation_commit changed.")
        expected_scope = "strict_epoch3" if epoch == 3 else "original_epoch4_10"
        if raw["source_scope"] != expected_scope:
            raise ResultsTDV2EMAError(f"{context}.source_scope is inconsistent.")
        if raw["source_status"] not in {"REUSED", "SUCCEEDED"}:
            raise ResultsTDV2EMAError(f"{context}.source_status is not terminal.")
        expected_id = _expected_cell_id(epoch, variant, mode)
        if raw["cell_id"] != expected_id or expected_id in cell_ids:
            raise ResultsTDV2EMAError(f"{context}.cell_id is invalid or duplicate.")
        cell_ids.add(expected_id)
        summary_run = by_epoch[str(epoch)][variant][mode]
        if (
            count != summary_run["success_count"]
            or not math.isclose(
                rate, summary_run["success_rate"], rel_tol=0.0, abs_tol=1e-12
            )
            or not math.isclose(
                percent, summary_run["success_rate_percent"], rel_tol=0.0, abs_tol=1e-12
            )
        ):
            raise ResultsTDV2EMAError(f"{context} disagrees with new-score summary.")
        normalized.append(
            {
                "epoch": epoch,
                "variant": variant,
                "method": raw["method"],
                "score_mode": mode,
                "g_first_weight": expected_weight,
                "success_count": count,
                "success_rate": rate,
                "success_rate_percent": percent,
                "checkpoint_epoch": epoch,
                "checkpoint_global_step": epoch * STEPS_PER_EPOCH,
                "checkpoint_sha256": checkpoint_sha,
                "training_manifest_sha256": training_manifest_sha,
                "evaluation_commit": raw["evaluation_commit"],
                "source_scope": raw["source_scope"],
                "source_status": raw["source_status"],
                "cell_id": expected_id,
            }
        )
    expected_grid = {
        (epoch, variant, mode)
        for epoch in EMA_EPOCHS
        for variant in VARIANTS
        for mode in SCORE_MODES
    }
    if identities != expected_grid:
        raise ResultsTDV2EMAError("New-score results grid is incomplete.")
    for epoch in EMA_EPOCHS:
        for variant in VARIANTS:
            pair = [
                row
                for row in normalized
                if row["epoch"] == epoch and row["variant"] == variant
            ]
            if len({row["checkpoint_sha256"] for row in pair}) != 1:
                raise ResultsTDV2EMAError(
                    f"E{epoch}/{variant} modes use different checkpoints."
                )
            if len({row["training_manifest_sha256"] for row in pair}) != 1:
                raise ResultsTDV2EMAError(
                    f"E{epoch}/{variant} modes use different training manifests."
                )
    return tuple(normalized)


def _validate_fixed_results(
    path: Path,
    *,
    summary: Mapping[str, Any],
) -> tuple[Mapping[str, Any], ...]:
    raw_rows = _read_csv(
        path, fields=FIXED_RESULTS_FIELDS, context="fixed checkpoint results CSV"
    )
    if len(raw_rows) != 24:
        raise ResultsTDV2EMAError("Fixed checkpoint results CSV must contain 24 rows.")
    summary_rows = {
        (row["version"], row["variant"], row["score_mode"]): row
        for row in summary["fixed_checkpoint_results"]
    }
    identities: set[tuple[str, str, str]] = set()
    job_ids: set[str] = set()
    normalized: list[Mapping[str, Any]] = []
    for index, raw in enumerate(raw_rows):
        context = f"fixed_results[{index}]"
        identity = (raw["version"], raw["variant"], raw["score_mode"])
        if identity not in _expected_fixed_grid() or identity in identities:
            raise ResultsTDV2EMAError(
                f"{context} has an invalid or duplicate identity."
            )
        identities.add(identity)
        version, variant, mode = identity
        if raw["method"] != f"actor_free_td_lewm_{version}_{variant}":
            raise ResultsTDV2EMAError(f"{context}.method is inconsistent.")
        if mode == "f_plus_g_first":
            weight = _finite_float(
                raw["g_first_weight"], context=f"{context}.g_first_weight"
            )
            if not math.isclose(
                weight, FIRST_ACTION_WEIGHT, rel_tol=0.0, abs_tol=1e-12
            ):
                raise ResultsTDV2EMAError(f"{context}.g_first_weight must be 0.25.")
        else:
            weight = None
            if raw["g_first_weight"] != "":
                raise ResultsTDV2EMAError(f"{context} must not record g_first_weight.")
        count, rate, percent = _validate_rate(
            raw["success_count"],
            raw["success_rate"],
            raw["success_rate_percent"],
            context=context,
        )
        checkpoint_sha = _sha256(
            raw["checkpoint_sha256"], context=f"{context}.checkpoint_sha256"
        )
        if raw["selection_sha256"] != FIXED_SELECTION_SHA256:
            raise ResultsTDV2EMAError(
                f"{context} does not use the locked fixed-launcher ranks digest."
            )
        if raw["episode_selection_file_sha256"] != SHARED_EPISODE_SELECTION_SHA256:
            raise ResultsTDV2EMAError(
                f"{context} does not use the shared 50-pair episode selection."
            )
        if raw["action_normalization_sha256"] != ACTION_NORMALIZATION_SHA256:
            raise ResultsTDV2EMAError(
                f"{context} action-normalization evidence changed."
            )
        if raw["evaluation_commit"] != EVALUATION_COMMIT:
            raise ResultsTDV2EMAError(f"{context}.evaluation_commit changed.")
        job_id = _text(raw["job_id"], context=f"{context}.job_id")
        if job_id in job_ids:
            raise ResultsTDV2EMAError(f"Duplicate fixed job_id: {job_id}.")
        job_ids.add(job_id)
        source_manifest = _text(
            raw["source_launcher_manifest"],
            context=f"{context}.source_launcher_manifest",
        )
        summary_row = summary_rows[identity]
        if (
            count != summary_row["success_count"]
            or not math.isclose(
                rate, summary_row["success_rate"], rel_tol=0.0, abs_tol=1e-12
            )
            or not math.isclose(
                percent, summary_row["success_rate_percent"], rel_tol=0.0, abs_tol=1e-12
            )
        ):
            raise ResultsTDV2EMAError(f"{context} disagrees with new-score summary.")
        normalized.append(
            {
                "version": version,
                "variant": variant,
                "method": raw["method"],
                "score_mode": mode,
                "g_first_weight": weight,
                "success_count": count,
                "success_rate": rate,
                "success_rate_percent": percent,
                "checkpoint_sha256": checkpoint_sha,
                "selection_sha256": raw["selection_sha256"],
                "episode_selection_file_sha256": raw["episode_selection_file_sha256"],
                "action_normalization_sha256": raw["action_normalization_sha256"],
                "evaluation_commit": raw["evaluation_commit"],
                "job_id": job_id,
                "source_launcher_manifest": source_manifest,
            }
        )
    if identities != _expected_fixed_grid():
        raise ResultsTDV2EMAError("Fixed checkpoint CSV is not the exact 24-cell grid.")
    return tuple(normalized)


def _validate_base_document(path: Path) -> Path:
    if not path.is_file():
        raise FileNotFoundError(f"Base Results TD DOCX does not exist: {path}")
    if path.suffix.lower() != ".docx":
        raise ResultsTDV2EMAError("Base Results TD document must be a DOCX.")
    payload = path.read_bytes()
    if not payload.startswith(b"PK"):
        raise ResultsTDV2EMAError("Base Results TD document is not OOXML.")
    digest = hashlib.sha256(payload).hexdigest()
    if digest != BASE_DOCUMENT_SHA256:
        raise ResultsTDV2EMAError(
            "Base Results TD document differs from the locked V0/V1 report "
            f"(expected {BASE_DOCUMENT_SHA256}, found {digest})."
        )
    return path


def load_validated_report_inputs(
    *,
    original_summary_path: str | Path,
    training_loss_csv_path: str | Path,
    new_summary_path: str | Path,
    new_results_csv_path: str | Path,
    fixed_results_csv_path: str | Path,
    base_document_path: str | Path,
) -> ValidatedReportInputs:
    """Validate the complete report input set without writing output."""

    original = _validate_original_summary(Path(original_summary_path))
    losses = _validate_training_losses(
        Path(training_loss_csv_path), original_summary=original
    )
    new_summary = _validate_new_summary(Path(new_summary_path))
    new_rows = _validate_new_results(Path(new_results_csv_path), summary=new_summary)
    fixed_rows = _validate_fixed_results(
        Path(fixed_results_csv_path), summary=new_summary
    )
    base = _validate_base_document(Path(base_document_path))
    return ValidatedReportInputs(
        original_summary=original,
        training_loss_rows=losses,
        new_summary=new_summary,
        new_rows=new_rows,
        fixed_rows=fixed_rows,
        base_document=base,
    )


def _result_text(row: Mapping[str, Any]) -> str:
    return f"{int(row['success_count'])}/{EPISODES} ({float(row['success_rate_percent']):.0f}%)"


def _row_lookup(
    rows: Iterable[Mapping[str, Any]],
    **identity: Any,
) -> Mapping[str, Any]:
    matches = [
        row
        for row in rows
        if all(row.get(key) == value for key, value in identity.items())
    ]
    if len(matches) != 1:
        raise ResultsTDV2EMAError(
            f"Expected one report row for {identity}, found {len(matches)}."
        )
    return matches[0]


def _best_variant_epochs(
    inputs: ValidatedReportInputs,
) -> tuple[Mapping[str, Any], ...]:
    output: list[Mapping[str, Any]] = []
    for variant in VARIANTS:
        for mode in SCORE_MODES:
            candidates = [
                row
                for row in inputs.new_rows
                if row["variant"] == variant and row["score_mode"] == mode
            ]
            best_count = max(int(row["success_count"]) for row in candidates)
            best_epochs = tuple(
                int(row["epoch"])
                for row in candidates
                if int(row["success_count"]) == best_count
            )
            epoch10 = _row_lookup(candidates, epoch=10)
            output.append(
                {
                    "variant": variant,
                    "score_mode": mode,
                    "best_epochs": best_epochs,
                    "best_count": best_count,
                    "best_percent": best_count * 2.0,
                    "epoch10_count": int(epoch10["success_count"]),
                    "epoch10_percent": float(epoch10["success_rate_percent"]),
                }
            )
    return tuple(output)


def _mode_display(mode: str) -> str:
    return {
        "f_only": "F-only",
        "g_only": "G-only",
        "f_plus_g": "F+G tail",
        "f_plus_g_first": "F + first-Q (α=.25)",
        "g_only_f_rollout_mean": "Mean-Q over F rollout",
    }[mode]


METHOD_FORMULA_ROWS = (
    (
        "C",
        "q_i^b = G_phi(s_i^b,e_i,m_i)^T m_i; q_i^Y = Y_i^T m_i",
        "L_C^b = mean(l_i^b) + lambda_C mean_goal[(q_i^b-q_i^Y)^2], lambda_C=1",
        "Only C adds a trainable scalar projection residual on goal-derived tasks.",
    ),
    (
        "D",
        "A_i = sg(Y_i^T m_i)",
        "L_D^b = mean_i[w_i(A) l_i^b]",
        "Detached target goal value reweights TD; tau=0.5.",
    ),
    (
        "F",
        "A_i = sg[Y_i^T m_i - mean_j(Y_i^T m_j)]",
        "L_F^b = mean_i[w_i(A) l_i^b]",
        "Matched future/task is contrasted with all goal tasks in the batch.",
    ),
    (
        "G1",
        "A_i = sg[q_i - sum_k softmax(-d_ik/tau_n) q_ik]",
        "L_G1^b = mean_i[w_i(A) l_i^b]",
        "K=8 other-episode latent-neighbour actions; candidates have no TD targets.",
    ),
    (
        "G2",
        "A_i = sg[q_i5 - (1/5) sum_{j=1}^5 q_ij]",
        "L_G2^b = mean_i[w_i(A) l_i^b]",
        "Five zero-suffix action prefixes; full-minus-prefix-mean signal.",
    ),
    (
        "G3",
        "A_i = sg[(1/4) sum_{j=1}^4(q_i,j+1-q_ij)]",
        "L_G3^b = mean_i[w_i(A) l_i^b]",
        "Five prefixes; mean adjacent marginal score gain.",
    ),
)


CHART_COLORS = ("#0B6E99", "#D95F02", "#1B9E77", "#7570B3", "#E7298A", "#555555")


def _chart_font(size: int, *, bold: bool = False) -> Any:
    """Load a deterministic local font for Pillow chart rendering."""

    from PIL import ImageFont

    candidates = (
        Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
        if bold
        else Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        if bold
        else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _draw_line_panel(
    draw: Any,
    *,
    box: tuple[int, int, int, int],
    title: str,
    x_values: Sequence[int],
    series: Sequence[tuple[str, Sequence[float], str]],
    y_min: float,
    y_max: float,
    y_ticks: Sequence[tuple[float, str]],
    y_transform: Any,
    y_label: str,
) -> None:
    """Draw one compact, explicit-geometry line panel with Pillow."""

    left, top, right, bottom = box
    plot_left = left + 92
    plot_top = top + 105
    plot_right = right - 28
    plot_bottom = bottom - 68
    draw.rounded_rectangle(box, radius=18, fill="#FAFBFC", outline="#D8DEE6", width=2)
    title_font = _chart_font(27, bold=True)
    label_font = _chart_font(20)
    tick_font = _chart_font(18)
    legend_font = _chart_font(18, bold=True)
    draw.text((left + 24, top + 18), title, fill="#0B2545", font=title_font)
    draw.text((plot_left, top + 63), y_label, fill="#5C6975", font=label_font)

    transformed_min = float(y_transform(y_min))
    transformed_max = float(y_transform(y_max))
    if not transformed_max > transformed_min:
        raise ResultsTDV2EMAError("Chart y-axis bounds must be increasing.")

    def x_position(value: int) -> float:
        return plot_left + (value - x_values[0]) * (plot_right - plot_left) / (
            x_values[-1] - x_values[0]
        )

    def y_position(value: float) -> float:
        transformed = float(y_transform(value))
        return plot_bottom - (transformed - transformed_min) * (
            plot_bottom - plot_top
        ) / (transformed_max - transformed_min)

    for tick_value, tick_label in y_ticks:
        if not y_min <= tick_value <= y_max:
            continue
        y = y_position(tick_value)
        draw.line((plot_left, y, plot_right, y), fill="#DDE3EA", width=2)
        label_box = draw.textbbox((0, 0), tick_label, font=tick_font)
        label_width = label_box[2] - label_box[0]
        draw.text(
            (plot_left - label_width - 12, y - 10),
            tick_label,
            fill="#5C6975",
            font=tick_font,
        )
    for value in x_values:
        x = x_position(value)
        draw.line((x, plot_top, x, plot_bottom), fill="#EEF1F5", width=1)
        label = str(value)
        label_box = draw.textbbox((0, 0), label, font=tick_font)
        label_width = label_box[2] - label_box[0]
        draw.text(
            (x - label_width / 2, plot_bottom + 13),
            label,
            fill="#5C6975",
            font=tick_font,
        )
    draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#7B8794", width=2)
    draw.line(
        (plot_left, plot_bottom, plot_right, plot_bottom), fill="#7B8794", width=2
    )

    legend_x = plot_right - 395
    legend_y = top + 65
    for index, (label, values, color) in enumerate(series):
        column = index % 3
        row = index // 3
        x = legend_x + column * 130
        y = legend_y + row * 28
        draw.line((x, y + 9, x + 28, y + 9), fill=color, width=5)
        draw.ellipse((x + 10, y + 3, x + 22, y + 15), fill=color, outline="white")
        draw.text((x + 36, y), label, fill="#26313D", font=legend_font)

        points = [
            (x_position(x_value), y_position(float(y_value)))
            for x_value, y_value in zip(x_values, values)
        ]
        draw.line(points, fill=color, width=5, joint="curve")
        for point_x, point_y in points:
            draw.ellipse(
                (point_x - 6, point_y - 6, point_x + 6, point_y + 6),
                fill=color,
                outline="white",
                width=2,
            )
    x_title = "Epoch" if x_values[0] == 1 else "Checkpoint epoch"
    x_box = draw.textbbox((0, 0), x_title, font=label_font)
    draw.text(
        (
            (plot_left + plot_right - (x_box[2] - x_box[0])) / 2,
            bottom - 34,
        ),
        x_title,
        fill="#5C6975",
        font=label_font,
    )


def _save_chart(image: Any) -> bytes:
    stream = io.BytesIO()
    image.save(stream, format="PNG", dpi=(180, 180), optimize=True)
    payload = stream.getvalue()
    if not payload.startswith(PNG_SIGNATURE):
        raise RuntimeError("Pillow did not produce a PNG chart.")
    return payload


def build_training_loss_chart(inputs: ValidatedReportInputs) -> bytes:
    """Render accepted training/validation total-loss curves as PNG bytes."""

    from PIL import Image, ImageDraw

    image = Image.new("RGB", (2160, 860), "white")
    draw = ImageDraw.Draw(image)
    title_font = _chart_font(34, bold=True)
    draw.text(
        (70, 28),
        "V2-EMA-SG training and validation loss",
        fill="#0B2545",
        font=title_font,
    )
    train_series = []
    validation_series = []
    all_values: list[float] = []
    for variant, color in zip(VARIANTS, CHART_COLORS):
        rows = sorted(
            (row for row in inputs.training_loss_rows if row["variant"] == variant),
            key=lambda row: int(row["epoch"]),
        )
        train_values = [float(row["train_total_loss"]) for row in rows]
        validation_values = [float(row["validation_total_loss"]) for row in rows]
        if any(value <= 0 for value in train_values + validation_values):
            raise ResultsTDV2EMAError("Log-scale loss chart requires positive losses.")
        train_series.append((variant.upper(), train_values, color))
        validation_series.append((variant.upper(), validation_values, color))
        all_values.extend(train_values + validation_values)
    log_min = math.floor(min(math.log10(value) for value in all_values) * 2) / 2
    log_max = math.ceil(max(math.log10(value) for value in all_values) * 2) / 2
    y_min = 10**log_min
    y_max = 10**log_max
    exponent_start = math.ceil(log_min)
    exponent_end = math.floor(log_max)
    ticks = tuple(
        (10.0**power, f"1e{power}") for power in range(exponent_start, exponent_end + 1)
    )
    if not ticks:
        ticks = ((y_min, f"{y_min:.2g}"), (y_max, f"{y_max:.2g}"))
    _draw_line_panel(
        draw,
        box=(55, 90, 1060, 825),
        title="Training total loss",
        x_values=TRAINING_EPOCHS,
        series=train_series,
        y_min=y_min,
        y_max=y_max,
        y_ticks=ticks,
        y_transform=math.log10,
        y_label="Loss (log scale)",
    )
    _draw_line_panel(
        draw,
        box=(1100, 90, 2105, 825),
        title="Validation total loss",
        x_values=TRAINING_EPOCHS,
        series=validation_series,
        y_min=y_min,
        y_max=y_max,
        y_ticks=ticks,
        y_transform=math.log10,
        y_label="Loss (log scale)",
    )
    return _save_chart(image)


def build_new_score_chart(inputs: ValidatedReportInputs) -> bytes:
    """Render success rate versus checkpoint epoch for both new score modes."""

    from PIL import Image, ImageDraw

    image = Image.new("RGB", (2160, 860), "white")
    draw = ImageDraw.Draw(image)
    title_font = _chart_font(34, bold=True)
    draw.text(
        (70, 28),
        "V2-EMA-SG new inference scores across epochs 3-10",
        fill="#0B2545",
        font=title_font,
    )
    for box, mode in zip(((55, 90, 1060, 825), (1100, 90, 2105, 825)), SCORE_MODES):
        series = []
        for variant, color in zip(VARIANTS, CHART_COLORS):
            rows = sorted(
                (
                    row
                    for row in inputs.new_rows
                    if row["variant"] == variant and row["score_mode"] == mode
                ),
                key=lambda row: int(row["epoch"]),
            )
            series.append(
                (
                    variant.upper(),
                    [float(row["success_rate_percent"]) for row in rows],
                    color,
                )
            )
        _draw_line_panel(
            draw,
            box=box,
            title=_mode_display(mode),
            x_values=EMA_EPOCHS,
            series=series,
            y_min=0.0,
            y_max=100.0,
            y_ticks=tuple((float(value), str(value)) for value in range(0, 101, 20)),
            y_transform=float,
            y_label="O50 success rate (%)",
        )
    return _save_chart(image)


def _epoch10_rows(inputs: ValidatedReportInputs) -> tuple[tuple[str, ...], ...]:
    methods = inputs.original_summary["methods"]
    rows: list[tuple[str, ...]] = []
    for variant in VARIANTS:
        original = methods[variant]["evaluations"]
        first = _row_lookup(
            inputs.new_rows,
            epoch=10,
            variant=variant,
            score_mode="f_plus_g_first",
        )
        mean_q = _row_lookup(
            inputs.new_rows,
            epoch=10,
            variant=variant,
            score_mode="g_only_f_rollout_mean",
        )
        rows.append(
            (
                variant.upper(),
                f"{original['f_only']['success_count']}/50 ({original['f_only']['success_rate'] * 100:.0f}%)",
                f"{original['g_only']['success_count']}/50 ({original['g_only']['success_rate'] * 100:.0f}%)",
                f"{original['f_plus_g']['success_count']}/50 ({original['f_plus_g']['success_rate'] * 100:.0f}%)",
                _result_text(first),
                _result_text(mean_q),
            )
        )
    return tuple(rows)


def _fixed_table_rows(inputs: ValidatedReportInputs) -> tuple[tuple[str, ...], ...]:
    rows: list[tuple[str, ...]] = []
    for version in FIXED_VERSIONS:
        for variant in VARIANTS:
            first = _row_lookup(
                inputs.fixed_rows,
                version=version,
                variant=variant,
                score_mode="f_plus_g_first",
            )
            rows.append(
                (
                    version.upper(),
                    variant.upper(),
                    _mode_display("f_plus_g_first"),
                    _result_text(first),
                    first["checkpoint_sha256"][:12] + "…",
                )
            )
    for variant in VARIANTS:
        mean_q = _row_lookup(
            inputs.fixed_rows,
            version="v2",
            variant=variant,
            score_mode="g_only_f_rollout_mean",
        )
        rows.append(
            (
                "V2",
                variant.upper(),
                _mode_display("g_only_f_rollout_mean"),
                _result_text(mean_q),
                mean_q["checkpoint_sha256"][:12] + "…",
            )
        )
    return tuple(rows)


def _fixed_matrix_rows(inputs: ValidatedReportInputs) -> tuple[tuple[str, ...], ...]:
    """Return the 24 fixed cells as a compact six-row comparison matrix."""

    rows: list[tuple[str, ...]] = []
    for variant in VARIANTS:
        first_by_version = {
            version: _row_lookup(
                inputs.fixed_rows,
                version=version,
                variant=variant,
                score_mode="f_plus_g_first",
            )
            for version in FIXED_VERSIONS
        }
        mean_q = _row_lookup(
            inputs.fixed_rows,
            version="v2",
            variant=variant,
            score_mode="g_only_f_rollout_mean",
        )
        rows.append(
            (
                variant.upper(),
                _result_text(first_by_version["v0"]),
                _result_text(first_by_version["v1"]),
                _result_text(first_by_version["v2"]),
                _result_text(mean_q),
            )
        )
    return tuple(rows)


def _best_epoch_table_rows(
    inputs: ValidatedReportInputs,
) -> tuple[tuple[str, ...], ...]:
    rows = []
    for value in _best_variant_epochs(inputs):
        epoch_label = "/".join(f"E{epoch}" for epoch in value["best_epochs"])
        rows.append(
            (
                value["variant"].upper(),
                _mode_display(value["score_mode"]),
                epoch_label,
                f"{value['best_count']}/50 ({value['best_percent']:.0f}%)",
                f"{value['epoch10_count']}/50 ({value['epoch10_percent']:.0f}%)",
                f"{value['best_count'] - value['epoch10_count']:+d}/50",
            )
        )
    return tuple(rows)


def _epoch10_winners(inputs: ValidatedReportInputs, mode: str) -> tuple[str, int]:
    rows = [
        row
        for row in inputs.new_rows
        if row["epoch"] == 10 and row["score_mode"] == mode
    ]
    best = max(int(row["success_count"]) for row in rows)
    labels = "/".join(
        row["variant"].upper() for row in rows if int(row["success_count"]) == best
    )
    return labels, best


def _fixed_winners(
    inputs: ValidatedReportInputs, version: str, mode: str
) -> tuple[str, int]:
    rows = [
        row
        for row in inputs.fixed_rows
        if row["version"] == version and row["score_mode"] == mode
    ]
    best = max(int(row["success_count"]) for row in rows)
    labels = "/".join(
        row["variant"].upper() for row in rows if int(row["success_count"]) == best
    )
    return labels, best


def _mean_success_percent(rows: Iterable[Mapping[str, Any]]) -> float:
    selected = tuple(rows)
    if not selected:
        raise ResultsTDV2EMAError("Cannot average an empty result selection.")
    return (
        100.0
        * sum(int(row["success_count"]) for row in selected)
        / (EPISODES * len(selected))
    )


def build_markdown_report(
    inputs: ValidatedReportInputs,
    *,
    training_chart_reference: str,
    score_chart_reference: str,
) -> str:
    """Build the complete Chinese Markdown report from validated inputs."""

    lines: list[str] = [
        "# Results TD — V2-EMA-SG 新推理评分（Cube seed 3072）",
        "",
        "本报告只接受严格归档后的 **96 个 V2-EMA checkpoint×评分单元**和 **24 个固定 V0/V1/V2 checkpoint 单元**。所有结果均为 O50；每格 50 个 start-goal pair。模型没有 Actor，也没有 reward loss。",
        "",
        "## 网络、TD target 与总训练损失",
        "",
        "在线 LeWM 产生真实 latent `z_t=E_θ(o_t)` 和可微预测 latent `ẑ_t=F_θ(H_{t-1},E^A_θ(a_{t-1}))`；同一个 TD-JEPA predictor `G_φ` 分别接收两条支路：",
        "",
        "$$s_t^{real}=z_t,\\qquad s_t^{pred}=\\hat z_t.$$",
        "",
        "两条支路共享一个完全 stop-gradient 的 EMA target：",
        "",
        "$$Y_t=\\operatorname{sg}\\left[\\bar z_{t+1}+\\gamma(1-d_t)G_{\\bar\\phi}(\\bar z_{t+1},\\bar e_{t+1},m_t)\\right],\\qquad \\gamma=0.95.$$",
        "",
        "其中 `z̄_{t+1}=E_{θ̄}(o_{t+1})` 来自真实下一帧的 EMA encoder，`ē_{t+1}=E^A_{θ̄}(a_{t+1})` 来自数据集已知 next action。它们不是 online LeWM 想象出来的下一状态/动作。基础 TD 误差为：",
        "",
        "$$\\ell_{t,b}^{TD}=\\left\\|G_\\phi(s_t^b,e_t,m_t)-Y_t\\right\\|_2^2,\\qquad b\\in\\{real,pred\\}.$$",
        "",
        "V2-EMA-SG 的 LeWM 一步预测本身也改用真实下一帧经过 EMA encoder 的 latent 作为 stop-gradient target：",
        "",
        "$$L_{pred}=\\operatorname{MSE}(F_\\theta(H_t,e_t),\\operatorname{sg}(\\bar z_{t+1})).$$",
        "",
        "完整训练目标是：",
        "",
        "$$L_{total}=L_{pred}+0.09L_{SIGReg}+\\rho(u)\\left(L_{method}^{real}+L_{method}^{pred}\\right),$$",
        "",
        "`ρ(u)` 在前 5% optimizer updates 从 0 线性升到 1。predicted-state 支路不 detach，因此 TD 梯度会回传到 online LeWM predictor 和共享 Action Encoder；EMA world model 与 EMA `G` 只生成 target。validation 使用共同的、未加方法特定权重的 base Hybrid TD，因此 total loss 高低不能直接跨方法排序。",
        "",
        "目标任务子集的统一权重为 `w_i=N_g softmax(A_i/τ)`（`τ=0.5`，stop-gradient，随机任务权重固定为 1，最后全 batch 均值归一为 1）。六个方法只改变下面的特殊信号/额外项：",
        "",
        "| 方法 | 特殊信号 | 每条 real/pred 支路的 loss | 作用 |",
        "| --- | --- | --- | --- |",
    ]
    for method, signal, loss, effect in METHOD_FORMULA_ROWS:
        lines.append(f"| {method} | `{signal}` | `{loss}` | {effect} |")

    lines.extend(
        [
            "",
            "## 两个新增推理评分",
            "",
            "令 `w(g)` 为 goal latent 归一化后的任务向量，`Q_G(z,A,g)=G(z,E^A(A),w(g))^T w(g)`。CEM 仍然只负责搜索候选动作序列、最小化 cost、执行第一块动作并重新规划；这里没有训练 Actor。",
            "",
            "### 1. F + first-Q（`f_plus_g_first`）",
            "",
            "$$J_{first}(A_{1:5})=\\|\\hat z_5-z_g\\|_2^2-0.25\\,Q_G(z_0,A_1,g).$$",
            "",
            "F 完整 rollout 五个 action blocks；`Q_G` 只评价当前真实 online latent `z_0` 与第一块候选动作 `A_1`。**没有 `γ^4` tail 折扣**。",
            "",
            "### 2. F rollout 上的 Mean-Q（`g_only_f_rollout_mean`）",
            "",
            "$$J_{mean}(A_{1:5})=-\\frac{1}{5}\\sum_{k=1}^{5}Q_G(z_{k-1}^{F},A_k,g),$$",
            "",
            "其中 `z_0^F=z_0`，`z_1^F,…,z_4^F` 由 online LeWM 的完整 F rollout 产生。F 在这里仅生成 imagined states；不使用 terminal goal-distance，`γ` 也不参与评分。",
            "",
            "## V2-EMA epoch 10：原三评分 + 两个新评分",
            "",
            "这五列绑定同一组 epoch-10 checkpoint，并使用共享正式 episode selection `e46ea81c…`，可以作为同一 50 pair set 上的推理评分消融。",
            "",
            "| 方法 | F-only | G-only | F+G tail | F + first-Q | Mean-Q rollout |",
            "| --- | ---: | ---: | ---: | ---: | ---: |",
        ]
    )
    for row in _epoch10_rows(inputs):
        lines.append("| " + " | ".join(row) + " |")

    lines.extend(
        [
            "",
            "## 固定 V0/V1/V2 checkpoint：完整 24 格",
            "",
            "固定 checkpoint 表与 EMA sweep **使用同一个 episode-selection.json（`e46ea81c…`）**，因此可以在同一 O50 pair set 上比较成功率。`88c20477…` 只是 fixed launcher 的 canonical valid-row-ranks digest，不代表另一组 pair。当前报告输入没有逐格 outcome 向量，因此不做逐-pair 显著性检验。V0/V1 只支持 first-Q；Mean-Q rollout 是 V2-only。",
            "",
            "| Version | 方法 | 评分 | O50 | Checkpoint SHA（短） |",
            "| --- | --- | --- | ---: | --- |",
        ]
    )
    for row in _fixed_table_rows(inputs):
        lines.append("| " + " | ".join(row) + " |")

    lines.extend(
        [
            "",
            "## V2-EMA 每个方法×评分的最佳 epoch（事后分析）",
            "",
            "> 下表是在 epoch 3–10 看完结果后事后选择 checkpoint，存在 selection bias。它只能用于理解训练轨迹和提出下一轮预注册 checkpoint 规则，不能当作无偏最终性能。正式 checkpoint 口径仍应单独报告 epoch 10。",
            "",
            "| 方法 | 评分 | 最佳 epoch | 最佳 O50 | Epoch 10 | 最佳−E10 |",
            "| --- | --- | --- | ---: | ---: | ---: |",
        ]
    )
    for row in _best_epoch_table_rows(inputs):
        lines.append("| " + " | ".join(row) + " |")

    first_labels, first_count = _epoch10_winners(inputs, "f_plus_g_first")
    mean_labels, mean_count = _epoch10_winners(inputs, "g_only_f_rollout_mean")
    fixed_v2_labels, fixed_v2_count = _fixed_winners(inputs, "v2", "f_plus_g_first")
    ema_first_mean = _mean_success_percent(
        row
        for row in inputs.new_rows
        if row["epoch"] == 10 and row["score_mode"] == "f_plus_g_first"
    )
    ema_rollout_mean = _mean_success_percent(
        row
        for row in inputs.new_rows
        if row["epoch"] == 10 and row["score_mode"] == "g_only_f_rollout_mean"
    )
    fixed_first_means = {
        version: _mean_success_percent(
            row
            for row in inputs.fixed_rows
            if row["version"] == version and row["score_mode"] == "f_plus_g_first"
        )
        for version in FIXED_VERSIONS
    }
    lines.extend(
        [
            "",
            "## 曲线",
            "",
            "训练 total loss 是方法特定目标；validation total loss 使用共同 base Hybrid TD，但仍包含 LeWM prediction/SIGReg。先看各方法自身的收敛趋势，不按绝对高度跨方法排名。",
            "",
            f"![V2-EMA-SG training and validation loss curves]({training_chart_reference})",
            "",
            "新评分曲线覆盖 epoch 3–10；每一个点都是同一 EMA 50-pair selection 的完整 O50。",
            "",
            f"![V2-EMA-SG new-score epoch curves]({score_chart_reference})",
            "",
            "## 关键结论与边界",
            "",
            f"- Epoch 10 的 `F + first-Q` 最好为 **{first_labels}: {first_count}/50 ({first_count * 2}%)**；Mean-Q rollout 最好为 **{mean_labels}: {mean_count}/50 ({mean_count * 2}%)**。",
            f"- Epoch 10 六方法均值：first-Q **{ema_first_mean:.1f}%**，Mean-Q **{ema_rollout_mean:.1f}%**。固定 checkpoint 的 first-Q 均值为 V0 **{fixed_first_means['v0']:.1f}%**、V1 **{fixed_first_means['v1']:.1f}%**、V2 **{fixed_first_means['v2']:.1f}%**；因此这组固定对照中 V1 最强。",
            f"- 固定 V2 checkpoint 的 first-Q 表中最好为 **{fixed_v2_labels}: {fixed_v2_count}/50 ({fixed_v2_count * 2}%)**；它与 EMA 表使用相同 O50 pair set，可以比较成功率，但本报告不声称逐-pair 统计显著性。",
            "- first-Q 只把 G 当作第一动作的 critic/readout；Mean-Q 则让 G 在 F rollout 的五个 predecessor-action 对上都参与排序。两者都不训练策略，也不会把 CEM 变成 Actor。",
            "- C/D/F/G1/G2/G3 的差异发生在训练期 TD loss 或其 detached 权重；推理时六者使用同一评分公式，只是 checkpoint 内学到的 `G` 不同。",
            "- 全部 120 格共享同一组 planning selection，但只有一个 training seed（3072）。96 格 epoch sweep 与最佳 epoch 表是结构/训练轨迹消融，不支持多随机种子总体最优或统计显著性结论。",
            "",
            "## 审计信息",
            "",
            f"- Training commit: `{TRAINING_COMMIT}`",
            f"- Evaluation commit: `{EVALUATION_COMMIT}`",
            f"- Shared 120-cell episode-selection file SHA-256: `{SHARED_EPISODE_SELECTION_SHA256}`",
            f"- Fixed-launcher canonical valid-row-ranks SHA-256: `{FIXED_LAUNCHER_RANKS_SHA256}`",
            f"- Action-normalization SHA-256: `{ACTION_NORMALIZATION_SHA256}`",
            "- EMA grid: epochs 3–10 × C/D/F/G1/G2/G3 × 2 scores = 96；fixed grid: 18 first-Q + 6 V2 Mean-Q = 24。",
            "",
        ]
    )
    return "\n".join(lines)


def _set_header_text(section: Any, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        _v1._set_run_font(run, size=8.5, color="6B7280")


def _font_paragraph(
    paragraph: Any, *, size: float = 10.5, color: str = "111827"
) -> Any:
    for run in paragraph.runs:
        _v1._set_run_font(run, size=size, color=color, bold=bool(run.bold))
    return paragraph


def _add_body(
    document: Any, text: str, *, color: str = "111827", bold: bool = False
) -> Any:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    _v1._set_run_font(run, size=10.5, color=color, bold=bold)
    return paragraph


def _add_formula(document: Any, label: str, formula: str, explanation: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(label + "  ")
    _v1._set_run_font(label_run, size=10.5, color="0B2545", bold=True)
    formula_run = paragraph.add_run(formula)
    _v1._set_run_font(formula_run, size=10.5, color="111827")
    detail = document.add_paragraph()
    detail.paragraph_format.left_indent = Pt(14)
    detail.paragraph_format.space_after = Pt(7)
    detail_run = detail.add_run(explanation)
    _v1._set_run_font(detail_run, size=9.8, color="4B5563")


def _add_heading(
    document: Any, text: str, *, level: int = 1, page_break: bool = False
) -> Any:
    from docx.shared import Pt

    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.page_break_before = page_break
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        _v1._set_run_font(
            run,
            size=16 if level == 1 else 12.5,
            color="0B2545" if level == 1 else "174A6E",
            bold=True,
        )
    return paragraph


def _add_picture(
    document: Any, payload: bytes, *, title: str, description: str
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(io.BytesIO(payload), width=Inches(9.2))
    shape._inline.docPr.set("title", title)
    shape._inline.docPr.set("descr", description)
    caption = document.add_paragraph(title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        _v1._set_run_font(run, size=9, color="4B5563")
        run.italic = True


def build_results_document(
    inputs: ValidatedReportInputs,
    *,
    training_chart_png: bytes,
    score_chart_png: bytes,
) -> bytes:
    """Append the validated V2-EMA-SG report and return DOCX bytes."""

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required; use the Codex workspace document runtime."
        ) from exc
    for name, payload in (
        ("training chart", training_chart_png),
        ("new-score chart", score_chart_png),
    ):
        if not payload.startswith(PNG_SIGNATURE):
            raise ResultsTDV2EMAError(f"{name} is not a valid PNG.")

    document = Document(str(inputs.base_document))
    _v1._configure_append_section(document)
    section = document.sections[-1]
    _set_header_text(
        section,
        "TD-JEPA · V2-EMA-SG · Cube O50 · 96 epoch-score cells + 24 fixed cells",
    )

    kicker = document.add_paragraph(style="Report Kicker")
    kicker.add_run("RESULTS TD / V2-EMA-SG NEW INFERENCE SCORES")
    _font_paragraph(kicker, size=9.5, color="5C6975")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("V2-EMA-SG new scores and checkpoint trajectory")
    _v1._set_run_font(run, size=24, color="0B2545", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Cube · seed 3072 · actor-free CEM · 96 EMA cells + 24 fixed-checkpoint cells"
    )
    _v1._set_run_font(run, size=12, color="4B5563")
    _add_body(
        document,
        "Interpretation boundary: one training seed; the EMA sweep and fixed-"
        "checkpoint study share one O50 episode selection. Best epoch is a "
        "post-hoc trajectory analysis, not an unbiased final-performance estimate; "
        "this report does not claim pairwise statistical significance.",
        color="7A5A00",
        bold=True,
    )

    _add_heading(document, "Network, TD target and total loss")
    _add_body(
        document,
        "Online LeWM supplies the real latent z_t and a differentiable predicted "
        "latent z-hat_t. One G_phi is trained on both real and predicted branches. "
        "The unique shared, trainable LeWM Action Encoder maps each raw 25D action "
        "block to 192D. There is no Actor, reward model or action loss.",
    )
    _add_formula(
        document,
        "TD target",
        "Y_t = sg[z-bar_(t+1) + gamma(1-d_t) G-bar(z-bar_(t+1), e-bar_(t+1), m_t)], gamma=0.95",
        "z-bar_(t+1) comes from the real next frame through the EMA encoder; "
        "e-bar_(t+1) comes from the known dataset next action through the EMA "
        "Action Encoder.",
    )
    _add_formula(
        document,
        "Branch TD",
        "l_(t,b) = ||G_phi(s_t^b,e_t,m_t)-Y_t||_2^2,  b in {real,pred}",
        "s_real=z_t and s_pred=z-hat_t. The predicted branch retains gradients to "
        "the online LeWM predictor and shared Action Encoder.",
    )
    _add_formula(
        document,
        "Total",
        "L_total = L_pred + 0.09 L_SIGReg + rho(u)[L_method^real + L_method^pred]",
        "L_pred uses the real next-frame EMA latent as a stop-gradient target; "
        "rho rises linearly from 0 to 1 over the first 5% of optimizer updates.",
    )

    _add_heading(document, "C / D / F / G1 / G2 / G3 training terms", page_break=True)
    _add_body(
        document,
        "D/F/G1/G2/G3 use detached signals only to reweight per-sample TD error. "
        "Goal-subset weights use tau=0.5 softmax and mean-one normalization. C "
        "also fits a goal-projection residual. Every method shares the same EMA "
        "Bellman target.",
    )
    _v1._add_table(
        document,
        headers=("Method", "Special signal A / q", "Per-branch loss", "Training role"),
        rows=METHOD_FORMULA_ROWS,
        widths=(1200, 4300, 4300, 4600),
    )

    _add_heading(document, "Two new inference scores", page_break=True)
    _add_body(
        document,
        "Q_G(z,A,g)=G(z,E_A(A),w(g))^T w(g). CEM minimizes the cost below, "
        "executes only A1, then replans. No Actor is trained.",
    )
    _add_formula(
        document,
        "F + first-Q",
        "J_first(A_1:5)=||z-hat_5-z_g||_2^2 - 0.25 Q_G(z_0,A_1,g)",
        "F performs the complete five-block rollout. first-Q has no gamma^4 tail "
        "discount and evaluates only real z_0 with the first candidate block A1.",
    )
    _add_formula(
        document,
        "Mean-Q rollout",
        "J_mean(A_1:5)=-(1/5) sum_(k=1)^5 Q_G(z_(k-1)^F,A_k,g)",
        "z_0 is the online encoder state; z_1 through z_4 come from the F rollout. "
        "F only generates imagined states; neither terminal goal distance nor "
        "gamma enters this score.",
    )

    _add_heading(document, "Epoch 10: five inference scores", page_break=True)
    _add_body(
        document,
        "All five columns bind the same V2-EMA epoch-10 checkpoint and shared "
        "formal episode selection (e46ea81c...). The first three columns come "
        "from the original formal archive; the last two come from the 96-cell sweep.",
    )
    _v1._add_table(
        document,
        headers=(
            "Method",
            "F-only",
            "G-only",
            "F+G tail",
            "F + first-Q",
            "Mean-Q rollout",
        ),
        rows=_epoch10_rows(inputs),
        widths=(1400, 2500, 2500, 2500, 2700, 2800),
    )

    _add_heading(document, "Fixed V0/V1/V2 checkpoints: all 24 cells", page_break=True)
    _add_body(
        document,
        "This table and the EMA sweep use the same episode-selection.json "
        "(e46ea81c...), so success rates are comparable on one O50 pair set. "
        "88c20477... is only the fixed launcher's canonical valid-row-ranks digest. "
        "The report input does not include per-cell outcome vectors, so no paired "
        "significance test is claimed. V0/V1 support first-Q only; Mean-Q is V2-only.",
        color="7A5A00",
        bold=True,
    )
    _v1._add_table(
        document,
        headers=("Method", "V0 first-Q", "V1 first-Q", "V2 first-Q", "V2 Mean-Q"),
        rows=_fixed_matrix_rows(inputs),
        widths=(1800, 3150, 3150, 3150, 3150),
    )

    _add_heading(document, "Best epoch by method and score (post-hoc)", page_break=True)
    _add_body(
        document,
        "Warning: each epoch-3-to-10 maximum was selected after observing the same "
        "O50 selection and is therefore selection-biased. Use it only to understand "
        "training trajectories or define a preregistered rule for the next study; "
        "epoch 10 remains the separately reported formal checkpoint.",
        color="9C2F17",
        bold=True,
    )
    _v1._add_table(
        document,
        headers=("Method", "Score", "Best epoch", "Best O50", "Epoch 10", "Best-E10"),
        rows=_best_epoch_table_rows(inputs),
        widths=(1200, 4300, 2200, 2500, 2400, 1800),
    )

    _add_heading(document, "Training and validation loss", page_break=True)
    _add_body(
        document,
        "Training total loss is method-specific. Validation uses the shared base "
        "Hybrid TD objective, although total loss still includes prediction/SIGReg. "
        "Use the curves to inspect within-method trajectories, not to rank methods "
        "by absolute loss height.",
    )
    _add_picture(
        document,
        training_chart_png,
        title="Figure 1. V2-EMA-SG training / validation total loss (epochs 1-10)",
        description="Two-panel line chart showing training and validation total loss for C, D, F, G1, G2 and G3 across epochs 1 through 10.",
    )

    _add_heading(document, "New-score trajectories across epochs", page_break=True)
    _add_body(
        document,
        "Every point is a complete O50 evaluation for epochs 3-10. The shared "
        "selection makes checkpoint trajectories comparable, but repeatedly "
        "inspecting the same 50 pairs is exactly why best-epoch selection is biased.",
    )
    _add_picture(
        document,
        score_chart_png,
        title="Figure 2. O50 success rate for the two new scores (epochs 3-10)",
        description="Two-panel line chart showing O50 success rates for first-action Q and rollout mean Q across V2-EMA-SG epochs 3 through 10.",
    )

    _add_heading(document, "Key conclusions and audit boundary", page_break=True)
    first_labels, first_count = _epoch10_winners(inputs, "f_plus_g_first")
    mean_labels, mean_count = _epoch10_winners(inputs, "g_only_f_rollout_mean")
    fixed_labels, fixed_count = _fixed_winners(inputs, "v2", "f_plus_g_first")
    ema_first_mean = _mean_success_percent(
        row
        for row in inputs.new_rows
        if row["epoch"] == 10 and row["score_mode"] == "f_plus_g_first"
    )
    ema_rollout_mean = _mean_success_percent(
        row
        for row in inputs.new_rows
        if row["epoch"] == 10 and row["score_mode"] == "g_only_f_rollout_mean"
    )
    fixed_first_means = {
        version: _mean_success_percent(
            row
            for row in inputs.fixed_rows
            if row["version"] == version and row["score_mode"] == "f_plus_g_first"
        )
        for version in FIXED_VERSIONS
    }
    conclusions = (
        f"Epoch 10 first-Q winner: {first_labels}, {first_count}/50 ({first_count * 2}%).",
        f"Epoch 10 Mean-Q winner: {mean_labels}, {mean_count}/50 ({mean_count * 2}%).",
        f"Across six methods, epoch-10 means are {ema_first_mean:.1f}% first-Q and "
        f"{ema_rollout_mean:.1f}% Mean-Q; fixed first-Q means are "
        f"V0 {fixed_first_means['v0']:.1f}%, V1 {fixed_first_means['v1']:.1f}% "
        f"and V2 {fixed_first_means['v2']:.1f}%.",
        f"Fixed V2 first-Q winner: {fixed_labels}, {fixed_count}/50 "
        f"({fixed_count * 2}%); it shares the EMA table's O50 pair set.",
        "C/D/F/G1/G2/G3 differ only in training-time TD loss or detached weights; "
        "their inference formulas are shared, while each checkpoint learns a "
        "different G.",
        "One training seed and one planning selection are available. These results "
        "do not establish multi-seed superiority or statistical significance.",
    )
    for index, conclusion in enumerate(conclusions, start=1):
        _add_body(document, f"{index}. {conclusion}")
    _v1._add_table(
        document,
        headers=("Audit field", "Locked value"),
        rows=(
            ("Training commit", TRAINING_COMMIT),
            ("Evaluation commit", EVALUATION_COMMIT),
            ("Shared 120-cell episode selection", SHARED_EPISODE_SELECTION_SHA256),
            ("Fixed launcher valid-row-ranks", FIXED_LAUNCHER_RANKS_SHA256),
            ("Action normalization", ACTION_NORMALIZATION_SHA256),
            ("Grid", "96 = 8 epochs × 6 variants × 2 modes; fixed = 24"),
        ),
        widths=(3400, 11000),
    )

    stream = io.BytesIO()
    document.save(stream)
    payload = stream.getvalue()
    if not payload.startswith(b"PK"):
        raise RuntimeError("python-docx did not produce a valid OOXML container.")
    return payload


def _atomic_write(path: Path, payload: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_name: str | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as stream:
            temporary_name = stream.name
            stream.write(payload)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary_name, path)
        temporary_name = None
    finally:
        if temporary_name is not None:
            Path(temporary_name).unlink(missing_ok=True)


def write_report_outputs(
    *,
    markdown: str,
    document: bytes,
    training_chart: bytes,
    score_chart: bytes,
    markdown_output: str | Path,
    docx_output: str | Path,
    chart_dir: str | Path,
) -> tuple[Path, ...]:
    markdown_path = Path(markdown_output)
    docx_path = Path(docx_output)
    chart_root = Path(chart_dir)
    if markdown_path.suffix.lower() != ".md":
        raise ResultsTDV2EMAError("Markdown output must end in .md.")
    if docx_path.suffix.lower() != ".docx":
        raise ResultsTDV2EMAError("DOCX output must end in .docx.")
    training_path = chart_root / "training_validation_loss_curves.png"
    score_path = chart_root / "new_score_epoch_curves.png"
    destinations = (markdown_path, docx_path, training_path, score_path)
    if len({path.resolve() for path in destinations}) != len(destinations):
        raise ResultsTDV2EMAError("Report output paths must be distinct.")
    _atomic_write(training_path, training_chart)
    _atomic_write(score_path, score_chart)
    _atomic_write(markdown_path, markdown.encode("utf-8"))
    _atomic_write(docx_path, document)
    return destinations


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--original-summary", default=str(DEFAULT_ORIGINAL_SUMMARY))
    parser.add_argument("--training-loss-csv", default=str(DEFAULT_TRAINING_LOSS_CSV))
    parser.add_argument("--new-summary", default=str(DEFAULT_NEW_SUMMARY))
    parser.add_argument("--new-results", default=str(DEFAULT_NEW_RESULTS_CSV))
    parser.add_argument("--fixed-results", default=str(DEFAULT_FIXED_RESULTS_CSV))
    parser.add_argument("--base-document", default=str(DEFAULT_BASE_DOCUMENT))
    parser.add_argument("--markdown-output", default=str(DEFAULT_MARKDOWN_OUTPUT))
    parser.add_argument("--docx-output", default=str(DEFAULT_DOCX_OUTPUT))
    parser.add_argument("--chart-dir", default=str(DEFAULT_CHART_DIR))
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Validate 96+24 cells, provenance, loss curves and base DOCX without writing.",
    )
    return parser.parse_args(argv)


def _relative_chart_reference(markdown_output: Path, chart_path: Path) -> str:
    return Path(os.path.relpath(chart_path, start=markdown_output.parent)).as_posix()


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    try:
        base = Path(args.base_document).resolve()
        outputs = {
            Path(args.markdown_output).resolve(),
            Path(args.docx_output).resolve(),
        }
        if base in outputs:
            raise ResultsTDV2EMAError("The locked V0/V1 base cannot be an output.")
        inputs = load_validated_report_inputs(
            original_summary_path=args.original_summary,
            training_loss_csv_path=args.training_loss_csv,
            new_summary_path=args.new_summary,
            new_results_csv_path=args.new_results,
            fixed_results_csv_path=args.fixed_results,
            base_document_path=args.base_document,
        )
        if args.validate_only:
            print(
                "Validated Results TD inputs: 96 EMA cells + 24 fixed cells + 60 loss rows."
            )
            return 0
        training_chart = build_training_loss_chart(inputs)
        score_chart = build_new_score_chart(inputs)
        markdown_path = Path(args.markdown_output)
        chart_root = Path(args.chart_dir)
        markdown = build_markdown_report(
            inputs,
            training_chart_reference=_relative_chart_reference(
                markdown_path, chart_root / "training_validation_loss_curves.png"
            ),
            score_chart_reference=_relative_chart_reference(
                markdown_path, chart_root / "new_score_epoch_curves.png"
            ),
        )
        document = build_results_document(
            inputs,
            training_chart_png=training_chart,
            score_chart_png=score_chart,
        )
        paths = write_report_outputs(
            markdown=markdown,
            document=document,
            training_chart=training_chart,
            score_chart=score_chart,
            markdown_output=args.markdown_output,
            docx_output=args.docx_output,
            chart_dir=args.chart_dir,
        )
    except (
        FileNotFoundError,
        ImportError,
        OSError,
        RuntimeError,
        ResultsTDV2EMAError,
    ) as exc:
        raise SystemExit(str(exc)) from exc
    print("Wrote validated Results TD V2-EMA-SG report:")
    for path in paths:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
