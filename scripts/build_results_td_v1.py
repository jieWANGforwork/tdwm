#!/usr/bin/env python3
"""Append controlled C–G3 V0/V1 results to the locked Results TD base.

This script is deliberately downstream of the result-bundle validator. It
requires a complete V1 ``summary.json``, the exact 50-pair outcomes CSV, a real
loss-chart PNG, the compact formal V0 summary, and the immutable legacy
7-method DOCX. Missing or inconsistent inputs fail before either output copy is
written; the builder never inserts pending/TBD/example result values.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import json
import math
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Sequence

REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_ARTIFACT_ROOT = (
    REPOSITORY_ROOT / "reports/artifacts/actor_free_td_lewm_v1_cube_seed3072"
)
DEFAULT_SUMMARY = DEFAULT_ARTIFACT_ROOT / "summary.json"
DEFAULT_PAIRED = DEFAULT_ARTIFACT_ROOT / "paired_outcomes.csv"
DEFAULT_LOSS_CHART = DEFAULT_ARTIFACT_ROOT / "training_loss_curves.png"
DEFAULT_V0_SUMMARY = (
    REPOSITORY_ROOT
    / "reports/artifacts/actor_free_td_lewm_v0_cube_seed3072/formal_o50_summary.json"
)
DEFAULT_BASE_DOCUMENT = (
    REPOSITORY_ROOT / "reports/results_td_actor_free_cube_seed3072.docx"
)
DEFAULT_REPOSITORY_OUTPUT = (
    REPOSITORY_ROOT / "reports/results_td_actor_free_td_lewm_v0_v1_cube_seed3072.docx"
)
PROJECT_DELIVERY_ROOT = (
    REPOSITORY_ROOT.parent.parent
    if REPOSITORY_ROOT.parent.name == "tmp"
    and (REPOSITORY_ROOT.parent.parent / "AGENTS.md").is_file()
    else REPOSITORY_ROOT
)
DEFAULT_PROJECT_COPY = PROJECT_DELIVERY_ROOT / "Results TD.docx"

VARIANT_ORDER = ("c", "d", "f", "g1", "g2", "g3")
SCORE_MODES = ("f_only", "g_only", "f_plus_g")
EPISODES = 50
TRAINING_SEED = 3072
TRAINING_EPOCHS = 10
TRAINING_STEPS = 127_960
PLANNING_SEED = 42
GOAL_OFFSET = 50
PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"
SHA256_PATTERN = re.compile(r"[0-9a-f]{64}")
PREDICTOR_PARAMETERS = 379_072
V0_PREDICTOR_PARAMETERS = 336_320
APPEND_TABLE_WIDTH_DXA = 14_400
BASE_DOCUMENT_SHA256 = (
    "d27ab14888bec3f96b3a5974f30afb73de07d18fe1c445b8cfc95c3072f6eaed"
)
LEGACY_TEXT_REPLACEMENTS = {
    "锁定协议：7 个方法均完成 10 epochs / 127,960 optimizer updates；每种方法评测 3 种分数，共 21 个 formal O50。goal offset=50；共同 selection SHA-256=e46ea81cce2e….": (
        "Locked protocol: all 7 methods completed 10 epochs / 127,960 optimizer "
        "updates; each method was evaluated with 3 scores, for 21 formal O50 runs. "
        "goal offset=50; common selection SHA-256=e46ea81cce2e…."
    ),
    "Combined O50 排名": "Combined O50 ranking",
    "表 1. 排名只使用预先定义的 combined 列；相同成功数共享名次，不按三列中的最佳值重新排序。": (
        "Table 1. Ranking uses only the pre-specified combined column; equal success "
        "counts share a rank. The best of the three score columns is not selected "
        "post hoc."
    ),
    "7×3 推理分数矩阵": "7×3 inference-score matrix",
    "表 2. 每格依次给出 successes/episodes、success rate 和 elapsed time（时:分:秒）；F-only 不忽略最后一个规划动作。": (
        "Table 2. Each cell reports successes/episodes, success rate and elapsed "
        "time (h:mm:ss); F-only includes the final planning action."
    ),
    "训练收敛与 checkpoint": "Training convergence and checkpoints",
    "比较限制：不同方法的 auxiliary loss 数量和定义不同；total loss 只用于判断各方法自身是否收敛，不能按曲线高低做跨方法性能排序。": (
        "Comparison limit: auxiliary-loss counts and definitions differ. Total loss "
        "is only a within-method convergence check and cannot rank methods across "
        "structures."
    ),
    "图 1. 归档器派生的训练与 validation total-loss 曲线；横向结果排名使用第 1 页的 combined O50。": (
        "Figure 1. Training and validation total-loss curves derived by the "
        "archiver; cross-method ranking uses combined O50 on page 1."
    ),
    "表 3. Final 指 epoch 10；Best validation 在各自 10 个 epoch 内选择。哈希和 commit 均来自归档摘要。": (
        "Table 3. Final is epoch 10; Best validation is selected within each "
        "method's 10 epochs. Hashes and commits come from the archive summary."
    ),
    "四个基础 Actor-Free TD-LeWM 结构": "Four base Actor-Free TD-LeWM structures",
    "共享 LeWM encoder/predictor；差异集中在 successor 输入、TD 梯度路径和 real/predicted 分支。": (
        "Shared LeWM encoder/predictor; differences are successor inputs, TD "
        "gradient paths and real/predicted branches."
    ),
    "共同目标：L_LeWM = L_pred + 0.09 L_SIGReg。Successor 目标使用真实 EMA next latent、EMA head 和数据集 next action，并对 target stop-gradient；所有方法均无 Actor。": (
        "Common objective: L_LeWM = L_pred + 0.09 L_SIGReg. The successor target "
        "uses the real EMA next latent, EMA head and dataset next action, with a "
        "stop-gradient target; every method is actor-free."
    ),
    "表 4. 方法描述取自归档 summary.method_spec；公式是固定方法定义，不包含手填实验结果。": (
        "Table 4. Method descriptions come from archive summary.method_spec; the "
        "equations are fixed definitions, not manually entered results."
    ),
    "三种推理分数的统一含义": "Unified meaning of the three inference scores",
    "控制口径：全部结构在同一 swm/OGBCube-v0、同一 selection、同一 planning seed 42 下比较；G-only 仍使用 LeWM rollout 构造 tail context，只是不把 F cost 加入排序。": (
        "Control protocol: all structures use the same swm/OGBCube-v0, selection "
        "and planning seed 42. G-only still uses a LeWM rollout to build tail "
        "context, but excludes F cost from ranking."
    ),
    "三种扩展、目标公式与审计来源": (
        "Three extensions, target equations and audit sources"
    ),
    "Goal readout TD、EMA imaginary bootstrap，以及直接 goal-conditioned critic 的并列对照。": (
        "Side-by-side comparison of goal-readout TD, EMA imaginary bootstrap and a "
        "direct goal-conditioned critic."
    ),
    "表 5. Direct 方法用 C-only / F+C；其余 successor 方法用 G-only / F+G。训练仍然不引入 Actor。": (
        "Table 5. Direct uses C-only / F+C; successor methods use G-only / F+G. "
        "Training remains actor-free."
    ),
    "目标与规划公式": "Target and planning equations",
    "SF TD： yₜ^SF = (1−γ)χ(z̄ₜ₊₁) + γ(1−dₜ)Ḡ(next)。\nGoal TD： C_g(H,a)=G(H,a)ᵀw(z_g)，yₜ^g=(1−γ)c(z̄ₜ₊₁,z_g)+γ(1−dₜ^g)C̄_g(next)。\nImaginary： Ĥₜ₊₁=sg(P̄(Hₜ,aₜ))，再从 Ĥₜ₊₁ bootstrap；Direct：C(H,a,z_g) 直接预测 scalar cost。\n规划：J_F=Σₖ₌₁ᴴγᵏ⁻¹c(z̃ₜ₊ₖ,z_g)；combined 为 F prefix 加 γᴴ⁻¹ 的 G/C tail。CEM 最小化 cost，并非训练 Actor 去最大化单步点积。": (
        "SF TD: yₜ^SF = (1−γ)χ(z̄ₜ₊₁) + γ(1−dₜ)Ḡ(next).\n"
        "Goal TD: C_g(H,a)=G(H,a)ᵀw(z_g), yₜ^g=(1−γ)c(z̄ₜ₊₁,z_g)+"
        "γ(1−dₜ^g)C̄_g(next).\n"
        "Imaginary: Ĥₜ₊₁=sg(P̄(Hₜ,aₜ)), then bootstrap from Ĥₜ₊₁; Direct: "
        "C(H,a,z_g) predicts scalar cost directly.\n"
        "Planning: J_F=Σₖ₌₁ᴴγᵏ⁻¹c(z̃ₜ₊ₖ,z_g); combined adds the discounted G/C "
        "tail to the F prefix. CEM minimizes cost; it does not train an Actor to "
        "maximize a one-step dot product."
    ),
    "Combined 结论：本次首位为 Hybrid 27/50 (54%)。该结论仅对应 combined score；F-only 与 G/C-only 是同一 checkpoint 的推理消融。边界：One training seed and one matched O50 planning selection. Rank only combined inference scores; this is not a multi-seed claim.": (
        "Combined conclusion: Hybrid ranks first at 27/50 (54%). This claim applies "
        "only to the combined score; F-only and G/C-only are inference ablations of "
        "the same checkpoint. Boundary: one training seed and one matched O50 "
        "planning selection. Rank only combined inference scores; this is not a "
        "multi-seed claim."
    ),
    "输入文件 SHA-256：selection=e46ea81cce2e6a9a5df05ba04893b4181cbd8979340111a012c30f1efa2d7ee7；summary.json=9fa31a12476ff4a332e555505fc720964d81c5640c77cc394beb5df538ae2da6；training_loss_curves.png=4cce0d5891b8b51e5e3dc7481bd9582de85e69c9b01a9fcb6d025ed271957a60；paired_outcomes.csv=3047c9a419f066548b6e942e29c49b35d286670dc60e5db27fcd845897a6f2ae。": (
        "Input SHA-256: selection=e46ea81cce2e6a9a5df05ba04893b4181cbd8979340111a012c30f1efa2d7ee7; "
        "summary.json=9fa31a12476ff4a332e555505fc720964d81c5640c77cc394beb5df538ae2da6; "
        "training_loss_curves.png=4cce0d5891b8b51e5e3dc7481bd9582de85e69c9b01a9fcb6d025ed271957a60; "
        "paired_outcomes.csv=3047c9a419f066548b6e942e29c49b35d286670dc60e5db27fcd845897a6f2ae."
    ),
    "Score 定义来源：F-only — LeWM rolls and scores all H predicted states with the normalized discounted latent-goal cost.  Combined — F prefix cost plus discounted successor G tail cost. / F prefix cost plus discounted direct critic C tail cost.": (
        "Score definitions: F-only — LeWM rolls and scores all H predicted states "
        "with the normalized discounted latent-goal cost. Combined — F prefix cost "
        "plus discounted successor G tail cost / discounted direct-critic C tail cost."
    ),
    "排名": "Rank",
    "方法": "Method",
    "族": "Family",
    "成功": "Successes",
    "率": "Rate",
    "耗时": "Elapsed",
    "尾部类型": "Tail type",
    "仅 Combined 排名": "Combined rank only",
    "网络结构": "Network",
    "训练损失": "Training loss",
    "特殊设计 / 梯度路径": "Special design / gradient path",
    "推理": "Inference",
    "L = L_LeWM + αᵤ(L_TD^pred + L_TD^real)\n同一 G 同时接收预测与真实 latent history": (
        "L = L_LeWM + αᵤ(L_TD^pred + L_TD^real)\n"
        "The same G consumes predicted and real latent histories"
    ),
    "L = L_LeWM + αᵤ L_TD^real\nG(Hₜ^real, aₜ) 与 LeWM predictor 平行": (
        "L = L_LeWM + αᵤ L_TD^real\nG(Hₜ^real, aₜ) is parallel to the LeWM predictor"
    ),
    "候选动作排序信号": "Candidate-action ranking signal",
    "L = L_LeWM + αᵤ(L_TD^real + L_TD^pred)\nbootstrap state 由 stop-gradient EMA-LeWM 想象一步": (
        "L = L_LeWM + αᵤ(L_TD^real + L_TD^pred)\n"
        "The bootstrap state is imagined one step by stop-gradient EMA-LeWM"
    ),
    "L = L_LeWM + αᵤ(L_C-TD^real + L_C-TD^pred)\nC(H,a,z_g) 直接输出 goal-conditioned scalar cost": (
        "L = L_LeWM + αᵤ(L_C-TD^real + L_C-TD^pred)\n"
        "C(H,a,z_g) directly outputs a goal-conditioned scalar cost"
    ),
}
EXPECTED_ACTION_ENCODER_SHA256 = (
    "2657b55140013b4b071cd8cdea63f1eac5c65c498d55331c7499744ef31a9cd3"
)


@dataclass(frozen=True)
class ValidatedReportInputs:
    summary: Mapping[str, Any]
    v0_summary: Mapping[str, Any]
    paired_rows: tuple[Mapping[str, str], ...]
    loss_chart: Path
    base_document: Path
    selection_sha256: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Rebuild Results TD from the immutable legacy report plus complete "
            "validated C–G3 V0/V1 evidence."
        )
    )
    parser.add_argument("--summary", default=str(DEFAULT_SUMMARY))
    parser.add_argument(
        "--paired",
        "--paired-outcomes",
        dest="paired",
        default=str(DEFAULT_PAIRED),
    )
    parser.add_argument("--loss-chart", default=str(DEFAULT_LOSS_CHART))
    parser.add_argument("--v0-summary", default=str(DEFAULT_V0_SUMMARY))
    parser.add_argument("--base-document", default=str(DEFAULT_BASE_DOCUMENT))
    parser.add_argument("--output", default=str(DEFAULT_REPOSITORY_OUTPUT))
    parser.add_argument("--project-copy", default=str(DEFAULT_PROJECT_COPY))
    parser.add_argument(
        "--no-project-copy",
        action="store_true",
        help="Write only --output, not the project-root Results TD.docx copy.",
    )
    return parser.parse_args()


def _mapping(value: Any, *, context: str) -> Mapping[str, Any]:
    if not isinstance(value, Mapping):
        raise ValueError(f"{context} must be an object.")
    return value


def _exact_int(value: Any, expected: int, *, context: str) -> int:
    if isinstance(value, bool) or not isinstance(value, int) or value != expected:
        raise ValueError(f"{context} must equal {expected}, found {value!r}.")
    return value


def _finite_number(value: Any, *, context: str) -> float:
    if isinstance(value, bool):
        raise ValueError(f"{context} must be numeric, found {value!r}.")
    try:
        number = float(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{context} must be numeric, found {value!r}.") from exc
    if not math.isfinite(number):
        raise ValueError(f"{context} must be finite, found {value!r}.")
    return number


def _nonempty_text(value: Any, *, context: str) -> str:
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{context} must be a non-empty string.")
    return value.strip()


def _success_percent(run: Mapping[str, Any], *, context: str) -> tuple[int, float]:
    successes = run.get("success_count")
    if isinstance(successes, bool) or not isinstance(successes, int):
        raise ValueError(f"{context}.success_count must be an integer.")
    if not 0 <= successes <= EPISODES:
        raise ValueError(f"{context}.success_count must lie in [0, {EPISODES}].")
    expected = successes * 100.0 / EPISODES
    rate = _finite_number(run.get("success_rate"), context=f"{context}.success_rate")
    if not math.isclose(rate, expected / 100.0, rel_tol=0.0, abs_tol=1e-12):
        raise ValueError(
            f"{context}.success_rate {rate} disagrees with {successes}/{EPISODES}."
        )
    if "success_rate_percent" in run:
        percent = _finite_number(
            run["success_rate_percent"],
            context=f"{context}.success_rate_percent",
        )
        if not math.isclose(percent, expected, rel_tol=0.0, abs_tol=1e-12):
            raise ValueError(
                f"{context}.success_rate_percent {percent} disagrees with "
                f"{successes}/{EPISODES}."
            )
    return successes, expected


def _validate_summary(source: Path) -> Mapping[str, Any]:
    if not source.is_file():
        raise FileNotFoundError(f"Validated summary does not exist: {source}")
    try:
        payload = json.loads(source.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Cannot parse validated summary {source}: {exc}") from exc
    summary = _mapping(payload, context="summary")
    _exact_int(summary.get("schema_version"), 1, context="summary.schema_version")
    study = _mapping(summary.get("study"), context="summary.study")
    if study.get("method_family") != "actor_free_td_lewm_v1":
        raise ValueError("summary.study.method_family must be actor_free_td_lewm_v1.")
    if study.get("environment") != "cube":
        raise ValueError("summary.study.environment must be cube.")
    expected_study = {
        "training_count": len(VARIANT_ORDER),
        "evaluation_count": len(VARIANT_ORDER) * len(SCORE_MODES),
        "episodes_per_evaluation": EPISODES,
        "training_seed": TRAINING_SEED,
        "planning_seed": PLANNING_SEED,
        "goal_offset": GOAL_OFFSET,
    }
    for key, expected in expected_study.items():
        _exact_int(study.get(key), expected, context=f"summary.study.{key}")
    if tuple(study.get("score_modes", ())) != SCORE_MODES:
        raise ValueError(f"summary.study.score_modes must equal {list(SCORE_MODES)}.")
    expected_horizons = {"f_only": 5, "g_only": 1, "f_plus_g": 5}
    if study.get("formal_horizon_by_score_mode") != expected_horizons:
        raise ValueError(
            "summary.study.formal_horizon_by_score_mode must equal "
            f"{expected_horizons}."
        )
    if study.get("ranking_metric") != "f_plus_g_success_count_only":
        raise ValueError("summary.study.ranking_metric must be F+G success count.")
    if study.get("single_training_seed") is not True:
        raise ValueError("summary.study.single_training_seed must be true.")
    if study.get("single_planning_selection") is not True:
        raise ValueError("summary.study.single_planning_selection must be true.")

    architecture = _mapping(summary.get("architecture"), context="summary.architecture")
    if architecture.get("lewm_frozen") is not True:
        raise ValueError("summary.architecture.lewm_frozen must be true.")
    if architecture.get("shared_action_encoder_frozen") is not True:
        raise ValueError(
            "summary.architecture.shared_action_encoder_frozen must be true."
        )
    architecture_expected = {
        "state_dim": 192,
        "raw_action_dim": 25,
        "action_embedding_dim": 192,
        "task_dim": 192,
        "output_dim": 192,
        "trainable_lewm_parameters": 0,
        "predictor_parameters": PREDICTOR_PARAMETERS,
    }
    for key, expected in architecture_expected.items():
        _exact_int(
            architecture.get(key),
            expected,
            context=f"summary.architecture.{key}",
        )
    action_encoder_sha = architecture.get("action_encoder_state_sha256")
    if action_encoder_sha != EXPECTED_ACTION_ENCODER_SHA256:
        raise ValueError(
            "summary.architecture.action_encoder_state_sha256 differs from "
            "the locked shared V1 Action Encoder."
        )
    if (
        architecture.get("action_processing")
        != "raw25_to_frozen_world_model.action_encoder_to_embedding192"
    ):
        raise ValueError("summary.architecture.action_processing is not V1 raw25→192.")
    if architecture.get("actor") != "none" or architecture.get("reward") != "none":
        raise ValueError("V1 Results TD requires actor=none and reward=none.")

    selection = _mapping(summary.get("selection"), context="summary.selection")
    _exact_int(
        selection.get("episode_count"),
        EPISODES,
        context="summary.selection.episode_count",
    )
    acceptance = _mapping(
        summary.get("training_acceptance"),
        context="summary.training_acceptance",
    )
    if acceptance.get("status") not in ("PASS", "PASS_WITH_WARNINGS"):
        raise ValueError(
            "summary.training_acceptance.status must be PASS or PASS_WITH_WARNINGS."
        )
    acceptance_warnings = acceptance.get("warnings", [])
    if not isinstance(acceptance_warnings, list) or any(
        not isinstance(item, str) or not item.strip() for item in acceptance_warnings
    ):
        raise ValueError("summary.training_acceptance.warnings must be strings.")
    if acceptance.get("status") == "PASS_WITH_WARNINGS" and not acceptance_warnings:
        raise ValueError(
            "PASS_WITH_WARNINGS requires preserved training-acceptance warnings."
        )
    if acceptance.get("status") == "PASS" and acceptance_warnings:
        raise ValueError("PASS training acceptance must not carry warnings.")
    if (
        acceptance.get("expected_action_encoder_sha256")
        != EXPECTED_ACTION_ENCODER_SHA256
    ):
        raise ValueError(
            "summary.training_acceptance expected Action Encoder hash is inconsistent."
        )

    methods = _mapping(summary.get("methods"), context="summary.methods")
    if set(methods) != set(VARIANT_ORDER):
        raise ValueError(
            "summary.methods must contain exactly "
            f"{list(VARIANT_ORDER)}, found {sorted(methods)}."
        )
    for variant in VARIANT_ORDER:
        method = _mapping(methods[variant], context=f"summary.methods.{variant}")
        _nonempty_text(
            method.get("display_name"),
            context=f"summary.methods.{variant}.display_name",
        )
        for key in ("network", "training_loss", "special_mechanism"):
            _nonempty_text(method.get(key), context=f"summary.methods.{variant}.{key}")
        inference = _mapping(
            method.get("inference"),
            context=f"summary.methods.{variant}.inference",
        )
        if set(inference) != set(SCORE_MODES):
            raise ValueError(
                f"summary.methods.{variant}.inference must contain three modes."
            )
        for mode in SCORE_MODES:
            _nonempty_text(
                inference.get(mode),
                context=f"summary.methods.{variant}.inference.{mode}",
            )
        training = _mapping(
            method.get("training"), context=f"summary.methods.{variant}.training"
        )
        _exact_int(
            training.get("seed"),
            TRAINING_SEED,
            context=f"summary.methods.{variant}.training.seed",
        )
        _exact_int(
            training.get("epochs"),
            TRAINING_EPOCHS,
            context=f"summary.methods.{variant}.training.epochs",
        )
        _exact_int(
            training.get("global_step"),
            TRAINING_STEPS,
            context=f"summary.methods.{variant}.training.global_step",
        )
        _nonempty_text(
            training.get("checkpoint_path"),
            context=f"summary.methods.{variant}.training.checkpoint_path",
        )
        curve = training.get("loss_curve")
        if not isinstance(curve, list) or len(curve) != TRAINING_EPOCHS:
            raise ValueError(
                f"summary.methods.{variant}.training.loss_curve must contain "
                f"{TRAINING_EPOCHS} epochs."
            )
        if training.get("train_loss_semantics") != "method_specific_objective":
            raise ValueError(
                f"summary.methods.{variant}.training.train_loss_semantics is wrong."
            )
        if training.get("validation_loss_semantics") != "common_base_td":
            raise ValueError(
                f"summary.methods.{variant}.training.validation_loss_semantics "
                "is wrong."
            )
        evaluations = _mapping(
            method.get("evaluations"),
            context=f"summary.methods.{variant}.evaluations",
        )
        if set(evaluations) != set(SCORE_MODES):
            raise ValueError(
                f"summary.methods.{variant}.evaluations must contain exactly "
                f"{list(SCORE_MODES)}."
            )
        for mode in SCORE_MODES:
            run = _mapping(
                evaluations[mode],
                context=f"summary.methods.{variant}.evaluations.{mode}",
            )
            if run.get("score_mode") != mode:
                raise ValueError(f"{variant}/{mode}.score_mode is inconsistent.")
            _exact_int(
                run.get("planning_horizon"),
                expected_horizons[mode],
                context=f"{variant}/{mode}.planning_horizon",
            )
            _success_percent(run, context=f"{variant}/{mode}")
            _nonempty_text(
                run.get("checkpoint_path"),
                context=f"{variant}/{mode}.checkpoint_path",
            )
            checkpoint_sha = run.get("checkpoint_sha256")
            if (
                not isinstance(checkpoint_sha, str)
                or SHA256_PATTERN.fullmatch(checkpoint_sha) is None
            ):
                raise ValueError(f"{variant}/{mode}.checkpoint_sha256 must be SHA-256.")
        checkpoint_paths = {
            evaluations[mode].get("checkpoint_path") for mode in SCORE_MODES
        }
        checkpoint_hashes = {
            evaluations[mode].get("checkpoint_sha256") for mode in SCORE_MODES
        }
        if len(checkpoint_paths) != 1 or len(checkpoint_hashes) != 1:
            raise ValueError(
                f"summary method {variant!r} does not use one checkpoint in all modes."
            )
        if next(iter(checkpoint_paths)) != training.get("checkpoint_path"):
            raise ValueError(
                f"summary method {variant!r} evaluation checkpoint differs from "
                "its training deployment checkpoint."
            )

    ranking = summary.get("ranking_by_f_plus_g")
    if not isinstance(ranking, list) or len(ranking) != len(VARIANT_ORDER):
        raise ValueError("summary.ranking_by_f_plus_g must contain six rows.")
    ranked_variants: list[str] = []
    expected_rank = 0
    previous_count: int | None = None
    for row_index, value in enumerate(ranking, start=1):
        row = _mapping(value, context=f"summary.ranking_by_f_plus_g[{row_index - 1}]")
        variant = row.get("variant")
        if variant not in VARIANT_ORDER:
            raise ValueError(
                f"Ranking row {row_index} has invalid variant {variant!r}."
            )
        ranked_variants.append(str(variant))
        run = _mapping(
            methods[variant]["evaluations"]["f_plus_g"],
            context=f"summary.methods.{variant}.evaluations.f_plus_g",
        )
        successes, percent = _success_percent(run, context=f"{variant}/f_plus_g")
        if row.get("success_count") != successes:
            raise ValueError(f"Ranking row {row_index} has inconsistent successes.")
        if successes != previous_count:
            expected_rank = row_index
            previous_count = successes
        _exact_int(
            row.get("rank"),
            expected_rank,
            context=f"summary.ranking_by_f_plus_g[{row_index - 1}].rank",
        )
        ranked_rate = _finite_number(
            row.get("success_rate"),
            context=f"summary.ranking_by_f_plus_g[{row_index - 1}].success_rate",
        )
        if not math.isclose(ranked_rate, percent / 100.0, rel_tol=0.0, abs_tol=1e-12):
            raise ValueError(f"Ranking row {row_index} has inconsistent rate.")
    if set(ranked_variants) != set(VARIANT_ORDER):
        raise ValueError("Combined ranking does not contain all six variants once.")
    ranked_successes = [int(row["success_count"]) for row in ranking]
    if ranked_successes != sorted(ranked_successes, reverse=True):
        raise ValueError("Combined ranking is not ordered by descending successes.")
    validation = _mapping(summary.get("validation"), context="summary.validation")
    for key in (
        "complete_6x3_bundle",
        "formal_o50_only",
        "common_selection_across_18_runs",
        "same_checkpoint_within_each_method",
        "frozen_lewm_and_action_encoder",
        "no_lewm_training_loss",
        "success_rates_match_episode_outcomes",
    ):
        if validation.get(key) is not True:
            raise ValueError(f"summary.validation.{key} must be true.")
    return summary


def _parse_bool(value: str, *, context: str) -> bool:
    normalized = value.strip().lower()
    if normalized == "true":
        return True
    if normalized == "false":
        return False
    raise ValueError(f"{context} must be true or false, found {value!r}.")


def _validate_paired(
    source: Path,
    *,
    summary: Mapping[str, Any],
) -> tuple[tuple[Mapping[str, str], ...], str]:
    if not source.is_file():
        raise FileNotFoundError(f"Paired outcomes CSV does not exist: {source}")
    with source.open(newline="", encoding="utf-8") as stream:
        reader = csv.DictReader(stream)
        rows = tuple(reader)
        fieldnames = tuple(reader.fieldnames or ())
    if len(rows) != EPISODES:
        raise ValueError(
            f"Paired outcomes must contain exactly {EPISODES} rows, found {len(rows)}."
        )
    base_columns = {
        "selection_position",
        "selection_sha256",
        "episode_index",
        "start_step",
        "goal_step",
        "valid_row_rank",
        "pair_hash",
    }
    outcome_columns = tuple(
        f"success_{variant}__{mode}"
        for variant in VARIANT_ORDER
        for mode in SCORE_MODES
    )
    missing = base_columns.union(outcome_columns) - set(fieldnames)
    if missing:
        raise ValueError(f"Paired outcomes CSV is missing columns: {sorted(missing)}")
    unexpected_success = {
        name for name in fieldnames if name.startswith("success_")
    } - set(outcome_columns)
    if unexpected_success:
        raise ValueError(
            f"Paired outcomes CSV has unexpected success columns: "
            f"{sorted(unexpected_success)}"
        )

    selection_hashes: set[str] = set()
    outcome_counts = {column: 0 for column in outcome_columns}
    seen_pairs: set[tuple[int, int, int]] = set()
    for expected_position, row in enumerate(rows):
        try:
            position = int(row["selection_position"])
            episode = int(row["episode_index"])
            start = int(row["start_step"])
            goal = int(row["goal_step"])
            int(row["valid_row_rank"])
        except (TypeError, ValueError) as exc:
            raise ValueError(
                f"Paired outcomes row {expected_position + 2} has invalid integers."
            ) from exc
        if position != expected_position:
            raise ValueError("Paired outcomes selection_position must be 0..49.")
        if goal - start != GOAL_OFFSET:
            raise ValueError(
                f"Paired row {expected_position} does not use goal offset {GOAL_OFFSET}."
            )
        pair = (episode, start, goal)
        if pair in seen_pairs:
            raise ValueError(f"Paired outcomes contains duplicate pair {pair}.")
        seen_pairs.add(pair)
        pair_hash = row["pair_hash"].strip().lower()
        if SHA256_PATTERN.fullmatch(pair_hash) is None:
            raise ValueError(f"Paired row {expected_position} has invalid pair_hash.")
        selection_sha = row["selection_sha256"].strip().lower()
        if SHA256_PATTERN.fullmatch(selection_sha) is None:
            raise ValueError(
                f"Paired row {expected_position} has invalid selection_sha256."
            )
        selection_hashes.add(selection_sha)
        for column in outcome_columns:
            outcome_counts[column] += int(
                _parse_bool(row[column], context=f"row {expected_position} {column}")
            )
    if len(selection_hashes) != 1:
        raise ValueError("All paired outcomes rows must share one selection SHA-256.")
    selection_sha = next(iter(selection_hashes))
    summary_selection = _mapping(
        summary.get("selection"), context="summary.selection"
    ).get("episode_selection_json_sha256")
    if summary_selection != selection_sha:
        raise ValueError(
            "Paired outcomes selection SHA-256 differs from summary.selection."
        )
    methods = _mapping(summary["methods"], context="summary.methods")
    for variant in VARIANT_ORDER:
        evaluations = _mapping(
            methods[variant]["evaluations"],
            context=f"summary.methods.{variant}.evaluations",
        )
        for mode in SCORE_MODES:
            expected = evaluations[mode]["success_count"]
            observed = outcome_counts[f"success_{variant}__{mode}"]
            if observed != expected:
                raise ValueError(
                    f"Paired outcome count {variant}/{mode}={observed} differs "
                    f"from summary={expected}."
                )
    return rows, selection_sha


def _validate_loss_chart(source: Path) -> Path:
    if not source.is_file():
        raise FileNotFoundError(f"Loss chart PNG does not exist: {source}")
    if source.stat().st_size <= len(PNG_SIGNATURE):
        raise ValueError(f"Loss chart PNG is empty or truncated: {source}")
    with source.open("rb") as stream:
        if stream.read(len(PNG_SIGNATURE)) != PNG_SIGNATURE:
            raise ValueError(f"Loss chart is not a PNG file: {source}")
    return source


def _validate_base_document(source: Path) -> Path:
    if not source.is_file():
        raise FileNotFoundError(f"Base Results TD document does not exist: {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError(f"Base Results TD document must be a DOCX: {source}")
    payload = source.read_bytes()
    if not payload.startswith(b"PK"):
        raise ValueError(f"Base Results TD document is not valid OOXML: {source}")
    digest = hashlib.sha256(payload).hexdigest()
    if digest != BASE_DOCUMENT_SHA256:
        raise ValueError(
            "Base Results TD document differs from the immutable legacy 7-method "
            f"source (expected {BASE_DOCUMENT_SHA256}, found {digest})."
        )
    return source


def _validate_v0_summary(
    source: Path,
    *,
    expected_selection_sha256: str,
) -> Mapping[str, Any]:
    if not source.is_file():
        raise FileNotFoundError(f"V0 formal O50 summary does not exist: {source}")
    try:
        payload = json.loads(source.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Cannot parse V0 formal O50 summary {source}: {exc}") from exc
    summary = _mapping(payload, context="v0_summary")
    _exact_int(summary.get("schema_version"), 1, context="v0_summary.schema_version")
    study = _mapping(summary.get("study"), context="v0_summary.study")
    expected_study: Mapping[str, Any] = {
        "method_family": "actor_free_td_lewm_v0",
        "training_commit": "79706d3",
        "training_seed": TRAINING_SEED,
        "planning_seed": PLANNING_SEED,
        "episodes_per_evaluation": EPISODES,
        "goal_offset": GOAL_OFFSET,
        "score_modes": list(SCORE_MODES),
    }
    for key, expected in expected_study.items():
        if study.get(key) != expected:
            raise ValueError(
                f"v0_summary.study.{key} must equal {expected!r}, "
                f"found {study.get(key)!r}."
            )
    _nonempty_text(study.get("source_root"), context="v0_summary.study.source_root")
    selection = _mapping(summary.get("selection"), context="v0_summary.selection")
    if selection.get("sha256") != expected_selection_sha256:
        raise ValueError(
            "V0 and V1 must use the same formal 50-pair selection SHA-256."
        )
    methods = _mapping(summary.get("methods"), context="v0_summary.methods")
    if set(methods) != set(VARIANT_ORDER):
        raise ValueError("V0 summary must contain exactly C/D/F/G1/G2/G3.")
    expected_horizons = {"f_only": 5, "g_only": 1, "f_plus_g": 5}
    for variant in VARIANT_ORDER:
        evaluations = _mapping(methods[variant], context=f"v0_summary.{variant}")
        if set(evaluations) != set(SCORE_MODES):
            raise ValueError(f"V0 {variant} must contain all three score modes.")
        for mode in SCORE_MODES:
            run = _mapping(evaluations[mode], context=f"v0_summary.{variant}.{mode}")
            successes = run.get("success_count")
            if isinstance(successes, bool) or not isinstance(successes, int):
                raise ValueError(f"V0 {variant}/{mode} success_count must be integer.")
            if not 0 <= successes <= EPISODES:
                raise ValueError(f"V0 {variant}/{mode} success_count is out of range.")
            expected_percent = successes * 100.0 / EPISODES
            percent = _finite_number(
                run.get("success_rate_percent"),
                context=f"v0_summary.{variant}.{mode}.success_rate_percent",
            )
            if not math.isclose(percent, expected_percent, rel_tol=0, abs_tol=1e-12):
                raise ValueError(f"V0 {variant}/{mode} rate disagrees with count.")
            _exact_int(
                run.get("planning_horizon"),
                expected_horizons[mode],
                context=f"v0_summary.{variant}.{mode}.planning_horizon",
            )
            digest = run.get("results_sha256")
            if not isinstance(digest, str) or SHA256_PATTERN.fullmatch(digest) is None:
                raise ValueError(f"V0 {variant}/{mode} results_sha256 is invalid.")
    return summary


def load_validated_report_inputs(
    *,
    summary_path: str | Path,
    paired_path: str | Path,
    loss_chart_path: str | Path,
    v0_summary_path: str | Path = DEFAULT_V0_SUMMARY,
    base_document_path: str | Path = DEFAULT_BASE_DOCUMENT,
) -> ValidatedReportInputs:
    """Validate all three required report inputs without writing output."""

    summary = _validate_summary(Path(summary_path))
    paired_rows, selection_sha = _validate_paired(Path(paired_path), summary=summary)
    loss_chart = _validate_loss_chart(Path(loss_chart_path))
    v0_summary = _validate_v0_summary(
        Path(v0_summary_path), expected_selection_sha256=selection_sha
    )
    base_document = _validate_base_document(Path(base_document_path))
    return ValidatedReportInputs(
        summary=summary,
        v0_summary=v0_summary,
        paired_rows=paired_rows,
        loss_chart=loss_chart,
        base_document=base_document,
        selection_sha256=selection_sha,
    )


def _format_result(run: Mapping[str, Any]) -> str:
    successes, percent = _success_percent(run, context="report evaluation")
    return f"{successes}/{EPISODES} ({percent:.0f}%)"


def _format_v0_result(run: Mapping[str, Any]) -> str:
    successes = int(run["success_count"])
    percent = float(run["success_rate_percent"])
    return f"{successes}/{EPISODES} ({percent:.0f}%)"


def _curve_final_and_best(method: Mapping[str, Any]) -> tuple[float, float, float, int]:
    training = _mapping(method["training"], context="report training")
    raw_curve = training["loss_curve"]
    points: list[tuple[int, float, float]] = []
    for index, raw in enumerate(raw_curve):
        row = _mapping(raw, context=f"report loss_curve[{index}]")
        epoch = row.get("epoch")
        if isinstance(epoch, bool) or not isinstance(epoch, int):
            raise ValueError("Report loss curve epochs must be integers.")
        train_value = row.get(
            "train_method_loss",
            row.get(
                "train_method_objective",
                row.get("train_method_objective_loss", row.get("train_loss")),
            ),
        )
        validation_value = row.get(
            "validation_base_td_loss",
            row.get(
                "validation_common_base_td",
                row.get("validation_common_base_td_loss", row.get("validation_loss")),
            ),
        )
        points.append(
            (
                epoch,
                _finite_number(train_value, context=f"loss_curve[{index}].train"),
                _finite_number(
                    validation_value,
                    context=f"loss_curve[{index}].validation",
                ),
            )
        )
    points.sort()
    if tuple(point[0] for point in points) != tuple(range(1, 11)):
        raise ValueError("Report loss curve must contain epochs 1..10 exactly once.")
    best = min(points, key=lambda point: point[2])
    return points[-1][1], points[-1][2], best[2], best[0]


def _set_cell_text(cell: Any, text: str, *, bold: bool = False) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if sum(widths) != APPEND_TABLE_WIDTH_DXA:
        raise ValueError(
            "Table widths must sum to "
            f"{APPEND_TABLE_WIDTH_DXA} DXA, found {sum(widths)}."
        )
    table.autofit = False
    table_properties = table._tbl.tblPr
    for tag, attributes in (
        (
            "w:tblW",
            {"w:w": str(APPEND_TABLE_WIDTH_DXA), "w:type": "dxa"},
        ),
        ("w:tblInd", {"w:w": "120", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        element = table_properties.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            table_properties.append(element)
        for key, value in attributes.items():
            element.set(qn(key), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            properties = cell._tc.get_or_add_tcPr()
            tc_width = properties.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                properties.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            margins = properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                properties.append(margins)
            for side, value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _prevent_row_split(row: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _add_table(
    document: Any,
    *,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True)
        _shade_cell(cell, "F2F4F7")
    _repeat_header(table.rows[0])
    _prevent_row_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        cells = row.cells
        for cell, value in zip(cells, values):
            _set_cell_text(cell, value)
    _set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_after = document.styles[
        "Normal"
    ].paragraph_format.space_after
    return table


def _set_run_font(run: Any, *, size: float, color: str, bold: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _translate_legacy_text_for_render(document: Any) -> None:
    """Translate the immutable base's CJK runs for reliable LibreOffice output."""

    def paragraphs() -> Any:
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    for paragraph in paragraphs():
        replacement = LEGACY_TEXT_REPLACEMENTS.get(paragraph.text)
        if replacement is None:
            continue
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in tuple(paragraph.runs[1:]):
                run._element.getparent().remove(run._element)
        else:
            paragraph.add_run(replacement)

    remaining = sorted(
        {
            paragraph.text
            for paragraph in paragraphs()
            if any("\u3400" <= character <= "\u9fff" for character in paragraph.text)
        }
    )
    if remaining:
        raise ValueError(
            "Legacy Results TD contains untranslated CJK text: " + "; ".join(remaining)
        )


def _configure_append_section(document: Any) -> None:
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.start_type = WD_SECTION.NEW_PAGE
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    if "Report Kicker" not in document.styles:
        kicker = document.styles.add_style("Report Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = document.styles["Report Kicker"]
    kicker.font.name = "Arial Unicode MS"
    kicker._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    kicker._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    kicker._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    kicker.font.size = Pt(9.5)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string("5C6975")
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(4)

    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        header_run = paragraph.add_run(
            "TD-JEPA · Actor-Free TD-LeWM raw-action V0 / action-encoder V1 · Cube O50"
        )
        _set_run_font(header_run, size=8.5, color="6B7280")
    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = paragraph.add_run("Validated result archive · Page ")
        _set_run_font(footer_run, size=8.5, color="6B7280")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        paragraph._p.append(field)


def build_results_document(inputs: ValidatedReportInputs) -> bytes:
    """Return one DOCX byte stream built only from validated inputs."""

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to build Results TD V1; use the Codex "
            "workspace document runtime."
        ) from exc

    summary = inputs.summary
    v0_summary = inputs.v0_summary
    study = _mapping(summary["study"], context="summary.study")
    methods = _mapping(summary["methods"], context="summary.methods")
    v0_methods = _mapping(v0_summary["methods"], context="v0_summary.methods")
    ranking = summary["ranking_by_f_plus_g"]
    acceptance = _mapping(
        summary["training_acceptance"], context="summary.training_acceptance"
    )
    acceptance_status = acceptance["status"]
    acceptance_warnings = tuple(acceptance.get("warnings", ()))
    document = Document(str(inputs.base_document))
    _configure_append_section(document)

    kicker = document.add_paragraph(style="Report Kicker")
    kicker.add_run("RESULTS TD / CONTROLLED C–G3 V0 → V1")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Raw-action V0 vs shared-action-encoder V1")
    _set_run_font(title_run, size=24, color="0B2545", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "Same C/D/F/G1/G2/G3 losses · same 50 pairs · actor-free CEM planning"
    )
    _set_run_font(subtitle_run, size=12.5, color="4B5563")
    document.add_paragraph(
        "Version boundary: the first four pages preserve the earlier 7-method "
        "Actor-Free TD-LeWM ablation. This appendix records the later controlled "
        "C-G3 study; its raw-action V0 is a separate 6-by-3 experiment."
    )
    document.add_heading("V0 / V1 definition", level=1)
    document.add_paragraph(
        "V0 sends the normalized raw 25D action block directly into G. V1 sends "
        "stopgrad(E_A(a)), the 192D embedding from the shared frozen LeWM Action "
        "Encoder. Both current and dataset next actions follow this rule; neither "
        "version contains an Actor or reward model. The six losses and planning "
        "definitions are unchanged, but predictor capacity is not matched: "
        f"{V0_PREDICTOR_PARAMETERS:,} parameters in V0 versus "
        f"{PREDICTOR_PARAMETERS:,} in V1."
    )

    document.add_heading("V0 formal O50 — raw 25D action", level=1)
    document.add_paragraph(
        "V0 has 18 completed formal evaluations: six methods × F-only/G-only/F+G. "
        "All rows use training seed 3072, planning seed 42, horizons 5/1/5 and the "
        "same 50 start–goal pairs used by V1."
    )
    v0_rows = []
    for variant in VARIANT_ORDER:
        evaluations = v0_methods[variant]
        f_success = int(evaluations["f_only"]["success_count"])
        combined_success = int(evaluations["f_plus_g"]["success_count"])
        v0_rows.append(
            (
                variant.upper(),
                _format_v0_result(evaluations["f_only"]),
                _format_v0_result(evaluations["g_only"]),
                _format_v0_result(evaluations["f_plus_g"]),
                f"{combined_success - f_success:+d}/50",
            )
        )
    _add_table(
        document,
        headers=("V0 method", "F-only", "G-only", "F+G", "Δ vs F"),
        rows=v0_rows,
        widths=(4800, 2400, 2400, 2400, 2400),
    )
    document.add_heading("V0 → V1 combined comparison", level=1)
    document.add_paragraph(
        "On this matched selection: F and G3 gain 8 pp; D and G2 are unchanged; "
        "C loses 4 pp; G1 loses 2 pp. V0 best is G1/G2 at 50%; V1 best is G3 at 54%."
    )
    comparison_rows = []
    for variant in VARIANT_ORDER:
        v0_count = int(v0_methods[variant]["f_plus_g"]["success_count"])
        v1_count = int(methods[variant]["evaluations"]["f_plus_g"]["success_count"])
        comparison_rows.append(
            (
                variant.upper(),
                f"{v0_count}/50 ({v0_count * 2}%)",
                f"{v1_count}/50 ({v1_count * 2}%)",
                f"{v1_count - v0_count:+d}/50",
            )
        )
    _add_table(
        document,
        headers=("Method", "V0 F+G", "V1 F+G", "V1 - V0"),
        rows=comparison_rows,
        widths=(3600, 3600, 3600, 3600),
    )
    kicker = document.add_paragraph(style="Report Kicker")
    kicker.paragraph_format.page_break_before = True
    kicker.add_run("RESULTS TD / ACTOR-FREE TD-JEPA V1 DETAILS")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Cube O50 — six methods × three score modes")
    _set_run_font(title_run, size=24, color="0B2545", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle_run = subtitle.add_run(
        "Shared frozen LeWM Action Encoder · actor-free CEM planning · seed 3072"
    )
    _set_run_font(subtitle_run, size=12.5, color="4B5563")
    metadata = (
        ("Study", _nonempty_text(study.get("id"), context="summary.study.id")),
        ("Evaluation", "18 formal O50 runs on one matched set of 50 pairs"),
        ("Training", "6 methods · 10 epochs · 127,960 updates per method"),
        ("Training acceptance", acceptance_status),
        ("Selection SHA-256", inputs.selection_sha256),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, size=10.5, color="111827", bold=True)
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, size=10.5, color="111827")

    lead = document.add_paragraph()
    lead.paragraph_format.space_before = Pt(10)
    lead.paragraph_format.space_after = Pt(10)
    lead_run = lead.add_run(
        "Interpretation boundary — one training seed and one planning selection. "
        "Use this report as a matched structural/score-mode ablation, not as a "
        "multi-seed claim of overall superiority."
    )
    _set_run_font(lead_run, size=10.5, color="7A5A00", bold=True)
    if acceptance_status == "PASS_WITH_WARNINGS":
        warning = document.add_paragraph()
        warning.paragraph_format.space_after = Pt(10)
        warning_run = warning.add_run(
            "Training provenance warning — launcher exit codes could not be "
            "recovered. Checkpoints, metrics, epochs, steps, and acceptance "
            "evidence passed the archive validator, but the missing exit-code "
            "provenance must not be represented as an unconditional PASS. "
            f"Archived warning: {'; '.join(acceptance_warnings)}"
        )
        _set_run_font(warning_run, size=10.5, color="9C2F17", bold=True)

    document.add_heading("Combined ranking", level=1)
    best = ranking[0]
    best_name = methods[best["variant"]]["display_name"]
    best_run = methods[best["variant"]]["evaluations"]["f_plus_g"]
    paragraph = document.add_paragraph(
        f"Highest combined result: {best_name}, {_format_result(best_run)}. "
        "Ties share the same success count; ordering is taken from the validated summary."
    )
    paragraph.paragraph_format.keep_with_next = True
    ranking_rows = []
    for row in ranking:
        variant = row["variant"]
        run = methods[variant]["evaluations"]["f_plus_g"]
        ranking_rows.append(
            (
                str(row["rank"]),
                methods[variant]["display_name"],
                variant.upper(),
                "F+G",
                _format_result(run),
            )
        )
    _add_table(
        document,
        headers=("Rank", "Method", "ID", "Score", "O50"),
        rows=ranking_rows,
        widths=(900, 6000, 900, 1800, 4800),
    )

    document.add_page_break()
    document.add_heading("Three inference scores", level=1)
    document.add_paragraph(
        "F-only uses the LeWM rollout score; G-only uses the learned successor "
        "readout; F+G uses the LeWM prefix with the successor tail. All three "
        "columns use the same checkpoint and the same 50 start–goal pairs within "
        "each method."
    )
    score_rows = []
    for variant in VARIANT_ORDER:
        method = methods[variant]
        evaluations = method["evaluations"]
        f_success = evaluations["f_only"]["success_count"]
        combined_success = evaluations["f_plus_g"]["success_count"]
        score_rows.append(
            (
                method["display_name"],
                _format_result(evaluations["f_only"]),
                _format_result(evaluations["g_only"]),
                _format_result(evaluations["f_plus_g"]),
                f"{combined_success - f_success:+d}/50",
            )
        )
    _add_table(
        document,
        headers=("Method", "F-only", "G-only", "F+G", "Δ vs F"),
        rows=score_rows,
        widths=(5200, 2300, 2300, 2300, 2300),
    )

    document.add_page_break()
    document.add_heading("Training curves", level=1)
    document.add_paragraph(
        "The left panel is each method's own training objective and is not "
        "cross-method comparable. The right panel is the common base TD validation "
        "objective used by all six methods."
    )
    chart_paragraph = document.add_paragraph()
    chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_paragraph.paragraph_format.keep_with_next = True
    chart_shape = chart_paragraph.add_run().add_picture(
        str(inputs.loss_chart), width=Inches(9.25)
    )
    chart_shape._inline.docPr.set(
        "descr",
        "Two-panel line chart for V1 C, D, F, G1, G2 and G3 across epochs 1 "
        "through 10. The left panel uses a log scale for each method-specific "
        "training objective; the right panel shows the common base-TD validation "
        "loss on a linear scale.",
    )
    chart_shape._inline.docPr.set("title", "V1 training and validation loss curves")
    caption = document.add_paragraph(
        "Figure 1. Train method objective (left) and validation common base TD (right)."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(9)

    training_rows = []
    for variant in VARIANT_ORDER:
        final_train, final_validation, best_validation, best_epoch = (
            _curve_final_and_best(methods[variant])
        )
        training_rows.append(
            (
                methods[variant]["display_name"],
                f"{final_train:.6g}",
                f"{final_validation:.6g}",
                f"{best_validation:.6g} (E{best_epoch})",
            )
        )
    _add_table(
        document,
        headers=(
            "Method",
            "E10 train\nmethod objective",
            "E10 validation\ncommon base TD",
            "Best common\nbase TD",
        ),
        rows=training_rows,
        widths=(5200, 3000, 3100, 3100),
    )

    method_definitions_heading = document.add_heading("Method definitions", level=1)
    method_definitions_heading.paragraph_format.page_break_before = True
    for variant in VARIANT_ORDER:
        method = methods[variant]
        method_heading = document.add_heading(
            f"{method['display_name']} ({variant.upper()})", level=2
        )
        method_heading.paragraph_format.keep_with_next = True
        inference = _mapping(
            method["inference"], context=f"summary.methods.{variant}.inference"
        )
        definitions = (
            ("Network", method["network"]),
            ("Training loss", method["training_loss"]),
            ("Special design", method["special_mechanism"]),
            (
                "Inference",
                "; ".join(
                    (
                        f"F-only: {inference['f_only']}",
                        f"G-only: {inference['g_only']}",
                        f"F+G: {inference['f_plus_g']}",
                    )
                ),
            ),
        )
        for definition_index, (label, value) in enumerate(definitions):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.keep_with_next = (
                definition_index < len(definitions) - 1
            )
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            paragraph.add_run(str(value))

    document.add_heading("Paired-outcome audit", level=1)
    document.add_paragraph(
        f"The paired archive contains exactly {len(inputs.paired_rows)} rows and "
        "18 Boolean outcome columns. Every pair uses goal_step − start_step = 50; "
        "the summary counts were recomputed from these columns before this document "
        "was written."
    )
    audit_rows = (
        ("Training seed", str(TRAINING_SEED)),
        ("Planning seed", str(PLANNING_SEED)),
        ("Episodes / mode", str(EPISODES)),
        ("Goal offset", str(GOAL_OFFSET)),
        ("Selection SHA-256", inputs.selection_sha256),
    )
    _add_table(
        document,
        headers=("Audit field", "Validated value"),
        rows=audit_rows,
        widths=(3600, 10800),
    )

    document.add_heading("Claim boundary", level=1)
    document.add_paragraph(
        "This report records one formal training seed and one shared O50 planning "
        "selection. It supports within-selection architecture and inference-score "
        "comparisons. It does not establish mean performance, variance, or "
        "statistical superiority across independent training seeds."
    )

    _translate_legacy_text_for_render(document)
    stream = io.BytesIO()
    document.save(stream)
    payload = stream.getvalue()
    if not payload.startswith(b"PK"):
        raise RuntimeError("python-docx did not produce a valid OOXML container.")
    return payload


def _atomic_write(path: Path, payload: bytes) -> None:
    if path.suffix.lower() != ".docx":
        raise ValueError(f"DOCX output path must end in .docx: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{path.name}.", suffix=".tmp", dir=path.parent
    )
    temporary = Path(temporary_name)
    try:
        with os.fdopen(descriptor, "wb") as stream:
            stream.write(payload)
            stream.flush()
            os.fsync(stream.fileno())
        temporary.replace(path)
    finally:
        if temporary.exists():
            temporary.unlink()


def write_results_documents(
    payload: bytes,
    *,
    repository_output: str | Path,
    project_copy: str | Path | None,
) -> tuple[Path, ...]:
    destinations = [Path(repository_output)]
    if project_copy is not None:
        candidate = Path(project_copy)
        if candidate.resolve() != destinations[0].resolve():
            destinations.append(candidate)
    for destination in destinations:
        _atomic_write(destination, payload)
    return tuple(destinations)


def main() -> None:
    args = parse_args()
    try:
        base_path = Path(args.base_document).resolve()
        destinations = [Path(args.output).resolve()]
        if not args.no_project_copy:
            destinations.append(Path(args.project_copy).resolve())
        if base_path in destinations:
            raise ValueError(
                "The immutable legacy Results TD base cannot also be an output."
            )
        inputs = load_validated_report_inputs(
            summary_path=args.summary,
            paired_path=args.paired,
            loss_chart_path=args.loss_chart,
            v0_summary_path=args.v0_summary,
            base_document_path=args.base_document,
        )
        payload = build_results_document(inputs)
        paths = write_results_documents(
            payload,
            repository_output=args.output,
            project_copy=None if args.no_project_copy else args.project_copy,
        )
    except (FileNotFoundError, OSError, RuntimeError, ValueError) as exc:
        raise SystemExit(str(exc)) from exc
    print("Wrote validated Results TD legacy + V0/V1 comparison:")
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
