#!/usr/bin/env python3
"""Build the complete Results TD ledger report from 477 audited O50 cells.

The existing Results TD reports were produced incrementally.  This builder is
the consolidation layer: it requires the exact legacy/V0/V1/V2/V2-EMA grid,
checks every identity and rate, then produces a standalone Markdown report and
a standalone DOCX which presents one 26-method, seven-score decision matrix.
The original 477-cell fixed-E10 analysis stays intact while an optional strict
eight-cell V1-C2/C3 endpoint ledger fills the added First-Q2 and State-V cells.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import io
import json
import math
import os
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from types import ModuleType
from typing import Any, Iterable, Mapping, Sequence

SCRIPT_DIR = Path(__file__).resolve().parent
REPOSITORY_ROOT = SCRIPT_DIR.parent


def _load_v2_report_builder() -> ModuleType:
    module_name = "_tdwm_results_td_v2_ema_builder"
    spec = importlib.util.spec_from_file_location(
        module_name, SCRIPT_DIR / "build_results_td_v2_ema_new_scores.py"
    )
    if spec is None or spec.loader is None:
        raise ImportError("Cannot load the V2-EMA Results TD builder.")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


v2_report = _load_v2_report_builder()


DEFAULT_LEDGER = (
    REPOSITORY_ROOT
    / "reports/artifacts/actor_free_td_lewm_complete_cube_seed3072/all_o50_results.csv"
)
DEFAULT_RECONCILIATION_LEDGER = DEFAULT_LEDGER.parent / "reconciliation_ledger.json"
DEFAULT_ENDPOINT_EXTENSION_LEDGER = (
    REPOSITORY_ROOT
    / "reports/artifacts/actor_free_td_lewm_v1_c2_c3_cube_seed3072/reconciliation_ledger.json"
)
DEFAULT_MARKDOWN_OUTPUT = (
    REPOSITORY_ROOT / "reports/actor_free_td_lewm_complete_cube_seed3072.md"
)
DEFAULT_DOCX_OUTPUT = (
    REPOSITORY_ROOT
    / "reports/results_td_actor_free_td_lewm_complete_cube_seed3072.docx"
)
DEFAULT_ROOT_DOCX_OUTPUT = REPOSITORY_ROOT.parents[1] / "Results TD.docx"
DEFAULT_BASE_DOCUMENT = v2_report.DEFAULT_BASE_DOCUMENT
DEFAULT_CHART_DIR = (
    REPOSITORY_ROOT
    / "reports/artifacts/actor_free_td_lewm_complete_cube_seed3072/figures"
)
ENDPOINT_TRAINING_CHART_FILENAME = "v1_c2_c3_training_validation_diagnostics.png"
DEFAULT_V0_TRAINING_CSV = DEFAULT_LEDGER.parent / "v0_training_loss_curves.csv"
DEFAULT_V2_TRAINING_CSV = DEFAULT_LEDGER.parent / "v2_training_loss_curves.csv"

EPISODES = 50
COMPLETE_CELL_COUNT = 477
COMPLETE_OUTCOME_COUNT = COMPLETE_CELL_COUNT * EPISODES
TRAINING_SEED = 3072
PLANNING_SEED = 42
STEPS_PER_EPOCH = 12_796
SELECTION_SHA256 = v2_report.SHARED_EPISODE_SELECTION_SHA256
ACTION_NORMALIZATION_SHA256 = v2_report.ACTION_NORMALIZATION_SHA256
FIXED_SELECTION_RANKS_SHA256 = (
    "88c204770f33c0b0220057d45b187766e3cfc54912e3f5ca49f2aa93d16437e9"
)
VARIANTS = ("c", "d", "f", "g1", "g2", "g3")
VERSIONS = ("v0", "v1", "v2", "v2_ema_sg")
VERSION_LABELS = {
    "v0": "V0",
    "v1": "V1",
    "v2": "V2",
    "v2_ema_sg": "V2-EMA",
}
# Compact method objectives for the fixed results matrix.  The exact stopped
# signals, weighting normalization and branch definitions remain in the
# detailed loss table later in the report.
METHOD_LOSS_LABELS = {
    "c": "L_C=mean(l)+mean_goal(q-qY)^2",
    "d": "L_D=mean_i w_i[sg(qY)]l_i",
    "f": "L_F=mean_i w_i(A_goal)l_i",
    "g1": "L_G1=mean_i w_i(A_neighbor)l_i",
    "g2": "L_G2=mean_i w_i(A_prefix-mean)l_i",
    "g3": "L_G3=mean_i w_i(A_prefix-gain)l_i",
    "c2": "L_C2=L_C+CE(p_F,p_Qfirst)",
    "c3": "L_C3=mean_i omega_tau(r_i)Huber_1(r_i)",
}
ORIGINAL_MODES = ("f_only", "g_only", "f_plus_g")
NEW_MODES = ("f_plus_g_first", "g_only_f_rollout_mean")
ALL_CONTROLLED_MODES = ORIGINAL_MODES + NEW_MODES
FIRST_Q2_MODE = "f_plus_g_first_q2"
STATE_V_MODE = "state_v_terminal"
MASTER_SCORE_MODES = ALL_CONTROLLED_MODES + (FIRST_Q2_MODE, STATE_V_MODE)
# C2 and C3 are continuations of V1-C and therefore sit beside C inside the
# V1 band.  Keeping the row membership explicit lets winner highlighting and
# version-label merging follow the actual group sizes rather than assuming
# every version contains six rows.
MASTER_ROW_GROUPS = (
    ("v0", VARIANTS),
    ("v1", ("c", "c2", "c3", "d", "f", "g1", "g2", "g3")),
    ("v2", VARIANTS),
    ("v2_ema_sg", VARIANTS),
)
MASTER_ROW_COUNT = sum(len(variants) for _version, variants in MASTER_ROW_GROUPS)
EMA_EPOCHS = tuple(range(3, 11))
LEGACY_METHODS = (
    "serial_decoupled",
    "serial_coupled",
    "hybrid",
    "parallel_real",
    "goal_hybrid",
    "imaginary_hybrid",
    "direct_goal_hybrid",
)
LEGACY_LABELS = {
    "serial_decoupled": "Serial Decoupled",
    "serial_coupled": "Serial Coupled",
    "hybrid": "Hybrid",
    "parallel_real": "Parallel Real",
    "goal_hybrid": "Goal Hybrid",
    "imaginary_hybrid": "Imaginary Hybrid",
    "direct_goal_hybrid": "Direct Goal Critic Hybrid",
}
MODE_LABELS = {
    "f_only": "F-only",
    "g_only": "G-only",
    "f_plus_g": "F+G tail",
    "c_only": "C-only",
    "f_plus_c": "F+C",
    "f_plus_g_first": "F + first-Q",
    "g_only_f_rollout_mean": "Mean-Q rollout",
    FIRST_Q2_MODE: "First-Q2",
    STATE_V_MODE: "State-V terminal",
}

# The seven columns in the endpoint-ready decision matrix share one protocol.
# Keep these definitions next to the display labels so the DOCX, Markdown and
# tests cannot silently drift apart.
COMMON_EVALUATION_PROTOCOL_ROWS = (
    ("Environment", "swm/OGBCube-v0"),
    ("Formal pairs", "The same 50 same-episode start-goal pairs; goal offset 50"),
    ("CEM", "300 candidates, 30 iterations, 30 elites, planning seed 42, warm start"),
    ("Execution", "Minimize candidate cost, execute only the first action block A1, then observe and replan"),
    ("Episode success", "Object-to-goal distance <= 0.04 m within 100 environment steps"),
    ("Checkpoint", "Base rows use fixed E10; C2 uses final E10; C3 uses final E12. No score-specific retraining"),
)

EVALUATION_METHOD_ROWS = (
    (
        "F-only",
        "F rolls A1...A5 from real z0 and produces imagined z1^F...z5^F; G is not called.",
        "J_F = ||z5^F - z_g||_2^2",
        "Uses terminal goal distance at z5; no Q and no gamma.",
    ),
    (
        "G-only",
        "H=1. G scores the real z0 and first candidate action A1; F is not rolled out.",
        "J_G = -Q_G(z0,A1,g)",
        "No explicit goal distance and no gamma; minimizing -Q maximizes Q.",
    ),
    (
        "F+G tail",
        "F rolls only A1...A4 to z4^F; G evaluates the fifth transition from z4^F with A5.",
        "J_tail = ||z4^F - z_g||_2^2 - gamma^4 Q_G(z4^F,A5,g)",
        "Uses z4 goal distance and the deepest imagined-state Q; gamma=0.95.",
    ),
    (
        "F + first-Q",
        "F completes the five-step rollout; G is read only once at the real z0 with A1.",
        "J_first = ||z5^F - z_g||_2^2 - 0.25 Q_G(z0,A1,g)",
        "Uses terminal goal distance; the Q term is not multiplied by gamma^4.",
    ),
    (
        "Mean-Q rollout",
        "F generates predecessors z0,z1^F,...,z4^F; G scores each aligned pair (z{k-1}^F,Ak).",
        "J_mean = -(1/5) sum[k=1..5] Q_G(z{k-1}^F,Ak,g)",
        "No terminal goal distance; z5 is not read by G and gamma is unused.",
    ),
    (
        "First-Q2",
        "F completes the five-step rollout; G is read once at real z0 with A1. Each candidate set normalizes F-cost and first-Q separately.",
        "J_first2 = zscore_candidates(J_F) - 0.25 zscore_candidates(Q_G(z0,A1,g))",
        "No gamma. Population z-score statistics are recomputed inside each CEM candidate set; they never persist across iterations or episodes.",
    ),
    (
        "State-V terminal",
        "Frozen F completes all five blocks to z5^F; only the EMA State-V critic reads (z5^F,z_g). G is not called.",
        "J_V = V_bar(z5^F,z_g)",
        "No terminal latent L2, G term, actor or gamma factor is added at inference; CEM minimizes predicted temporal cost-to-go.",
    ),
)
LEDGER_SCORE_LABELS = {
    "f_only": "F-only",
    "g_only": "G-only",
    "f_plus_g": "F+G",
    "c_only": "C-only",
    "f_plus_c": "F+C",
    "f_plus_g_first": "F+G First (alpha=0.25)",
    "g_only_f_rollout_mean": "Mean-Q over F rollout",
    FIRST_Q2_MODE: "First-Q2 (candidate-z, alpha=0.25)",
    STATE_V_MODE: "EMA State-V at F terminal",
}

FIELDS = (
    "cell_id",
    "family",
    "version",
    "method_id",
    "method_label",
    "variant",
    "epoch",
    "global_step",
    "checkpoint_sha256",
    "score_mode",
    "score_label",
    "success_count",
    "episode_count",
    "success_rate",
    "success_rate_percent",
    "training_seed",
    "planning_seed",
    "selection_sha256",
    "action_normalization_sha256",
    "training_commit",
    "evaluation_commit",
    "planning_horizon",
    "g_first_weight",
    "status",
    "comparison_role",
    "source_kind",
    "source_path",
    "source_results_sha256",
    "source_protocol_sha256",
    "outcomes_sha256",
)


class CompleteResultsError(ValueError):
    """The complete Results TD evidence is missing or inconsistent."""


@dataclass(frozen=True)
class ResultCell:
    values: Mapping[str, Any]

    def __getitem__(self, key: str) -> Any:
        return self.values[key]


@dataclass(frozen=True)
class FixedAnalysis:
    """Derived fixed-E10 comparisons used by every report surface."""

    mode_means: Mapping[str, float]
    mode_best_counts: Mapping[str, int]
    mode_winners: Mapping[str, tuple[str, ...]]
    overall_best_count: int
    overall_winners: tuple[str, ...]
    best_mean_modes: tuple[str, ...]
    best_mean_percent: float
    version_mode_means: Mapping[tuple[str, str], float]
    training_row_means: Mapping[tuple[str, str], float]
    best_training_mean_percent: float
    best_training_mean_winners: tuple[str, ...]
    ema_variant_means: Mapping[str, float]
    ema_best_variants: tuple[str, ...]


@dataclass(frozen=True)
class LedgerEvidence:
    """Companion files that preserve the full scalar and per-pair audit trail."""

    csv_path: str
    csv_sha256: str
    json_path: str
    json_sha256: str
    cell_count: int
    outcome_count: int


@dataclass(frozen=True)
class EndpointAnalysis:
    """Evidence-derived C2/C3 summary; absent runs remain explicit placeholders."""

    available: bool
    cell_count: int
    c2_shared_deltas: Mapping[str, int]
    c2_improved: int
    c2_tied: int
    c2_harmed: int
    c2_mean_delta_percent: float | None
    first_q2_parent_count: int | None
    first_q2_c2_count: int | None
    state_v_c3_count: int | None
    state_v_c3_epoch3_count: int | None
    state_v_c3_epoch3_vs_epoch12_contingency: Mapping[str, int] | None
    state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p: float | None


@dataclass(frozen=True)
class EndpointTrainingEvidence:
    """Hash-bound C2/C3 epoch aggregates used by the endpoint diagnostic chart."""

    c2_metrics_path: str
    c2_metrics_sha256: str
    c2_epochs: tuple[int, ...]
    c2_train_loss: tuple[float, ...]
    c2_validation_loss: tuple[float, ...]
    c2_alignment_loss: tuple[float, ...]
    c2_alignment_top1: tuple[float, ...]
    c3_metrics_path: str
    c3_metrics_sha256: str
    c3_epochs: tuple[int, ...]
    c3_train_loss: tuple[float, ...]
    c3_validation_loss: tuple[float, ...]
    c3_mc_mae: tuple[float, ...]
    c3_td_residual_mae: tuple[float, ...]
    c3_spearman: tuple[float, ...]


def _sha(value: str, *, context: str, allow_empty: bool = False) -> str:
    value = value.strip()
    if allow_empty and value == "":
        return value
    if len(value) != 64 or any(
        character not in "0123456789abcdef" for character in value
    ):
        raise CompleteResultsError(f"{context} must be a lowercase SHA-256.")
    return value


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _display_path(path: Path) -> str:
    resolved = path.resolve()
    try:
        return resolved.relative_to(REPOSITORY_ROOT.resolve()).as_posix()
    except ValueError:
        return str(resolved)


def load_ledger_evidence(csv_path: str | Path, json_path: str | Path) -> LedgerEvidence:
    """Validate and fingerprint the two files that preserve the complete ledger."""

    csv_source = Path(csv_path)
    json_source = Path(json_path)
    if not csv_source.is_file():
        raise FileNotFoundError(f"Complete CSV ledger does not exist: {csv_source}")
    if not json_source.is_file():
        raise FileNotFoundError(
            f"Reconciliation JSON ledger does not exist: {json_source}"
        )

    with csv_source.open(newline="", encoding="utf-8") as stream:
        csv_cell_count = sum(1 for _ in csv.DictReader(stream))
    if csv_cell_count != COMPLETE_CELL_COUNT:
        raise CompleteResultsError(
            "Companion CSV must contain "
            f"{COMPLETE_CELL_COUNT} cells, found {csv_cell_count}."
        )

    try:
        reconciliation = json.loads(json_source.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise CompleteResultsError(
            f"Reconciliation ledger is not valid JSON: {json_source}."
        ) from exc
    if not isinstance(reconciliation, dict):
        raise CompleteResultsError("Reconciliation ledger root must be an object.")
    if reconciliation.get("cell_count") != COMPLETE_CELL_COUNT:
        raise CompleteResultsError(
            f"Reconciliation ledger cell_count must be {COMPLETE_CELL_COUNT}."
        )
    if reconciliation.get("episode_count_per_cell") != EPISODES:
        raise CompleteResultsError(
            f"Reconciliation ledger must preserve {EPISODES} outcomes per cell."
        )
    if reconciliation.get("outcome_count") != COMPLETE_OUTCOME_COUNT:
        raise CompleteResultsError(
            f"Reconciliation ledger outcome_count must be {COMPLETE_OUTCOME_COUNT}."
        )
    json_cells = reconciliation.get("cells")
    if not isinstance(json_cells, list) or len(json_cells) != COMPLETE_CELL_COUNT:
        raise CompleteResultsError(
            "Reconciliation ledger cells must contain exactly "
            f"{COMPLETE_CELL_COUNT} entries."
        )
    for index, cell in enumerate(json_cells):
        if not isinstance(cell, dict):
            raise CompleteResultsError(
                f"reconciliation.cells[{index}] is not an object."
            )
        outcomes = cell.get("outcomes")
        if (
            not isinstance(outcomes, list)
            or len(outcomes) != EPISODES
            or any(type(value) is not bool for value in outcomes)
        ):
            raise CompleteResultsError(
                f"reconciliation.cells[{index}].outcomes must contain "
                f"{EPISODES} Booleans."
            )

    return LedgerEvidence(
        csv_path=_display_path(csv_source),
        csv_sha256=_file_sha256(csv_source),
        json_path=_display_path(json_source),
        json_sha256=_file_sha256(json_source),
        cell_count=csv_cell_count,
        outcome_count=sum(len(cell["outcomes"]) for cell in json_cells),
    )


def _load_epoch_metric_series(
    path: Path,
    *,
    method_key: str,
    expected_epoch_count: int,
    columns: Sequence[str],
) -> Mapping[str, tuple[float, ...]]:
    """Load one finite epoch aggregate per required Lightning metric."""

    context = f"endpoint training metrics {method_key}"
    try:
        with path.open(newline="", encoding="utf-8") as stream:
            reader = csv.DictReader(stream)
            fieldnames = tuple(reader.fieldnames or ())
            rows = tuple(dict(row) for row in reader)
    except (OSError, UnicodeError, csv.Error) as exc:
        raise CompleteResultsError(f"{context} is unreadable: {path}.") from exc
    required = {"epoch", *columns}
    missing = sorted(required - set(fieldnames))
    if missing:
        raise CompleteResultsError(
            f"{context} is missing required columns: {missing}."
        )
    if not rows:
        raise CompleteResultsError(f"{context} contains no rows.")

    expected_epochs = tuple(range(expected_epoch_count))
    parsed: dict[str, tuple[float, ...]] = {}
    for column in columns:
        by_epoch: dict[int, float] = {}
        for row_index, row in enumerate(rows):
            raw_value = row.get(column, "")
            if raw_value in (None, ""):
                continue
            raw_epoch = row.get("epoch", "")
            if raw_epoch in (None, ""):
                raise CompleteResultsError(
                    f"{context} row {row_index} has {column} without an epoch."
                )
            epoch = _int(str(raw_epoch), context=f"{context}.epoch")
            if epoch not in expected_epochs:
                raise CompleteResultsError(
                    f"{context}.{column} has out-of-range epoch {epoch}."
                )
            if epoch in by_epoch:
                raise CompleteResultsError(
                    f"{context}.{column} has duplicate aggregate at epoch {epoch + 1}."
                )
            by_epoch[epoch] = _float(
                str(raw_value), context=f"{context}.{column}.epoch_{epoch + 1}"
            )
        if tuple(sorted(by_epoch)) != expected_epochs:
            missing_epochs = [
                epoch + 1 for epoch in expected_epochs if epoch not in by_epoch
            ]
            raise CompleteResultsError(
                f"{context}.{column} is missing logical epochs {missing_epochs}."
            )
        parsed[column] = tuple(by_epoch[epoch] for epoch in expected_epochs)
    return parsed


def _require_metric_range(
    values: Sequence[float],
    *,
    context: str,
    lower: float,
    upper: float | None = None,
    lower_inclusive: bool = True,
) -> None:
    for value in values:
        lower_ok = value >= lower if lower_inclusive else value > lower
        upper_ok = upper is None or value <= upper
        if not lower_ok or not upper_ok:
            interval = (
                f"{'[' if lower_inclusive else '('}{lower}, {upper}]"
                if upper is not None
                else f">{'=' if lower_inclusive else ''} {lower}"
            )
            raise CompleteResultsError(
                f"{context} must stay in {interval}; found {value}."
            )


def load_endpoint_training_evidence(
    endpoint_ledger_path: str | Path,
) -> EndpointTrainingEvidence | None:
    """Load C2/C3 curves only when both hash-bound training sources are declared.

    The strict endpoint loader audits the archive as a whole.  This report-level
    reader repeats the metrics-specific checks so a malformed, duplicated or
    incomplete epoch aggregate can never be turned into a plausible-looking
    chart.  An endpoint ledger with no declared training sources is valid but
    produces no chart.
    """

    supplied = Path(endpoint_ledger_path).expanduser().resolve()
    ledger_path = (
        supplied / "reconciliation_ledger.json" if supplied.is_dir() else supplied
    )
    try:
        ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        raise
    except (OSError, UnicodeError, json.JSONDecodeError) as exc:
        raise CompleteResultsError(
            f"Endpoint extension ledger is unreadable: {ledger_path}."
        ) from exc
    if not isinstance(ledger, Mapping):
        raise CompleteResultsError("Endpoint extension ledger root must be an object.")
    raw_sources = ledger.get("training_sources")
    if not isinstance(raw_sources, Mapping):
        raise CompleteResultsError(
            "Endpoint extension training_sources must be an object."
        )
    if not raw_sources:
        return None
    expected_methods = {"v1_c2", "v1_c3"}
    if set(raw_sources) != expected_methods:
        raise CompleteResultsError(
            "Declared endpoint training sources must contain both v1_c2 and v1_c3."
        )

    metrics_paths: dict[str, Path] = {}
    metrics_hashes: dict[str, str] = {}
    for method_key in sorted(expected_methods):
        source = raw_sources[method_key]
        if not isinstance(source, Mapping):
            raise CompleteResultsError(
                f"Endpoint training source {method_key} must be an object."
            )
        expected_directory = f"training/{method_key}"
        if source.get("archive_directory") != expected_directory:
            raise CompleteResultsError(
                f"Endpoint training source {method_key} must use {expected_directory}."
            )
        raw_hashes = source.get("files_sha256")
        if not isinstance(raw_hashes, Mapping) or "metrics.csv" not in raw_hashes:
            raise CompleteResultsError(
                f"Endpoint training source {method_key} does not declare metrics.csv."
            )
        metrics_sha = _sha(
            str(raw_hashes["metrics.csv"]),
            context=f"endpoint training source {method_key}.metrics.csv",
        )
        metrics_path = ledger_path.parent / expected_directory / "metrics.csv"
        if not metrics_path.is_file():
            raise CompleteResultsError(
                f"Declared endpoint training metrics are missing: {metrics_path}."
            )
        actual_sha = _file_sha256(metrics_path)
        if actual_sha != metrics_sha:
            raise CompleteResultsError(
                f"Declared endpoint training metrics hash changed for {method_key}."
            )
        metrics_paths[method_key] = metrics_path
        metrics_hashes[method_key] = metrics_sha

    c2 = _load_epoch_metric_series(
        metrics_paths["v1_c2"],
        method_key="v1_c2",
        expected_epoch_count=10,
        columns=(
            "train/loss_epoch",
            "validation/loss",
            "train/first_q_alignment_loss_epoch",
            "train/first_q_alignment_top1_agreement_epoch",
        ),
    )
    c3 = _load_epoch_metric_series(
        metrics_paths["v1_c3"],
        method_key="v1_c3",
        expected_epoch_count=12,
        columns=(
            "train/loss_epoch",
            "validation/loss",
            "validation/mc_mae",
            "validation/td_residual_mae",
            "validation/spearman",
        ),
    )
    for context, values in (
        ("C2 train loss", c2["train/loss_epoch"]),
        ("C2 validation loss", c2["validation/loss"]),
        ("C2 alignment loss", c2["train/first_q_alignment_loss_epoch"]),
        ("C3 train loss", c3["train/loss_epoch"]),
        ("C3 validation loss", c3["validation/loss"]),
    ):
        _require_metric_range(
            values, context=context, lower=0.0, lower_inclusive=False
        )
    for context, values in (
        ("C3 MC MAE", c3["validation/mc_mae"]),
        ("C3 TD residual MAE", c3["validation/td_residual_mae"]),
    ):
        _require_metric_range(values, context=context, lower=0.0)
    _require_metric_range(
        c2["train/first_q_alignment_top1_agreement_epoch"],
        context="C2 First-Q top-1 agreement",
        lower=0.0,
        upper=1.0,
    )
    _require_metric_range(
        c3["validation/spearman"],
        context="C3 validation Spearman",
        lower=-1.0,
        upper=1.0,
    )
    return EndpointTrainingEvidence(
        c2_metrics_path=_display_path(metrics_paths["v1_c2"]),
        c2_metrics_sha256=metrics_hashes["v1_c2"],
        c2_epochs=tuple(range(1, 11)),
        c2_train_loss=c2["train/loss_epoch"],
        c2_validation_loss=c2["validation/loss"],
        c2_alignment_loss=c2["train/first_q_alignment_loss_epoch"],
        c2_alignment_top1=c2[
            "train/first_q_alignment_top1_agreement_epoch"
        ],
        c3_metrics_path=_display_path(metrics_paths["v1_c3"]),
        c3_metrics_sha256=metrics_hashes["v1_c3"],
        c3_epochs=tuple(range(1, 13)),
        c3_train_loss=c3["train/loss_epoch"],
        c3_validation_loss=c3["validation/loss"],
        c3_mc_mae=c3["validation/mc_mae"],
        c3_td_residual_mae=c3["validation/td_residual_mae"],
        c3_spearman=c3["validation/spearman"],
    )


def _int(value: str, *, context: str) -> int:
    try:
        parsed = int(value)
    except ValueError as exc:
        raise CompleteResultsError(f"{context} must be an integer.") from exc
    if str(parsed) != value.strip():
        raise CompleteResultsError(f"{context} is not an exact integer: {value!r}.")
    return parsed


def _float(value: str, *, context: str) -> float:
    try:
        parsed = float(value)
    except ValueError as exc:
        raise CompleteResultsError(f"{context} must be numeric.") from exc
    if not math.isfinite(parsed):
        raise CompleteResultsError(f"{context} must be finite.")
    return parsed


def _optional_int(value: str, *, context: str) -> int | None:
    if value.strip() == "":
        return None
    return _int(value, context=context)


def _expected_identities() -> set[tuple[str, str, int, str]]:
    expected: set[tuple[str, str, int, str]] = set()
    for method in LEGACY_METHODS:
        modes = (
            ("f_only", "c_only", "f_plus_c")
            if method == "direct_goal_hybrid"
            else ORIGINAL_MODES
        )
        expected.update(("legacy", method, 10, mode) for mode in modes)
    for version in ("v0", "v1"):
        expected.update(
            (version, variant, 10, mode)
            for variant in VARIANTS
            for mode in ALL_CONTROLLED_MODES
        )
    expected.update(
        ("v2", variant, epoch, mode)
        for epoch in EMA_EPOCHS
        for variant in VARIANTS
        for mode in ORIGINAL_MODES
    )
    expected.update(
        ("v2", variant, 10, mode) for variant in VARIANTS for mode in NEW_MODES
    )
    expected.update(
        ("v2_ema_sg", variant, epoch, mode)
        for epoch in EMA_EPOCHS
        for variant in VARIANTS
        for mode in ALL_CONTROLLED_MODES
    )
    return expected


def load_complete_ledger(path: str | Path) -> tuple[ResultCell, ...]:
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(f"Complete O50 ledger does not exist: {source}")
    with source.open(newline="", encoding="utf-8") as stream:
        reader = csv.DictReader(stream)
        if tuple(reader.fieldnames or ()) != FIELDS:
            raise CompleteResultsError(
                f"Complete ledger columns changed: {reader.fieldnames!r}."
            )
        raw_rows = tuple(dict(row) for row in reader)
    if len(raw_rows) != COMPLETE_CELL_COUNT:
        raise CompleteResultsError(
            "Complete ledger must contain "
            f"{COMPLETE_CELL_COUNT} cells, found {len(raw_rows)}."
        )

    cells: list[ResultCell] = []
    identities: set[tuple[str, str, int, str]] = set()
    cell_ids: set[str] = set()
    for index, row in enumerate(raw_rows):
        context = f"ledger[{index}]"
        version = row["version"].strip()
        method_key = row["variant"].strip()
        epoch = _int(row["epoch"], context=f"{context}.epoch")
        global_step = _int(row["global_step"], context=f"{context}.global_step")
        mode = row["score_mode"].strip()
        identity = (version, method_key, epoch, mode)
        if identity in identities:
            raise CompleteResultsError(f"Duplicate result identity: {identity}.")
        identities.add(identity)
        cell_id = row["cell_id"].strip()
        if not cell_id or cell_id in cell_ids:
            raise CompleteResultsError(f"{context}.cell_id is empty or duplicated.")
        cell_ids.add(cell_id)
        successes = _int(row["success_count"], context=f"{context}.success_count")
        episodes = _int(row["episode_count"], context=f"{context}.episode_count")
        rate = _float(row["success_rate"], context=f"{context}.success_rate")
        percent = _float(
            row["success_rate_percent"], context=f"{context}.success_rate_percent"
        )
        if episodes != EPISODES or not 0 <= successes <= episodes:
            raise CompleteResultsError(f"{context} must be a valid O50 result.")
        if not math.isclose(rate, successes / episodes, rel_tol=0.0, abs_tol=1e-12):
            raise CompleteResultsError(f"{context}.success_rate disagrees with count.")
        if not math.isclose(percent, 100.0 * rate, rel_tol=0.0, abs_tol=1e-9):
            raise CompleteResultsError(f"{context}.success_rate_percent is stale.")
        if (
            _int(row["training_seed"], context=f"{context}.training_seed")
            != TRAINING_SEED
        ):
            raise CompleteResultsError(f"{context}.training_seed changed.")
        if (
            _int(row["planning_seed"], context=f"{context}.planning_seed")
            != PLANNING_SEED
        ):
            raise CompleteResultsError(f"{context}.planning_seed changed.")
        if row["selection_sha256"] != SELECTION_SHA256:
            raise CompleteResultsError(f"{context} uses a different O50 selection.")
        if row["action_normalization_sha256"] != ACTION_NORMALIZATION_SHA256:
            raise CompleteResultsError(f"{context} action normalization changed.")
        _sha(row["checkpoint_sha256"], context=f"{context}.checkpoint_sha256")
        _sha(row["source_results_sha256"], context=f"{context}.source_results_sha256")
        _sha(
            row["source_protocol_sha256"],
            context=f"{context}.source_protocol_sha256",
            allow_empty=True,
        )
        _sha(row["outcomes_sha256"], context=f"{context}.outcomes_sha256")
        if global_step != epoch * STEPS_PER_EPOCH:
            raise CompleteResultsError(f"{context}.global_step disagrees with epoch.")
        if row["status"] != "VERIFIED":
            raise CompleteResultsError(f"{context}.status must be VERIFIED.")
        expected_family = (
            "actor_free_td_lewm"
            if version == "legacy"
            else f"actor_free_td_lewm_{version}"
        )
        expected_method_id = (
            "actor_free_td_lewm"
            if version == "legacy"
            else f"{expected_family}_{method_key}"
        )
        if row["family"] != expected_family:
            raise CompleteResultsError(
                f"{context}.family is inconsistent with version."
            )
        if row["method_id"] != expected_method_id:
            raise CompleteResultsError(f"{context}.method_id is inconsistent.")
        if (
            not row["method_label"].strip()
            or row["score_label"] != LEDGER_SCORE_LABELS[mode]
        ):
            raise CompleteResultsError(
                f"{context} has a missing or stale display label."
            )
        if not row["training_commit"].strip() or not row["evaluation_commit"].strip():
            raise CompleteResultsError(f"{context} is missing commit provenance.")
        if not row["source_kind"].strip() or not row["source_path"].strip():
            raise CompleteResultsError(f"{context} is missing source provenance.")
        alpha = row["g_first_weight"].strip()
        if mode == "f_plus_g_first":
            if not math.isclose(
                _float(alpha, context=f"{context}.g_first_weight"),
                0.25,
                rel_tol=0.0,
                abs_tol=1e-12,
            ):
                raise CompleteResultsError(f"{context}.g_first_weight changed.")
        elif alpha != "":
            raise CompleteResultsError(f"{context} has an unexpected g_first_weight.")
        planning_horizon = _optional_int(
            row["planning_horizon"], context=f"{context}.planning_horizon"
        )
        if planning_horizon is not None:
            expected_horizon = (
                5 if version == "legacy" else (1 if mode == "g_only" else 5)
            )
            if planning_horizon != expected_horizon:
                raise CompleteResultsError(f"{context}.planning_horizon changed.")
        normalized = dict(row)
        normalized.update(
            {
                "epoch": epoch,
                "global_step": global_step,
                "success_count": successes,
                "episode_count": episodes,
                "success_rate": rate,
                "success_rate_percent": percent,
                "planning_horizon": planning_horizon,
            }
        )
        cells.append(ResultCell(normalized))
    expected = _expected_identities()
    if identities != expected:
        missing = sorted(expected - identities)
        extra = sorted(identities - expected)
        raise CompleteResultsError(
            f"Complete grid mismatch; missing={missing[:8]!r}, extra={extra[:8]!r}."
        )
    checkpoints: dict[tuple[str, str, int], set[str]] = {}
    for cell in cells:
        version = str(cell["version"])
        method_key = str(cell["variant"])
        key = (version, method_key, int(cell["epoch"]))
        checkpoints.setdefault(key, set()).add(str(cell["checkpoint_sha256"]))
    inconsistent = {
        key: values for key, values in checkpoints.items() if len(values) != 1
    }
    if inconsistent:
        raise CompleteResultsError(
            f"Score modes do not share a checkpoint: {next(iter(inconsistent.items()))!r}."
        )
    fixed_mean_q = {
        (str(cell["version"]), str(cell["variant"]))
        for cell in cells
        if int(cell["epoch"]) == 10
        and cell["score_mode"] == "g_only_f_rollout_mean"
        and cell["version"] in VERSIONS
    }
    expected_fixed_mean_q = {
        (version, variant) for version in VERSIONS for variant in VARIANTS
    }
    if fixed_mean_q != expected_fixed_mean_q:
        raise CompleteResultsError(
            "Fixed E10 Mean-Q coverage must contain exactly 24 version/variant cells."
        )
    return tuple(cells)


def _find(
    cells: Iterable[ResultCell],
    *,
    version: str,
    epoch: int,
    mode: str,
    variant: str | None = None,
    method_id: str | None = None,
) -> ResultCell:
    matches = [
        cell
        for cell in cells
        if cell["version"] == version
        and cell["epoch"] == epoch
        and cell["score_mode"] == mode
        and (variant is None or cell["variant"] == variant)
        and (method_id is None or cell["method_id"] == method_id)
    ]
    if len(matches) != 1:
        raise CompleteResultsError(
            f"Expected one cell for {version}/{variant or method_id}/E{epoch}/{mode}, "
            f"found {len(matches)}."
        )
    return matches[0]


def _result(cell: ResultCell) -> str:
    return f"{cell['success_count']}/50 ({cell['success_rate_percent']:.0f}%)"


def _percent(cell: ResultCell) -> float:
    return float(cell["success_rate_percent"])


def _mean(cells: Iterable[ResultCell]) -> float:
    selected = tuple(cells)
    if not selected:
        raise CompleteResultsError("Cannot average an empty result set.")
    return sum(_percent(cell) for cell in selected) / len(selected)


def _fixed_cell_label(cell: ResultCell) -> str:
    return f"{VERSION_LABELS[str(cell['version'])]}-{str(cell['variant']).upper()}"


def _fixed_analysis(cells: Sequence[ResultCell]) -> FixedAnalysis:
    fixed = tuple(
        _find(cells, version=version, variant=variant, epoch=10, mode=mode)
        for version in VERSIONS
        for variant in VARIANTS
        for mode in ALL_CONTROLLED_MODES
    )
    if len(fixed) != len(VERSIONS) * len(VARIANTS) * len(ALL_CONTROLLED_MODES):
        raise CompleteResultsError("Fixed E10 decision grid is incomplete.")
    version_mode_means = {
        (version, mode): _mean(
            _find(cells, version=version, variant=variant, epoch=10, mode=mode)
            for variant in VARIANTS
        )
        for version in VERSIONS
        for mode in ALL_CONTROLLED_MODES
    }
    mode_means = {
        mode: _mean(cell for cell in fixed if cell["score_mode"] == mode)
        for mode in ALL_CONTROLLED_MODES
    }
    mode_best_counts: dict[str, int] = {}
    mode_winners: dict[str, tuple[str, ...]] = {}
    for mode in ALL_CONTROLLED_MODES:
        selected = tuple(cell for cell in fixed if cell["score_mode"] == mode)
        best = max(int(cell["success_count"]) for cell in selected)
        mode_best_counts[mode] = best
        mode_winners[mode] = tuple(
            _fixed_cell_label(cell)
            for cell in selected
            if int(cell["success_count"]) == best
        )
    overall_best = max(int(cell["success_count"]) for cell in fixed)
    overall_winners = tuple(
        f"{_fixed_cell_label(cell)} + {MODE_LABELS[str(cell['score_mode'])]}"
        for cell in fixed
        if int(cell["success_count"]) == overall_best
    )
    best_mean = max(mode_means.values())
    best_mean_modes = tuple(
        mode for mode in ALL_CONTROLLED_MODES if mode_means[mode] == best_mean
    )
    training_row_means = {
        (version, variant): _mean(
            _find(cells, version=version, variant=variant, epoch=10, mode=mode)
            for mode in ALL_CONTROLLED_MODES
        )
        for version in VERSIONS
        for variant in VARIANTS
    }
    best_training_mean = max(training_row_means.values())
    best_training_mean_winners = tuple(
        f"{VERSION_LABELS[version]}-{variant.upper()}"
        for version in VERSIONS
        for variant in VARIANTS
        if training_row_means[(version, variant)] == best_training_mean
    )
    ema_variant_means = {
        variant: sum(
            _percent(
                _find(
                    cells,
                    version="v2_ema_sg",
                    variant=variant,
                    epoch=10,
                    mode=mode,
                )
            )
            for mode in ALL_CONTROLLED_MODES
        )
        / len(ALL_CONTROLLED_MODES)
        for variant in VARIANTS
    }
    ema_best_mean = max(ema_variant_means.values())
    return FixedAnalysis(
        mode_means=mode_means,
        mode_best_counts=mode_best_counts,
        mode_winners=mode_winners,
        overall_best_count=overall_best,
        overall_winners=overall_winners,
        best_mean_modes=best_mean_modes,
        best_mean_percent=best_mean,
        version_mode_means=version_mode_means,
        training_row_means=training_row_means,
        best_training_mean_percent=best_training_mean,
        best_training_mean_winners=best_training_mean_winners,
        ema_variant_means=ema_variant_means,
        ema_best_variants=tuple(
            variant
            for variant in VARIANTS
            if ema_variant_means[variant] == ema_best_mean
        ),
    )


def _joined(values: Sequence[str]) -> str:
    return ", ".join(values)


def _compact_join(values: Sequence[str], *, limit: int = 3) -> str:
    """Keep tied-winner summaries bounded; the matrix retains every highlight."""

    items = tuple(values)
    if len(items) <= limit:
        return _joined(items)
    return f"{len(items)} tied: {_joined(items[:limit])}, …"


def _winner_result(analysis: FixedAnalysis, mode: str) -> str:
    count = analysis.mode_best_counts[mode]
    return f"{_compact_join(analysis.mode_winners[mode])}: {count}/50 ({count * 2}%)"


def _column_winner_summary(analysis: FixedAnalysis) -> str:
    return "; ".join(
        f"{MODE_LABELS[mode]} = {_winner_result(analysis, mode)}"
        for mode in ALL_CONTROLLED_MODES
    )


def _signed_pp(value: float) -> str:
    return f"{value:+.1f} pp"


def _markdown_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> list[str]:
    return [
        "| " + " | ".join(headers) + " |",
        "| "
        + " | ".join("---" if index == 0 else "---:" for index in range(len(headers)))
        + " |",
        *("| " + " | ".join(row) + " |" for row in rows),
    ]


def _legacy_rows(cells: Sequence[ResultCell]) -> tuple[tuple[str, ...], ...]:
    rows = []
    for method in LEGACY_METHODS:
        secondary = "c_only" if method == "direct_goal_hybrid" else "g_only"
        combined = "f_plus_c" if method == "direct_goal_hybrid" else "f_plus_g"
        rows.append(
            (
                LEGACY_LABELS[method],
                _result(
                    _find(
                        cells, version="legacy", variant=method, epoch=10, mode="f_only"
                    )
                ),
                _result(
                    _find(
                        cells,
                        version="legacy",
                        variant=method,
                        epoch=10,
                        mode=secondary,
                    )
                ),
                _result(
                    _find(
                        cells, version="legacy", variant=method, epoch=10, mode=combined
                    )
                ),
            )
        )
    return tuple(rows)


ENDPOINT_EXPECTED_IDENTITIES = frozenset(
    {
        ("v1_c", FIRST_Q2_MODE),
        *(("v1_c2", mode) for mode in (*ALL_CONTROLLED_MODES, FIRST_Q2_MODE)),
        ("v1_c3", STATE_V_MODE),
    }
)
_ENDPOINT_MODE_ALIASES = {
    "first_q2": FIRST_Q2_MODE,
    "state_v": STATE_V_MODE,
}


def _endpoint_field(cell: Any, field: str) -> Any:
    if isinstance(cell, Mapping):
        return cell[field]
    return getattr(cell, field)


def _endpoint_index(endpoint_cells: Sequence[Any] | None) -> dict[tuple[str, str], Any]:
    """Normalize the strict optional eight-cell endpoint extension.

    The dedicated endpoint-ledger loader performs the complete source/protocol/
    outcome audit.  This second, deliberately small guard prevents a caller
    from handing the report builder a partial or duplicated in-memory subset.
    """

    if not endpoint_cells:
        return {}
    indexed: dict[tuple[str, str], Any] = {}
    for cell in endpoint_cells:
        method = str(_endpoint_field(cell, "method_key"))
        raw_mode = str(_endpoint_field(cell, "score_mode"))
        mode = _ENDPOINT_MODE_ALIASES.get(raw_mode, raw_mode)
        identity = (method, mode)
        if identity in indexed:
            raise CompleteResultsError(f"Duplicate endpoint result identity: {identity}.")
        indexed[identity] = cell
        count = int(_endpoint_field(cell, "success_count"))
        percent = float(_endpoint_field(cell, "success_rate_percent"))
        if not 0 <= count <= EPISODES or not math.isclose(percent, count * 2.0):
            raise CompleteResultsError(
                f"Endpoint result {identity} has inconsistent O50 rate."
            )
    identities = frozenset(indexed)
    if identities != ENDPOINT_EXPECTED_IDENTITIES:
        missing = sorted(ENDPOINT_EXPECTED_IDENTITIES - identities)
        extra = sorted(identities - ENDPOINT_EXPECTED_IDENTITIES)
        raise CompleteResultsError(
            f"Endpoint extension must contain exactly 8 cells; missing={missing}, extra={extra}."
        )
    return indexed


def _c3_epoch3_diagnostic_outcomes(diagnostic: Any | None) -> tuple[bool, ...] | None:
    if diagnostic is None:
        return None
    method = str(_endpoint_field(diagnostic, "method_key"))
    raw_mode = str(_endpoint_field(diagnostic, "score_mode"))
    mode = _ENDPOINT_MODE_ALIASES.get(raw_mode, raw_mode)
    epoch = int(_endpoint_field(diagnostic, "checkpoint_epoch"))
    step = int(_endpoint_field(diagnostic, "checkpoint_global_step"))
    if (method, mode, epoch, step) != ("v1_c3", STATE_V_MODE, 3, 3_000):
        raise CompleteResultsError(
            "C3 early diagnostic must be V1-C3 State-V at E3/global-step 3000."
        )
    outcomes = _endpoint_field(diagnostic, "outcomes")
    if (
        not isinstance(outcomes, (list, tuple))
        or len(outcomes) != EPISODES
        or any(not isinstance(value, bool) for value in outcomes)
    ):
        raise CompleteResultsError(
            "C3 early diagnostic must contain exactly 50 Boolean outcomes."
        )
    outcome_tuple = tuple(outcomes)
    count = int(_endpoint_field(diagnostic, "success_count"))
    if count != sum(outcome_tuple):
        raise CompleteResultsError(
            "C3 early diagnostic success_count differs from its outcomes."
        )
    return outcome_tuple


def _exact_mcnemar_p_two_sided(left_only: int, right_only: int) -> float:
    discordant = left_only + right_only
    if discordant == 0:
        return 1.0
    tail = sum(
        math.comb(discordant, successes)
        for successes in range(min(left_only, right_only) + 1)
    ) / (2**discordant)
    return min(1.0, 2.0 * tail)


def _master_row_metadata() -> tuple[tuple[str, str], ...]:
    return tuple(
        (version, variant)
        for version, variants in MASTER_ROW_GROUPS
        for variant in variants
    )


def _master_group_slices() -> tuple[slice, ...]:
    slices: list[slice] = []
    start = 0
    for _version, variants in MASTER_ROW_GROUPS:
        stop = start + len(variants)
        slices.append(slice(start, stop))
        start = stop
    return tuple(slices)


def _master_endpoint_identity(version: str, variant: str, mode: str) -> tuple[str, str] | None:
    if version != "v1":
        return None
    method = f"v1_{variant}"
    identity = (method, mode)
    return identity if identity in ENDPOINT_EXPECTED_IDENTITIES else None


def _fixed_master_counts(
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None = None,
    *,
    epoch: int = 10,
) -> tuple[tuple[int | None, ...], ...]:
    """Numeric 26-by-7 matrix; unavailable endpoint cells are ``None``."""

    endpoint = _endpoint_index(endpoint_cells)
    rows: list[tuple[int | None, ...]] = []
    for version, variant in _master_row_metadata():
        values: list[int | None] = []
        for mode in MASTER_SCORE_MODES:
            if variant in VARIANTS and mode in ALL_CONTROLLED_MODES:
                values.append(
                    int(
                        _find(
                            cells,
                            version=version,
                            variant=variant,
                            epoch=epoch,
                            mode=mode,
                        )["success_count"]
                    )
                )
                continue
            identity = _master_endpoint_identity(version, variant, mode)
            extension_cell = endpoint.get(identity) if identity is not None else None
            values.append(
                None
                if extension_cell is None
                else int(_endpoint_field(extension_cell, "success_count"))
            )
        rows.append(tuple(values))
    if len(rows) != MASTER_ROW_COUNT:
        raise CompleteResultsError(f"Master matrix must contain {MASTER_ROW_COUNT} rows.")
    return tuple(rows)


def _fixed_master_rows(
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None = None,
    *,
    epoch: int = 10,
) -> tuple[tuple[str, ...], ...]:
    """Return the single 26-method, seven-score comparison matrix."""

    counts = _fixed_master_counts(cells, endpoint_cells, epoch=epoch)
    return tuple(
        (
            VERSION_LABELS[version],
            variant.upper(),
            *("—" if value is None else f"{value}/50 ({value * 2}%)" for value in row),
        )
        for (version, variant), row in zip(_master_row_metadata(), counts)
    )


def _fixed_master_version_column_maxima(
    count_rows: Sequence[Sequence[int | None]],
) -> tuple[tuple[int | None, ...], ...]:
    """Return per-column maxima independently inside each dynamic version band."""

    if len(count_rows) != MASTER_ROW_COUNT:
        raise CompleteResultsError(
            f"Fixed master matrix must contain {MASTER_ROW_COUNT} rows."
        )
    maxima: list[tuple[int | None, ...]] = []
    for group_slice in _master_group_slices():
        block = count_rows[group_slice]
        if any(len(row) != len(MASTER_SCORE_MODES) for row in block):
            raise CompleteResultsError("Fixed master score width changed.")
        group_maxima: list[int | None] = []
        for column in range(len(MASTER_SCORE_MODES)):
            available = tuple(
                int(row[column]) for row in block if row[column] is not None
            )
            group_maxima.append(max(available) if available else None)
        maxima.append(tuple(group_maxima))
    return tuple(maxima)


def _fixed_master_markdown_rows(
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None = None,
    *,
    epoch: int = 10,
) -> tuple[tuple[str, ...], ...]:
    """Format row/within-version winners while leaving missing cells neutral."""

    display_rows = _fixed_master_rows(cells, endpoint_cells, epoch=epoch)
    count_rows = _fixed_master_counts(cells, endpoint_cells, epoch=epoch)
    version_column_maxima = _fixed_master_version_column_maxima(count_rows)
    group_by_row = tuple(
        group_index
        for group_index, (_version, variants) in enumerate(MASTER_ROW_GROUPS)
        for _variant in variants
    )
    formatted: list[tuple[str, ...]] = []
    for row_index, (display_row, count_row) in enumerate(zip(display_rows, count_rows)):
        version_index = group_by_row[row_index]
        available = tuple(value for value in count_row if value is not None)
        row_maximum = max(available) if available else None
        scores = []
        for index, (display, count) in enumerate(zip(display_row[2:], count_row)):
            if count is None:
                scores.append(display)
                continue
            row_best = row_maximum is not None and count == row_maximum
            column_maximum = version_column_maxima[version_index][index]
            column_best = column_maximum is not None and count == column_maximum
            text = f"**{display}**" if row_best else display
            if column_best:
                text = f"◆ {text}"
            scores.append(text)
        method = display_row[1].lower()
        formatted.append((*display_row[:2], METHOD_LOSS_LABELS[method], *scores))
    return tuple(formatted)


def _fixed_version_winner_rows(
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None = None,
    *,
    epoch: int = 10,
) -> tuple[tuple[str, ...], ...]:
    """Summarize blue fills, ignoring unavailable cells in each version band."""

    counts = _fixed_master_counts(cells, endpoint_cells, epoch=epoch)
    metadata = _master_row_metadata()
    maxima = _fixed_master_version_column_maxima(counts)
    rows: list[tuple[str, ...]] = []
    for group_index, ((version, _variants), group_slice) in enumerate(
        zip(MASTER_ROW_GROUPS, _master_group_slices())
    ):
        winners: list[str] = []
        for column, best in enumerate(maxima[group_index]):
            if best is None:
                winners.append("—")
                continue
            labels = "/".join(
                metadata[row_index][1].upper()
                for row_index in range(group_slice.start, group_slice.stop)
                if counts[row_index][column] == best
            )
            winners.append(f"{labels} {best}/50")
        rows.append((VERSION_LABELS[version], *winners))
    return tuple(rows)


def _endpoint_analysis(
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None,
    c3_epoch3_diagnostic: Any | None = None,
) -> EndpointAnalysis:
    endpoint = _endpoint_index(endpoint_cells)
    if not endpoint:
        if c3_epoch3_diagnostic is not None:
            raise CompleteResultsError(
                "C3 E3 diagnostic cannot be reported without the strict endpoint ledger."
            )
        return EndpointAnalysis(
            False,
            0,
            {},
            0,
            0,
            0,
            None,
            None,
            None,
            None,
            None,
            None,
            None,
        )
    deltas: dict[str, int] = {}
    for mode in ALL_CONTROLLED_MODES:
        parent = int(
            _find(cells, version="v1", variant="c", epoch=10, mode=mode)[
                "success_count"
            ]
        )
        c2 = int(_endpoint_field(endpoint[("v1_c2", mode)], "success_count"))
        deltas[mode] = c2 - parent
    parent_first = int(
        _endpoint_field(endpoint[("v1_c", FIRST_Q2_MODE)], "success_count")
    )
    c2_first = int(
        _endpoint_field(endpoint[("v1_c2", FIRST_Q2_MODE)], "success_count")
    )
    deltas[FIRST_Q2_MODE] = c2_first - parent_first
    values = tuple(deltas.values())
    epoch12_cell = endpoint[("v1_c3", STATE_V_MODE)]
    epoch12_count = int(_endpoint_field(epoch12_cell, "success_count"))
    epoch3_outcomes = _c3_epoch3_diagnostic_outcomes(c3_epoch3_diagnostic)
    epoch3_count: int | None = None
    contingency: Mapping[str, int] | None = None
    mcnemar_p: float | None = None
    if epoch3_outcomes is not None:
        epoch12_outcomes = _endpoint_field(epoch12_cell, "outcomes")
        if (
            not isinstance(epoch12_outcomes, (list, tuple))
            or len(epoch12_outcomes) != EPISODES
            or any(not isinstance(value, bool) for value in epoch12_outcomes)
        ):
            raise CompleteResultsError(
                "C3 E12 endpoint must expose 50 Boolean outcomes for paired E3 analysis."
            )
        epoch12_outcomes = tuple(epoch12_outcomes)
        epoch3_count = sum(epoch3_outcomes)
        both_success = sum(a and b for a, b in zip(epoch3_outcomes, epoch12_outcomes))
        epoch3_only = sum(a and not b for a, b in zip(epoch3_outcomes, epoch12_outcomes))
        epoch12_only = sum(not a and b for a, b in zip(epoch3_outcomes, epoch12_outcomes))
        both_failure = sum(
            not a and not b for a, b in zip(epoch3_outcomes, epoch12_outcomes)
        )
        contingency = {
            "both_success": both_success,
            "epoch3_only": epoch3_only,
            "epoch12_only": epoch12_only,
            "both_failure": both_failure,
        }
        mcnemar_p = _exact_mcnemar_p_two_sided(epoch3_only, epoch12_only)
    return EndpointAnalysis(
        available=True,
        cell_count=len(endpoint),
        c2_shared_deltas=deltas,
        c2_improved=sum(value > 0 for value in values),
        c2_tied=sum(value == 0 for value in values),
        c2_harmed=sum(value < 0 for value in values),
        c2_mean_delta_percent=2.0 * sum(values) / len(values),
        first_q2_parent_count=parent_first,
        first_q2_c2_count=c2_first,
        state_v_c3_count=epoch12_count,
        state_v_c3_epoch3_count=epoch3_count,
        state_v_c3_epoch3_vs_epoch12_contingency=contingency,
        state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p=mcnemar_p,
    )


def _trajectory_rows(
    cells: Sequence[ResultCell], version: str, mode: str
) -> tuple[tuple[str, ...], ...]:
    epochs = EMA_EPOCHS if version in ("v2", "v2_ema_sg") else (10,)
    return tuple(
        (
            f"E{epoch}",
            *(
                _result(
                    _find(
                        cells, version=version, variant=variant, epoch=epoch, mode=mode
                    )
                )
                for variant in VARIANTS
            ),
        )
        for epoch in epochs
    )


def _global_posthoc_rows(cells: Sequence[ResultCell]) -> tuple[tuple[str, ...], ...]:
    rows = []
    for version, modes in (("v2", ORIGINAL_MODES), ("v2_ema_sg", ALL_CONTROLLED_MODES)):
        for mode in modes:
            candidates = [
                cell
                for cell in cells
                if cell["version"] == version and cell["score_mode"] == mode
            ]
            best = max(int(cell["success_count"]) for cell in candidates)
            winners = [
                cell for cell in candidates if int(cell["success_count"]) == best
            ]
            labels = ", ".join(
                f"{cell['variant'].upper()}-E{cell['epoch']}" for cell in winners
            )
            formal_e10 = max(
                int(cell["success_count"])
                for cell in candidates
                if int(cell["epoch"]) == 10
            )
            rows.append(
                (
                    "V2" if version == "v2" else "V2-EMA",
                    MODE_LABELS[mode],
                    labels,
                    f"{best}/50 ({best * 2}%)",
                    f"{formal_e10}/50 ({formal_e10 * 2}%)",
                )
            )
    return tuple(rows)


def _ema_new_stability_rows(cells: Sequence[ResultCell]) -> tuple[tuple[str, ...], ...]:
    rows = []
    for variant in VARIANTS:
        for mode in NEW_MODES:
            selected = [
                cell
                for cell in cells
                if cell["version"] == "v2_ema_sg"
                and cell["variant"] == variant
                and cell["score_mode"] == mode
            ]
            values = [_percent(cell) for cell in selected]
            mean = sum(values) / len(values)
            sigma = math.sqrt(
                sum((value - mean) ** 2 for value in values) / len(values)
            )
            best = max(values)
            best_epochs = ",".join(
                f"E{cell['epoch']}" for cell in selected if _percent(cell) == best
            )
            e10 = _percent(
                _find(cells, version="v2_ema_sg", variant=variant, epoch=10, mode=mode)
            )
            rows.append(
                (
                    variant.upper(),
                    MODE_LABELS[mode],
                    f"{mean:.1f}%",
                    f"{sigma:.1f}",
                    f"{best_epochs} / {best:.0f}%",
                    f"{e10:.0f}%",
                )
            )
    return tuple(rows)


def _epoch10_mode_mean(cells: Sequence[ResultCell], version: str, mode: str) -> float:
    return _mean(
        cell
        for cell in cells
        if cell["version"] == version
        and int(cell["epoch"]) == 10
        and cell["score_mode"] == mode
    )


def _fixed_version_mean_rows(
    cells: Sequence[ResultCell],
) -> tuple[tuple[str, ...], ...]:
    """Summarize fixed-E10 version quality without mixing in post-hoc epochs."""

    rows = []
    for version, label in (
        ("v0", "V0 raw action"),
        ("v1", "V1 action encoder"),
        ("v2", "V2 joint fine-tune"),
        ("v2_ema_sg", "V2-EMA-SG"),
    ):
        values = tuple(
            _epoch10_mode_mean(cells, version, mode) for mode in ALL_CONTROLLED_MODES
        )
        rows.append(
            (
                label,
                *(f"{value:.1f}%" for value in values),
                f"{sum(values) / len(values):.1f}%",
            )
        )
    return tuple(rows)


def _fixed_score_mean_rows(
    cells: Sequence[ResultCell],
) -> tuple[tuple[str, ...], ...]:
    """Compare readouts only on versions for which each readout was evaluated."""

    rows = []
    for mode in ALL_CONTROLLED_MODES:
        versions = VERSIONS
        per_version = [_epoch10_mode_mean(cells, version, mode) for version in versions]
        fixed_cells = [
            cell
            for cell in cells
            if cell["version"] in versions
            and int(cell["epoch"]) == 10
            and cell["score_mode"] == mode
        ]
        best = max(int(cell["success_count"]) for cell in fixed_cells)
        winner_labels = tuple(
            f"{str(cell['version']).replace('v2_ema_sg', 'V2-EMA').upper()}-"
            f"{str(cell['variant']).upper()}"
            for cell in fixed_cells
            if int(cell["success_count"]) == best
        )
        rows.append(
            (
                MODE_LABELS[mode],
                "/".join(
                    version.replace("v2_ema_sg", "V2-EMA").upper()
                    for version in versions
                ),
                str(len(fixed_cells)),
                f"{sum(per_version) / len(per_version):.1f}%",
                f"{_compact_join(winner_labels)}: {best}/50 ({best * 2}%)",
            )
        )
    return tuple(rows)


def _ema_variant_mean_rows(
    cells: Sequence[ResultCell],
) -> tuple[tuple[str, ...], ...]:
    rows = []
    for variant in VARIANTS:
        values = [
            _percent(
                _find(
                    cells,
                    version="v2_ema_sg",
                    variant=variant,
                    epoch=10,
                    mode=mode,
                )
            )
            for mode in ALL_CONTROLLED_MODES
        ]
        best_index = max(range(len(values)), key=values.__getitem__)
        rows.append(
            (
                variant.upper(),
                f"{sum(values) / len(values):.1f}%",
                MODE_LABELS[ALL_CONTROLLED_MODES[best_index]],
                f"{values[best_index]:.0f}%",
            )
        )
    return tuple(rows)


def build_training_chart_from_archive(path: str | Path, *, title: str) -> bytes:
    """Render the common total-loss fields from a reconciled 60-row training CSV."""

    source = Path(path)
    with source.open(newline="", encoding="utf-8") as stream:
        rows = tuple(dict(row) for row in csv.DictReader(stream))
    required = {"variant", "epoch", "train/loss_epoch", "validation/loss"}
    if len(rows) != 60 or not rows or not required.issubset(rows[0]):
        raise CompleteResultsError(f"Training curve archive is incomplete: {source}.")

    from PIL import Image, ImageDraw

    image = Image.new("RGB", (2160, 860), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (70, 28),
        title,
        fill="#0B2545",
        font=v2_report._chart_font(34, bold=True),
    )
    train_series = []
    validation_series = []
    all_values: list[float] = []
    for variant, color in zip(VARIANTS, v2_report.CHART_COLORS):
        variant_rows = sorted(
            (row for row in rows if row["variant"] == variant),
            key=lambda row: _int(row["epoch"], context=f"{source}.{variant}.epoch"),
        )
        if [int(row["epoch"]) for row in variant_rows] != list(range(1, 11)):
            raise CompleteResultsError(
                f"{source} has an incomplete {variant} trajectory."
            )
        train_values = [
            _float(row["train/loss_epoch"], context=f"{source}.{variant}.train")
            for row in variant_rows
        ]
        validation_values = [
            _float(row["validation/loss"], context=f"{source}.{variant}.validation")
            for row in variant_rows
        ]
        if any(value <= 0 for value in train_values + validation_values):
            raise CompleteResultsError(f"{source} contains non-positive loss values.")
        train_series.append((variant.upper(), train_values, color))
        validation_series.append((variant.upper(), validation_values, color))
        all_values.extend(train_values + validation_values)
    log_min = math.floor(min(math.log10(value) for value in all_values) * 2) / 2
    log_max = math.ceil(max(math.log10(value) for value in all_values) * 2) / 2
    y_min = 10**log_min
    y_max = 10**log_max
    ticks = tuple(
        (10.0**power, f"1e{power}")
        for power in range(math.ceil(log_min), math.floor(log_max) + 1)
    )
    if not ticks:
        ticks = ((y_min, f"{y_min:.2g}"), (y_max, f"{y_max:.2g}"))
    for box, panel_title, series in (
        ((55, 90, 1060, 825), "Training total loss", train_series),
        ((1100, 90, 2105, 825), "Validation total loss", validation_series),
    ):
        v2_report._draw_line_panel(
            draw,
            box=box,
            title=panel_title,
            x_values=tuple(range(1, 11)),
            series=series,
            y_min=y_min,
            y_max=y_max,
            y_ticks=ticks,
            y_transform=math.log10,
            y_label="Loss (log scale)",
        )
    return v2_report._save_chart(image)


def _linear_chart_scale(
    values: Sequence[float],
    *,
    include_zero: bool = False,
    percent: bool = False,
) -> tuple[float, float, tuple[tuple[float, str], ...]]:
    if not values or any(not math.isfinite(value) for value in values):
        raise CompleteResultsError("Chart scale requires finite values.")
    low = min(values)
    high = max(values)
    span = high - low
    padding = span * 0.12 if span > 0 else max(abs(high) * 0.12, 0.1)
    y_min = min(0.0, low - padding) if include_zero else low - padding
    y_max = high + padding
    if not y_max > y_min:
        y_max = y_min + 1.0
    ticks = []
    for index in range(5):
        value = y_min + index * (y_max - y_min) / 4
        label = f"{value:.1f}%" if percent else f"{value:.3g}"
        ticks.append((value, label))
    return y_min, y_max, tuple(ticks)


def _endpoint_training_summary_rows(
    evidence: EndpointTrainingEvidence,
) -> tuple[tuple[str, ...], ...]:
    def relative_change(values: Sequence[float]) -> str:
        start, end = float(values[0]), float(values[-1])
        return f"{100.0 * (end / start - 1.0):+.1f}%"

    def scalar_row(
        method: str,
        metric: str,
        values: Sequence[float],
        *,
        digits: int,
        change: str | None = None,
    ) -> tuple[str, ...]:
        return (
            method,
            metric,
            f"{values[0]:.{digits}f}",
            f"{values[-1]:.{digits}f}",
            relative_change(values) if change is None else change,
        )

    c2_top1_delta = 100.0 * (
        evidence.c2_alignment_top1[-1] - evidence.c2_alignment_top1[0]
    )
    c3_spearman_delta = evidence.c3_spearman[-1] - evidence.c3_spearman[0]
    return (
        scalar_row("V1-C2", "Train total loss", evidence.c2_train_loss, digits=2),
        scalar_row(
            "V1-C2", "Validation base TD loss", evidence.c2_validation_loss, digits=2
        ),
        scalar_row(
            "V1-C2", "First-Q ranking CE", evidence.c2_alignment_loss, digits=4
        ),
        scalar_row(
            "V1-C2",
            "First-Q top-1 agreement",
            tuple(100.0 * value for value in evidence.c2_alignment_top1),
            digits=2,
            change=f"{c2_top1_delta:+.2f} pp; random 6.25%",
        ),
        scalar_row("V1-C3", "Train TD loss", evidence.c3_train_loss, digits=4),
        scalar_row(
            "V1-C3", "Validation TD loss", evidence.c3_validation_loss, digits=4
        ),
        scalar_row("V1-C3", "Validation MC MAE", evidence.c3_mc_mae, digits=3),
        scalar_row(
            "V1-C3",
            "Validation TD residual MAE",
            evidence.c3_td_residual_mae,
            digits=3,
        ),
        scalar_row(
            "V1-C3",
            "Validation Spearman",
            evidence.c3_spearman,
            digits=4,
            change=f"{c3_spearman_delta:+.4f}",
        ),
    )


def build_endpoint_training_chart(evidence: EndpointTrainingEvidence) -> bytes:
    """Render one compact 2x2 C2/C3 convergence and calibration figure."""

    from PIL import Image, ImageDraw

    image = Image.new("RGB", (2160, 1200), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (70, 24),
        "V1-C2 and V1-C3 training and validation diagnostics",
        fill="#0B2545",
        font=v2_report._chart_font(34, bold=True),
    )

    c2_loss_values = (
        *evidence.c2_train_loss,
        *evidence.c2_validation_loss,
        *evidence.c2_alignment_loss,
    )
    log_min = math.floor(min(math.log10(value) for value in c2_loss_values))
    log_max = math.ceil(max(math.log10(value) for value in c2_loss_values))
    c2_log_ticks = tuple(
        (10.0**power, f"1e{power}") for power in range(log_min, log_max + 1)
    )
    v2_report._draw_line_panel(
        draw,
        box=(45, 78, 1058, 618),
        title="C2 loss scales",
        x_values=evidence.c2_epochs,
        series=(
            ("Train", evidence.c2_train_loss, v2_report.CHART_COLORS[0]),
            ("Val TD", evidence.c2_validation_loss, v2_report.CHART_COLORS[1]),
            ("Rank CE", evidence.c2_alignment_loss, v2_report.CHART_COLORS[2]),
        ),
        y_min=10.0**log_min,
        y_max=10.0**log_max,
        y_ticks=c2_log_ticks,
        y_transform=math.log10,
        y_label="Loss (log scale)",
    )

    c2_top1 = tuple(100.0 * value for value in evidence.c2_alignment_top1)
    c2_top1_values = (*c2_top1, 6.25)
    y_min, y_max, ticks = _linear_chart_scale(c2_top1_values)
    v2_report._draw_line_panel(
        draw,
        box=(1102, 78, 2115, 618),
        title="C2 planner-ranking diagnostic",
        x_values=evidence.c2_epochs,
        series=(
            ("Top-1", c2_top1, v2_report.CHART_COLORS[0]),
            ("Random", (6.25,) * len(c2_top1), v2_report.CHART_COLORS[5]),
        ),
        y_min=y_min,
        y_max=y_max,
        y_ticks=ticks,
        y_transform=lambda value: value,
        y_label="Agreement (%)",
    )

    c3_loss_values = (*evidence.c3_train_loss, *evidence.c3_validation_loss)
    y_min, y_max, ticks = _linear_chart_scale(c3_loss_values)
    v2_report._draw_line_panel(
        draw,
        box=(45, 650, 1058, 1175),
        title="C3 expectile-Huber TD loss",
        x_values=evidence.c3_epochs,
        series=(
            ("Train", evidence.c3_train_loss, v2_report.CHART_COLORS[0]),
            ("Validation", evidence.c3_validation_loss, v2_report.CHART_COLORS[1]),
        ),
        y_min=y_min,
        y_max=y_max,
        y_ticks=ticks,
        y_transform=lambda value: value,
        y_label="Loss",
    )

    c3_error_values = (*evidence.c3_mc_mae, *evidence.c3_td_residual_mae)
    y_min, y_max, ticks = _linear_chart_scale(c3_error_values)
    v2_report._draw_line_panel(
        draw,
        box=(1102, 650, 2115, 1175),
        title="C3 validation error",
        x_values=evidence.c3_epochs,
        series=(
            ("MC MAE", evidence.c3_mc_mae, v2_report.CHART_COLORS[0]),
            ("TD MAE", evidence.c3_td_residual_mae, v2_report.CHART_COLORS[1]),
        ),
        y_min=y_min,
        y_max=y_max,
        y_ticks=ticks,
        y_transform=lambda value: value,
        y_label=(
            f"MAE; Spearman {evidence.c3_spearman[0]:.3f} -> "
            f"{evidence.c3_spearman[-1]:.3f}"
        ),
    )
    return v2_report._save_chart(image)


def build_markdown(
    cells: Sequence[ResultCell],
    ledger_evidence: LedgerEvidence,
    endpoint_cells: Sequence[Any] | None = None,
    c3_epoch3_diagnostic: Any | None = None,
    *,
    endpoint_training_evidence: EndpointTrainingEvidence | None = None,
    endpoint_training_chart_path: str | None = None,
) -> str:
    analysis = _fixed_analysis(cells)
    endpoint_analysis = _endpoint_analysis(
        cells, endpoint_cells, c3_epoch3_diagnostic
    )
    f_plus_g_winner = _winner_result(analysis, "f_plus_g")
    overall_winner = (
        f"{_compact_join(analysis.overall_winners)}: {analysis.overall_best_count}/50 "
        f"({analysis.overall_best_count * 2}%)"
    )
    best_score_names = _joined(
        tuple(MODE_LABELS[mode] for mode in analysis.best_mean_modes)
    )
    best_training_mean_names = _joined(analysis.best_training_mean_winners)
    ema_best_variants = _compact_join(
        tuple(variant.upper() for variant in analysis.ema_best_variants)
    )
    ema_best_mean = analysis.ema_variant_means[analysis.ema_best_variants[0]]
    v1_f = analysis.version_mode_means[("v1", "f_only")]
    v2_f = analysis.version_mode_means[("v2", "f_only")]
    v1_tail = analysis.version_mode_means[("v1", "f_plus_g")]
    v2_tail = analysis.version_mode_means[("v2", "f_plus_g")]
    ema_mode_summary = "、".join(
        f"{MODE_LABELS[mode]} {analysis.version_mode_means[('v2_ema_sg', mode)]:.1f}%"
        for mode in ALL_CONTROLLED_MODES
    )
    mean_q_winner = _winner_result(analysis, "g_only_f_rollout_mean")
    tail_deltas = tuple(
        _percent(
            _find(cells, version=version, variant=variant, epoch=10, mode="f_plus_g")
        )
        - _percent(
            _find(cells, version=version, variant=variant, epoch=10, mode="f_only")
        )
        for version in VERSIONS
        for variant in VARIANTS
    )
    tail_gain_count = sum(delta > 0 for delta in tail_deltas)
    tail_tie_count = sum(delta == 0 for delta in tail_deltas)
    tail_harm_count = sum(delta < 0 for delta in tail_deltas)
    tail_mean_delta = sum(tail_deltas) / len(tail_deltas)
    version_overall_means = {
        version: sum(
            analysis.version_mode_means[(version, mode)]
            for mode in ALL_CONTROLLED_MODES
        )
        / len(ALL_CONTROLLED_MODES)
        for version in VERSIONS
    }
    lines = [
        "# Results TD — 全部 Actor-Free TD-LeWM 实验总账（Cube seed 3072）",
        "",
        f"本报告保留 **{COMPLETE_CELL_COUNT} 个已核验正式 O50 基础单元**及其原分析；唯一主矩阵另预留 V1-C2/C3 的 8 个严格 endpoint 单元。基础方法固定 E10，C2 固定最终 E10，C3 固定最终 E12。每格均为同一组 50 个 start-goal pair；训练 seed=3072，planning seed=42。模型均不训练 Actor。",
        "",
        "## 一句话结论",
        "",
        f"- **按原先固定的主评分列 F+G，描述性领先配置为 {f_plus_g_winner}。**",
        f"- **所有固定 E10 单格的最高结果为 {overall_winner}。**",
        f"- **按四版本、24 个训练配置的固定 E10 均值，描述性领先测试评分为 {best_score_names}（{analysis.best_mean_percent:.1f}%）。** 单 seed 下不把它表述为统计稳健最优。",
        f"- **若把五种评分等权平均，描述性领先训练配置为 {best_training_mean_names}（并列 {analysis.best_training_mean_percent:.1f}%）。**",
        f"- **按六个训练方法 × 五种评分的版本均值，V1 action encoder 最高（{version_overall_means['v1']:.1f}%）。**",
        f"- V1→V2 联合微调后，F-only 均值由 {v1_f:.1f}% 变为 {v2_f:.1f}%（{_signed_pp(v2_f - v1_f)}），F+G 由 {v1_tail:.1f}% 变为 {v2_tail:.1f}%（{_signed_pp(v2_tail - v1_tail)}）；这首先提示 world-model/control representation 变化，而不只是 G 的读出形式。",
        "",
        "## 完整全账伴随文件",
        "",
    ]
    lines += _markdown_table(
        ("文件", "保留内容", "路径", "SHA-256"),
        (
            (
                "CSV scalar ledger",
                f"{ledger_evidence.cell_count} 个 O50 单元",
                f"`{ledger_evidence.csv_path}`",
                f"`{ledger_evidence.csv_sha256}`",
            ),
            (
                "JSON reconciliation ledger",
                f"{ledger_evidence.cell_count} 格 × {EPISODES} outcomes = {ledger_evidence.outcome_count:,}",
                f"`{ledger_evidence.json_path}`",
                f"`{ledger_evidence.json_sha256}`",
            ),
        ),
    )
    lines += [
        "",
        f"原 24×5 分析固定使用 E10；全部 {ledger_evidence.cell_count} 格和 {ledger_evidence.outcome_count:,} 个逐-pair 布尔结果仍由上述伴随文件完整保留。C2/C3 的 8 格由独立严格 endpoint ledger 接入，不改写原账。",
        "",
        "## 结果覆盖与版本定义",
        "",
    ]
    coverage = (
        ("Legacy", "7", "E10", "F / G(C) / combined", "21"),
        ("V0 raw action", "6", "E10", "all five scores", "30"),
        ("V1 action encoder", "6", "E10", "all five scores", "30"),
        ("V2 joint fine-tune", "6", "E3-E10", "3 original + E10 first/Mean", "156"),
        ("V2-EMA-SG", "6", "E3-E10", "all five scores", "240"),
        (
            "V1-C2/C3 endpoint extension",
            "2",
            "C2 E10 / C3 E12",
            "First-Q2 + State-V integrated into seven-column matrix",
            "8" if endpoint_analysis.available else "pending strict ledger",
        ),
    )
    lines += _markdown_table(
        ("版本/家族", "方法数", "Checkpoint", "评分覆盖", "O50 格数"), coverage
    )
    lines += [
        "",
        "## 方法、网络和训练 loss",
        "",
        "旧结构消融比较 Successor/critic head 与 LeWM predictor 的连接方式：Serial Decoupled、Serial Coupled、Hybrid、Parallel Real、Goal Hybrid、Imaginary Hybrid、Direct Goal Critic Hybrid。其总目标均为 `L_LeWM + α_u L_TD`，区别在 real/predicted 支路、是否让 TD 梯度进入 LeWM、是否使用 goal projection/imaginary bootstrap/direct scalar critic。DOCX 继续保留前一版的详细结构、loss、训练曲线与 V0/V1 逐方法说明。",
        "",
        "C–G3 家族共享同一个 TD-JEPA predictor `G`。V0 输入归一化 raw action；V1 改用冻结的 LeWM Action Encoder；V2 联合微调 LeWM/Action Encoder/G；V2-EMA-SG 进一步用 EMA world model、EMA action encoder 与 EMA G 构造完全 stop-gradient target。",
        "",
        "基础 target 与总目标：",
        "",
        r"$$Y_t=\operatorname{sg}[\bar z_{t+1}+\gamma(1-d_t)G_{\bar\phi}(\bar z_{t+1},\bar e_{t+1},m_t)],\quad \gamma=0.95,$$",
        "",
        r"$$L_{total}=L_{pred}+0.09L_{SIGReg}+\rho(u)(L_{method}^{real}+L_{method}^{pred}).$$",
        "",
        "其中分支 `b∈{real,pred}`；逐样本基础 TD 残差为 `l_i^b=||G_φ(s_i^b,e_i,m_i)-Y_i||²`，`m_i` 是 goal/random task vector，`ρ(u)` 是 TD warm-up 权重。C 额外训练 goal scalar residual；D/F/G1/G2/G3 只用 stop-gradient 的优势信号重加权 `l_i^b`。",
        "",
    ]
    lines += _markdown_table(
        ("方法", "特殊训练信号", "支路 loss", "作用"),
        tuple(
            tuple(str(value) for value in row) for row in v2_report.METHOD_FORMULA_ROWS
        )
        + (
            (
                "C2 (V1 only)",
                "Frozen-F terminal goal-cost ranking over 16 candidate action sequences",
                "L_C2=L_C+CE(p_F,p_Q); p_F=softmax(-z_cand(J_F)), p_Q=softmax(z_cand(Q_G(z0,A1,g)))",
                "Initialize every parameter from V1-C E10, freeze LeWM/Action Encoder, and fine-tune only G so First-Q follows the planner ranking",
            ),
            (
                "C3 (V1 only)",
                "Same-episode temporal distance in primitive-step units with an EMA State-V bootstrap",
                "L_C3=E[omega_tau(r)Huber_1(r)], r=V_psi(z,g)-sg(y), tau=0.03; y=delta inside n_eff, otherwise c_gamma(n_eff)+gamma^n_eff V_bar(z_succ,g)",
                "Freeze the complete V1-C parent, including both G copies; train only a nonnegative MRN State-V critic (gamma=0.98, n<=50 primitives)",
            ),
        ),
    )
    lines += [
        "",
        "## 七种测试方法怎么测",
        "",
        "统一约定：`z0` 是当前真实图像经部署 encoder 得到的 latent，`z_g` 是 goal 图像的 latent，`z_k^F` 是 LeWM rollout 的 imagined latent。V1/V2/V2-EMA 先用共享 Action Encoder 得到 `e_k=E_A(A_k)`；V0 直接把归一化 25D action block 输入 G。`Q_G(z,A,g)=G(z,e,w(g))^T w(g)`。CEM 始终最小化 cost。",
        "",
    ]
    lines += _markdown_table(
        ("统一评测字段", "固定设置"), COMMON_EVALUATION_PROTOCOL_ROWS
    )
    lines += ["", "七个评分列的实际计算：", ""]
    lines += _markdown_table(
        ("评分列", "F/G 的实际路径", "CEM 最小化的 cost", "goal/Q 使用位置"),
        EVALUATION_METHOD_ROWS,
    )
    lines += [
        "",
        "V2-EMA 的 EMA world model、EMA Action Encoder 和 EMA G 只构造训练 target；正式 CEM 测试仍部署 online F、online Action Encoder 和 online G。Legacy 7 方法的旧 `G/C-only` 会先由 F 构造 H-1 tail context，不等同于 C-G3 主矩阵里严格 H=1 的 `G-only`。",
        "",
        "## Legacy 7 方法：完整 21 格",
        "",
    ]
    lines += _markdown_table(
        ("方法", "F-only", "G/C-only", "Combined"), _legacy_rows(cells)
    )
    lines += [
        "",
        "## 26 个方法 × 7 种评分的唯一主结果矩阵",
        "",
        "横向读每一行，可以同时看到训练 loss，并比较同一个训练方法已有的评分；纵向读每一列时，以版本为边界比较该版本内所有可用方法。Markdown 中 **粗体**是行最佳，`◆` 是同版本列最佳；并列全部标记。缺失格显示 `—`，不参加任何最大值；DOCX 使用黄底表示行最佳、蓝底表示同版本列最佳、青色底表示两者同时成立。",
        "",
        "Loss 列采用紧凑记号：`l_i` 是逐样本 successor TD 残差，`qY=Y^T m`；D–G3 的 `w_i(·)` 是由括号内 stop-gradient 信号形成的归一化样本权重。V0/V1 只有 real 分支；V2/V2-EMA 的总目标为 `L_pred+0.09L_SIGReg+ρ(L_method^real+L_method^pred)`。精确信号、goal 子集和权重定义见前面的“方法、网络和训练 loss”表。",
        "",
    ]
    lines += _markdown_table(
        (
            "版本",
            "训练方法",
            "训练 loss",
            "F-only",
            "G-only",
            "F+G tail",
            "First-Q",
            "Mean-Q",
            "First-Q2",
            "State-V",
        ),
        _fixed_master_markdown_rows(cells, endpoint_cells),
    )
    lines += [
        "",
        "### 每个版本内部的逐列赢家",
        "",
    ]
    lines += _markdown_table(
        (
            "版本",
            "F-only",
            "G-only",
            "F+G tail",
            "First-Q",
            "Mean-Q",
            "First-Q2",
            "State-V",
        ),
        _fixed_version_winner_rows(cells, endpoint_cells),
    )
    lines += [
        "",
        f"**跨四版本的全局逐列赢家（只用于补充分析，不对应 DOCX 蓝框）：** {_column_winner_summary(analysis)}。",
        "",
        "### V1-C2/C3 endpoint 证据",
        "",
    ]
    if endpoint_analysis.available:
        endpoint = _endpoint_index(endpoint_cells)
        comparison_rows = []
        for mode in (*ALL_CONTROLLED_MODES, FIRST_Q2_MODE):
            if mode == FIRST_Q2_MODE:
                parent_count = endpoint_analysis.first_q2_parent_count
            else:
                parent_count = int(
                    _find(
                        cells,
                        version="v1",
                        variant="c",
                        epoch=10,
                        mode=mode,
                    )["success_count"]
                )
            c2_count = int(
                _endpoint_field(endpoint[("v1_c2", mode)], "success_count")
            )
            assert parent_count is not None
            comparison_rows.append(
                (
                    MODE_LABELS[mode],
                    f"{parent_count}/50 ({parent_count * 2}%)",
                    f"{c2_count}/50 ({c2_count * 2}%)",
                    f"{c2_count - parent_count:+d}/50 ({(c2_count - parent_count) * 2:+d} pp)",
                )
            )
        lines += _markdown_table(
            ("共享评分", "V1-C parent", "V1-C2", "C2-parent"),
            tuple(comparison_rows),
        )
        lines += [
            "",
            f"C2 在 6 个可比评分上为 {endpoint_analysis.c2_improved} 升 / {endpoint_analysis.c2_tied} 平 / {endpoint_analysis.c2_harmed} 降，平均变化 {endpoint_analysis.c2_mean_delta_percent:+.1f} pp。C3 的独立 State-V endpoint 为 {endpoint_analysis.state_v_c3_count}/50 ({endpoint_analysis.state_v_c3_count * 2}%)。这些结论只在严格 8-cell ledger 到齐后由实际 outcome 生成。",
            "",
        ]
        if endpoint_analysis.state_v_c3_epoch3_count is not None:
            contingency = (
                endpoint_analysis.state_v_c3_epoch3_vs_epoch12_contingency
            )
            assert contingency is not None
            lines += [
                f"C3 的同一组 50 个 pair 早期诊断为 E3 {endpoint_analysis.state_v_c3_epoch3_count}/50 ({endpoint_analysis.state_v_c3_epoch3_count * 2}%)，最终 endpoint 为 E12 {endpoint_analysis.state_v_c3_count}/50 ({endpoint_analysis.state_v_c3_count * 2}%)。配对列联为：两者均成功 {contingency['both_success']}、仅 E3 成功 {contingency['epoch3_only']}、仅 E12 成功 {contingency['epoch12_only']}、两者均失败 {contingency['both_failure']}；exact McNemar 双侧 p={endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p:.6g}。E3 只作为诊断，不进入 8-cell 主表或 endpoint 计数。",
                "",
            ]
        parent_f_only = int(
            _find(
                cells,
                version="v1",
                variant="c",
                epoch=10,
                mode="f_only",
            )["success_count"]
        )
        parent_first_q = int(
            _find(
                cells,
                version="v1",
                variant="c",
                epoch=10,
                mode="f_plus_g_first",
            )["success_count"]
        )
        assert endpoint_analysis.state_v_c3_count is not None
        assert endpoint_analysis.first_q2_parent_count is not None
        assert endpoint_analysis.first_q2_c2_count is not None
        first_q_delta_pp = (
            2 * endpoint_analysis.c2_shared_deltas["f_plus_g_first"]
        )
        first_q2_delta_pp = 2 * (
            endpoint_analysis.first_q2_c2_count
            - endpoint_analysis.first_q2_parent_count
        )
        c3_vs_f_pp = 2 * (
            endpoint_analysis.state_v_c3_count - parent_f_only
        )
        c3_vs_first_q_pp = 2 * (
            endpoint_analysis.state_v_c3_count - parent_first_q
        )
        lines += [
            "#### Endpoint 专项解释与下一步",
            "",
            f"- **C2 的结果是混合的，而不是稳定的 ranking 增益。** 六种共享评分合计 {endpoint_analysis.c2_improved} 升 / {endpoint_analysis.c2_tied} 平 / {endpoint_analysis.c2_harmed} 降，平均变化 {endpoint_analysis.c2_mean_delta_percent:+.1f} pp；其中 First-Q 为 {first_q_delta_pp:+d} pp，First-Q2 为 {first_q2_delta_pp:+d} pp。即便个别格上升，单 seed 下也没有跨读出一致、可称为稳健的排序改善。",
            f"- **C3 的最终 State-V 是 {endpoint_analysis.state_v_c3_count}/50。** 相对同一个 V1-C parent 的 F-only 为 {c3_vs_f_pp:+d} pp，相对 V1-C 的 First-Q 为 {c3_vs_first_q_pp:+d} pp；它说明独立时间价值读出有信号，但尚未稳定超过 parent 的最佳 first-action readout。",
        ]
        if (
            endpoint_analysis.state_v_c3_epoch3_count is not None
            and endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p
            is not None
        ):
            p_value = (
                endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p
            )
            significance = (
                "没有达到常用 0.05 阈值"
                if p_value >= 0.05
                else "达到常用 0.05 阈值"
            )
            lines += [
                f"- E3 与 E12 的同-pair exact McNemar p={p_value:.6g}，{significance}。因此不能在看过正式 O50 后把 E3 当成新的正式 endpoint；下一轮应在独立 dev pairs 上选择 epoch，或事先登记 early-stop 规则。",
            ]
        lines += [
            "- C3 训练主要读取真实 encoder latent，推理却在 `F^5` imagined terminal latent 上读 State-V。下一轮应把 stop-gradient 的 F-imagined states 按受控比例混入 State-V 训练，直接缩小这一 terminal-state OOD 间隙。",
            "- 以上比较均为一个 training seed 和同一组 50 pair 的描述性消融，不构成多 seed 总体最优或因果证明。",
            "",
        ]
    else:
        lines += [
            "V1-C First-Q2、V1-C2 的六种评分和 V1-C3 State-V 共 8 格尚未提供严格 endpoint ledger；主矩阵保留中性 `—`，本节不生成数值比较或改写原 477-cell 结论。",
            "",
        ]
    lines += [
        "## 训练 / validation loss 证据",
        "",
        "训练总 loss 含不同辅助项，绝对数值不能直接给 C–G3 排名，只用于判断各自是否收敛。Legacy 与 V1 曲线保留在历史来源文档；V0、V2、V2-EMA 的逐 epoch 数值和全部 E3–E10 O50 轨迹继续保留在总账 artifacts 中，但不再塞进主结果表。",
        "",
    ]
    if endpoint_training_evidence is not None:
        lines += [
            "V1-C2/C3 的下表和图只读取 endpoint archive 中 hash-bound 的 `training/v1_c2/metrics.csv` 与 `training/v1_c3/metrics.csv`；每个必需指标每个 epoch 必须恰有一个有限 aggregate，否则报告构建失败。",
            "",
        ]
        lines += _markdown_table(
            ("方法", "指标", "首个 epoch", "最终 epoch", "变化"),
            _endpoint_training_summary_rows(endpoint_training_evidence),
        )
        if endpoint_training_chart_path is not None:
            lines += [
                "",
                f"![V1-C2 and V1-C3 training and validation diagnostics]({endpoint_training_chart_path})",
            ]
        c2_scale_ratio = 100.0 * (
            endpoint_training_evidence.c2_alignment_loss[-1]
            / endpoint_training_evidence.c2_train_loss[-1]
        )
        lines += [
            "",
            f"C2 的 validation base TD loss 从 {endpoint_training_evidence.c2_validation_loss[0]:.2f} 降至 {endpoint_training_evidence.c2_validation_loss[-1]:.2f}，但 First-Q ranking CE 仅从 {endpoint_training_evidence.c2_alignment_loss[0]:.4f} 到 {endpoint_training_evidence.c2_alignment_loss[-1]:.4f}，top-1 agreement 从 {100 * endpoint_training_evidence.c2_alignment_top1[0]:.2f}% 到 {100 * endpoint_training_evidence.c2_alignment_top1[-1]:.2f}%。E10 的 ranking CE 数值只相当于 train total loss 的 {c2_scale_ratio:.4f}%；这不能单独证明梯度不足，但与其几乎不动和 endpoint 无净增益一致。下一版应先做 loss 标准化，再使用自适应权重或梯度平衡，而不是继续固定权重 1。",
            f"C3 的 validation TD loss 从 {endpoint_training_evidence.c3_validation_loss[0]:.4f} 降至 {endpoint_training_evidence.c3_validation_loss[-1]:.4f}，MC MAE 从 {endpoint_training_evidence.c3_mc_mae[0]:.3f} 到 {endpoint_training_evidence.c3_mc_mae[-1]:.3f}，TD residual MAE 从 {endpoint_training_evidence.c3_td_residual_mae[0]:.3f} 到 {endpoint_training_evidence.c3_td_residual_mae[-1]:.3f}，Spearman 从 {endpoint_training_evidence.c3_spearman[0]:.4f} 到 {endpoint_training_evidence.c3_spearman[-1]:.4f}。优化指标在缓慢改善，但正式控制 success 并未随训练后段单调提高，说明 value calibration 与 planner utility 仍需分开验证。",
            "",
            f"证据指纹：C2 `{endpoint_training_evidence.c2_metrics_path}` (`{endpoint_training_evidence.c2_metrics_sha256}`)；C3 `{endpoint_training_evidence.c3_metrics_path}` (`{endpoint_training_evidence.c3_metrics_sha256}`)。",
            "",
        ]
    elif endpoint_analysis.available:
        lines += [
            "严格 endpoint 分数已接入，但 archive 没有同时声明 C2 与 C3 的训练 metrics；本报告不绘制或推断缺失的训练曲线。",
            "",
        ]
    else:
        lines += [
            "尚未提供 endpoint ledger，因此 C2/C3 训练曲线保持缺失，不生成占位图或数值结论。",
            "",
        ]
    lines += [
        "## 最佳训练方法与最佳测试评分",
        "",
        "### 固定 E10 四版本均值",
        "",
    ]
    lines += _markdown_table(
        ("训练版本", "F-only", "G-only", "F+G", "First-Q", "Mean-Q", "五评分均值"),
        _fixed_version_mean_rows(cells),
    )
    lines += [
        "",
        "### 五种评分的四版本汇总",
        "",
    ]
    lines += _markdown_table(
        ("评分方式", "覆盖版本", "固定格数", "四版本均值", "最高固定单格"),
        _fixed_score_mean_rows(cells),
    )
    lines += [
        "",
        "### 1. 哪个训练方法最好",
        "",
        f"不存在脱离测试评分定义的唯一训练赢家。按原研究固定的 F+G 主列，领先配置为 **{f_plus_g_winner}**；若把五种评分等权平均，则 **{best_training_mean_names} 并列领先（{analysis.best_training_mean_percent:.1f}%）**；若寻找最高单格，则为 **{overall_winner}**。在 V2-EMA E10 内，五评分均值最高的训练变体为 **{ema_best_variants}（{ema_best_mean:.1f}%）**。这些都是描述性单 seed 结果。",
        f"从版本整体看，V1 action encoder 的六方法 × 五评分均值最高（{version_overall_means['v1']:.1f}%）。",
        "",
        "### 2. 哪个测试方法最好",
        "",
        f"V2-EMA E10 六个训练方法的均值为：{ema_mode_summary}。跨 V0/V1/V2/V2-EMA 的固定 E10，**{best_score_names}** 的 24 配置均值最高（{analysis.best_mean_percent:.1f}%），因此它是当前描述性默认主测试方式。Mean-Q 的最高固定配置为 {mean_q_winner}。",
        "",
        "### 3. 原因分析",
        "",
        f"- V0→V1 后 G-only 均值从 {analysis.version_mode_means[('v0', 'g_only')]:.1f}% 到 {analysis.version_mode_means[('v1', 'g_only')]:.1f}%，First-Q 从 {analysis.version_mode_means[('v0', 'f_plus_g_first')]:.1f}% 到 {analysis.version_mode_means[('v1', 'f_plus_g_first')]:.1f}%；这与共享 Action Encoder 改善 G 读出一致。",
        "- V1→V2 后 F-only 与 F+G 同时大幅下降，与 TD 梯度进入 online LeWM/Action Encoder 后产生 latent/control representation drift 的假设一致；单 seed 不能证明因果，问题也不能只归因于 critic。",
        f"- 当前均值领先读出是 {best_score_names}；不同读出暴露于真实状态、imagined states 与 G 尺度的程度不同，OOD/rollout 误差仍是需要用独立 dev pairs 验证的解释。",
        f"- Mean-Q 在 24 个固定训练配置上的均值为 {analysis.mode_means['g_only_f_rollout_mean']:.1f}%，最高格为 {mean_q_winner}；它是否应进入主评测由这一完整覆盖结果决定，不再按旧的 V2-only 结论处理。",
        "",
        "## 负面结果与下一轮目标",
        "",
    ]
    lines += _markdown_table(
        ("发现", "证据", "含义/下一步"),
        (
            (
                "V2 world model 退化",
                f"V1→V2 F-only {v1_f:.1f}%→{v2_f:.1f}%",
                "先恢复 F，再谈 G 增益",
            ),
            (
                "EMA 未根治",
                "V2→EMA 五评分变化："
                + "，".join(
                    f"{MODE_LABELS[mode]} {_signed_pp(analysis.version_mode_means[('v2_ema_sg', mode)] - analysis.version_mode_means[('v2', mode)])}"
                    for mode in ALL_CONTROLLED_MODES
                ),
                "冻结或低 LR 微调 encoder/world",
            ),
            (
                "tail 效果异质",
                f"F+G 对 F-only：{tail_gain_count} 升 / {tail_tie_count} 平 / {tail_harm_count} 降；均值 {_signed_pp(tail_mean_delta)}",
                f"不能默认 tail 必然增益；以 {best_score_names} 为候选并校准 G",
            ),
            (
                "Mean-Q 完整覆盖",
                f"24 配置均值 {analysis.mode_means['g_only_f_rollout_mean']:.1f}%；{mean_q_winner}",
                "按完整四版本结果决定主评测或消融地位",
            ),
            (
                "checkpoint 选择偏差",
                "同一 O50 上看 E3–E10 再取最大",
                "dev pairs 选 epoch/alpha，正式 O50 只跑锁定配置",
            ),
            (
                "单 seed",
                "全部训练 seed=3072",
                "至少 3 个训练 seeds，保存逐-pair outcome",
            ),
        ),
    )
    lines += [
        "",
        "下一轮优先目标：",
        "",
        f"1. 把联合模型固定 E10 的 F-only 六法均值从 {analysis.version_mode_means[('v2_ema_sg', 'f_only')]:.1f}% 恢复到至少 V1 的 {v1_f:.1f}%。",
        f"2. 预注册 `{_compact_join(analysis.overall_winners)}` 与 `{_compact_join(analysis.mode_winners['f_plus_g'])} + F+G tail` 作为主基线；正式 O50 不再事后选 epoch。",
        "3. 若继续 joint training，先冻结 encoder/world 或给 TD 极低学习率，再分阶段解冻；增加对 V1 latent/prediction 的 anchor，并限制 TD 梯度进入 F。",
        "4. 在独立 dev pair set 上选择 α、epoch 与 Q 校准；对 CEM 候选/imagined-state 分布加入 conservative/calibration 训练，抑制 OOD 高估。",
        "5. 至少 3 个、最好 5 个 training seeds；用 paired bootstrap/McNemar 分析固定配置。",
        "",
        "## 审计边界",
        "",
        f"- {COMPLETE_CELL_COUNT}/{COMPLETE_CELL_COUNT} 格共享 episode-selection 文件 SHA-256 `{SELECTION_SHA256}` 与 action normalization SHA-256 `{ACTION_NORMALIZATION_SHA256}`。",
        f"- fixed 新评分 launcher 另有 valid-row-ranks SHA-256 `{FIXED_SELECTION_RANKS_SHA256}`；它是规范化索引哈希，不是 episode-selection 文件哈希，二者不能混写。",
        f"- 每格成功数都由 50 个布尔 outcome 重算；CSV `{ledger_evidence.csv_path}`（SHA-256 `{ledger_evidence.csv_sha256}`）与 JSON `{ledger_evidence.json_path}`（SHA-256 `{ledger_evidence.json_sha256}`）共同保留 {ledger_evidence.cell_count} 格 / {ledger_evidence.outcome_count:,} 个 outcomes。",
        "- EMA E3 的 G1/F+G 与 G2/F-only 使用隔离 retry attempt_02；原失败调度证据保留，不把失败单元伪装成原调度成功。",
        "- 原固定 E10 Mean-Q 覆盖 V0/V1/V2/V2-EMA × C/D/F/G1/G2/G3，共 24 格且无缺格；新增 First-Q2/State-V 不适用处用中性 `—`，不参与赢家计算。",
        "- 主结果表只展示 E10；E3–E10 全轨迹仍保存在 `all_o50_results.csv`，没有因版式精简而删除。",
        "- 只有一个 training seed；所有跨版本结果都是描述性结构消融，不声称多 seed 总体最优或统计显著。",
        "",
    ]
    return "\n".join(lines)


def _add_heading(
    document: Any, text: str, *, level: int = 1, page_break: bool = False
) -> Any:
    return v2_report._add_heading(document, text, level=level, page_break=page_break)


def _add_body(
    document: Any, text: str, *, color: str = "111827", bold: bool = False
) -> Any:
    return v2_report._add_body(document, text, color=color, bold=bold)


def _add_table(
    document: Any,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
) -> Any:
    return v2_report._v1._add_table(document, headers=headers, rows=rows, widths=widths)


def _set_cell_border(
    cell: Any,
    *,
    edges: Sequence[str],
    color: str,
    size: int,
) -> None:
    """Apply an explicit Word cell border without replacing unrelated edges."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in edges:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _add_fixed_master_legend(document: Any) -> Any:
    """Add a literal legend for the background-fill winner encodings."""

    table = _add_table(
        document,
        ("Marker", "Meaning", "Comparison scope"),
        (
            ("Yellow fill", "Best available value in this row", "Seven score columns; neutral missing cells are ignored; all ties are marked"),
            ("Blue fill", "Best available value in this column", "Methods inside the same dynamic version band only; all ties are marked"),
            ("Teal fill", "Both winner conditions", "Best in the row and best in the version-specific column"),
            ("Version band", "Version boundary", "Separates V0, V1, V2 and V2-EMA; it does not encode performance"),
        ),
        (2200, 5200, 7000),
    )
    v2_report._v1._shade_cell(table.rows[1].cells[0], "FFF2CC")
    v2_report._v1._shade_cell(table.rows[2].cells[0], "DDEBF7")
    v2_report._v1._shade_cell(table.rows[3].cells[0], "B7DEE8")
    v2_report._v1._shade_cell(table.rows[4].cells[0], "EDE9FE")
    return table


def _add_fixed_master_table(
    document: Any,
    cells: Sequence[ResultCell],
    endpoint_cells: Sequence[Any] | None = None,
) -> Any:
    """Add the 26-by-7 matrix with dynamic version bands and neutral gaps."""

    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    display_rows = _fixed_master_rows(cells, endpoint_cells)
    count_rows = _fixed_master_counts(cells, endpoint_cells)
    version_column_maxima = _fixed_master_version_column_maxima(count_rows)
    table_rows = tuple(
        (*row[:2], METHOD_LOSS_LABELS[row[1].lower()], *row[2:])
        for row in display_rows
    )
    table = _add_table(
        document,
        (
            "Version",
            "Method",
            "Training loss",
            "F-only",
            "G-only",
            "F+G tail",
            "First-Q",
            "Mean-Q",
            "First-Q2",
            "State-V",
        ),
        table_rows,
        (900, 700, 2300, 1500, 1500, 1500, 1500, 1500, 1500, 1500),
    )
    for column, cell in enumerate(table.rows[0].cells):
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if column < 3 else WD_ALIGN_PARAGRAPH.CENTER
        )
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            properties = cell._tc.get_or_add_tcPr()
            margins = properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                properties.append(margins)
            for side in ("top", "bottom"):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), "55")
                element.set(qn("w:type"), "dxa")
    group_styles = (
        ("E2E8F0", "475569"),
        ("D1FAE5", "047857"),
        ("FEF3C7", "B45309"),
        ("EDE9FE", "6D28D9"),
    )
    group_by_row = tuple(
        group_index
        for group_index, (_version, variants) in enumerate(MASTER_ROW_GROUPS)
        for _variant in variants
    )
    group_slices = _master_group_slices()
    for row_index, (row, counts) in enumerate(zip(table.rows[1:], count_rows)):
        version_index = group_by_row[row_index]
        group_fill, group_accent = group_styles[version_index]
        available = tuple(value for value in counts if value is not None)
        row_maximum = max(available) if available else None
        v2_report._v1._shade_cell(row.cells[0], group_fill)
        v2_report._v1._shade_cell(row.cells[1], "F8FAFC")
        v2_report._v1._shade_cell(row.cells[2], "F8FAFC")
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True
        loss_paragraph = row.cells[2].paragraphs[0]
        loss_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        loss_paragraph.runs[0].font.size = Pt(7.5)
        group_slice = group_slices[version_index]
        if row_index == group_slice.start:
            for cell in row.cells:
                _set_cell_border(
                    cell, edges=("top",), color=group_accent, size=24
                )
        if row_index == group_slice.stop - 1:
            for cell in row.cells:
                _set_cell_border(
                    cell, edges=("bottom",), color=group_accent, size=24
                )
        for score_index, count in enumerate(counts):
            cell = row.cells[score_index + 3]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(8.5)
            if count is None:
                # Missing endpoint measurements are displayed but deliberately
                # receive neither winner fill nor winner typography.
                continue
            row_best = row_maximum is not None and count == row_maximum
            column_maximum = version_column_maxima[version_index][score_index]
            column_best = column_maximum is not None and count == column_maximum
            if row_best and column_best:
                v2_report._v1._shade_cell(cell, "B7DEE8")
            elif column_best:
                v2_report._v1._shade_cell(cell, "DDEBF7")
            elif row_best:
                v2_report._v1._shade_cell(cell, "FFF2CC")
            run.bold = row_best or column_best

    # Merge the repeated version labels into one colored band. V1 has eight
    # rows while the other bands have six, so no fixed block size is used.
    for version_index, ((version, _variants), group_slice) in enumerate(
        zip(MASTER_ROW_GROUPS, group_slices)
    ):
        first = 1 + group_slice.start
        last = group_slice.stop
        merged = table.cell(first, 0).merge(table.cell(last, 0))
        merged.text = VERSION_LABELS[version]
        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        merged_run = merged.paragraphs[0].runs[0]
        merged_run.bold = True
        merged_run.font.size = Pt(10)
        merged_run.font.color.rgb = RGBColor.from_string(group_styles[version_index][1])
        v2_report._v1._shade_cell(merged, group_styles[version_index][0])
        _set_cell_border(
            merged,
            edges=("top", "bottom"),
            color=group_styles[version_index][1],
            size=24,
        )
    return table


def _configure_primary_document(document: Any) -> None:
    """Configure a standalone landscape report with no stale pages in front."""

    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    if "Report Kicker" not in document.styles:
        kicker = document.styles.add_style("Report Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = document.styles["Report Kicker"]
    kicker.font.name = "Arial Unicode MS"
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    kicker.font.size = Pt(9.5)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string("5C6975")
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(4)

    _set_section_running_matter(
        section,
        header_text=(
            f"Results TD complete ledger · Cube O50 · {COMPLETE_CELL_COUNT} verified cells"
        ),
        footer_prefix="Complete fixed-E10 decision report · Page ",
    )

    core = document.core_properties
    core.title = "Results TD Complete Experiment Ledger"
    core.subject = (
        f"{COMPLETE_CELL_COUNT} verified base O50 cells and 26-by-7 endpoint-ready comparison"
    )
    core.keywords = (
        "TD-JEPA, Actor-Free TD-LeWM, fixed endpoints, seven scores, "
        f"{COMPLETE_CELL_COUNT} cells, {COMPLETE_OUTCOME_COUNT} outcomes"
    )


def _set_section_running_matter(
    section: Any, *, header_text: str, footer_prefix: str
) -> None:
    """Populate default, even and first-page headers for renderer consistency."""

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for header in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header_text)
        v2_report._v1._set_run_font(run, size=8.5, color="6B7280")
    for footer in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(footer_prefix)
        v2_report._v1._set_run_font(run, size=8.5, color="6B7280")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        paragraph._p.append(field)


def _use_minimal_appendix_running_matter(document: Any) -> None:
    """Keep preserved pages intact and make every appended page style identical."""

    for section in document.sections:
        header_text = " ".join(
            paragraph.text for paragraph in section.header.paragraphs
        )
        if not header_text.startswith("Results TD complete ledger"):
            continue
        _set_section_running_matter(section, header_text="", footer_prefix="")


def build_docx(
    cells: Sequence[ResultCell],
    ledger_evidence: LedgerEvidence,
    endpoint_cells: Sequence[Any] | None = None,
    *,
    c3_epoch3_diagnostic: Any | None = None,
    endpoint_training_evidence: EndpointTrainingEvidence | None = None,
    endpoint_training_chart: bytes | None = None,
    v0_training_chart: bytes,
    v2_training_chart: bytes,
    training_chart: bytes,
    score_chart: bytes | None = None,
    base_document: str | Path | None = None,
) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required from the workspace runtime."
        ) from exc

    analysis = _fixed_analysis(cells)
    endpoint_analysis = _endpoint_analysis(
        cells, endpoint_cells, c3_epoch3_diagnostic
    )
    f_plus_g_winner = _winner_result(analysis, "f_plus_g")
    overall_winner = (
        f"{_compact_join(analysis.overall_winners)}: {analysis.overall_best_count}/50 "
        f"({analysis.overall_best_count * 2}%)"
    )
    best_score_names = _joined(
        tuple(MODE_LABELS[mode] for mode in analysis.best_mean_modes)
    )
    best_training_mean_names = _joined(analysis.best_training_mean_winners)
    ema_best_variants = _compact_join(
        tuple(variant.upper() for variant in analysis.ema_best_variants)
    )
    ema_best_mean = analysis.ema_variant_means[analysis.ema_best_variants[0]]
    v1_f = analysis.version_mode_means[("v1", "f_only")]
    v2_f = analysis.version_mode_means[("v2", "f_only")]
    mean_q_winner = _winner_result(analysis, "g_only_f_rollout_mean")
    tail_deltas = tuple(
        _percent(
            _find(cells, version=version, variant=variant, epoch=10, mode="f_plus_g")
        )
        - _percent(
            _find(cells, version=version, variant=variant, epoch=10, mode="f_only")
        )
        for version in VERSIONS
        for variant in VARIANTS
    )
    tail_gain_count = sum(delta > 0 for delta in tail_deltas)
    tail_tie_count = sum(delta == 0 for delta in tail_deltas)
    tail_harm_count = sum(delta < 0 for delta in tail_deltas)
    tail_mean_delta = sum(tail_deltas) / len(tail_deltas)
    if base_document is None:
        document = Document()
        _configure_primary_document(document)
    else:
        document = Document(str(base_document))
        # Preserve the previous Legacy and V0/V1 pages verbatim, then start a
        # fresh landscape section for the complete 477-cell decision report.
        v2_report._v1._configure_append_section(document)
        _set_section_running_matter(
            document.sections[-1],
            header_text=(
                f"Results TD complete ledger · Cube O50 · {COMPLETE_CELL_COUNT} verified cells"
            ),
            footer_prefix="Validated result archive · Page ",
        )
        for paragraph in document.paragraphs:
            if paragraph.text.startswith("Locked protocol:"):
                notice = paragraph.add_run(
                    " The complete seven-score cross-version section follows the preserved "
                    "Legacy and V0/V1 reference pages."
                )
                v2_report._v1._set_run_font(
                    notice, size=9.5, color="7A5A00", bold=True
                )
                break
        core = document.core_properties
        core.title = "Results TD Complete Experiment Ledger"
        core.subject = (
            f"{COMPLETE_CELL_COUNT} verified base O50 cells and 26-by-7 endpoint-ready comparison"
        )
        core.keywords = (
            "TD-JEPA, Actor-Free TD-LeWM, fixed endpoints, seven scores, "
            f"{COMPLETE_CELL_COUNT} cells, {COMPLETE_OUTCOME_COUNT} outcomes"
        )

    kicker = document.add_paragraph(style="Report Kicker")
    kicker.add_run("RESULTS TD / COMPLETE EXPERIMENT LEDGER")
    v2_report._font_paragraph(kicker, size=9.5, color="5C6975")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Complete 26-method, seven-score results and analysis")
    v2_report._v1._set_run_font(run, size=24, color="0B2545", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        f"Cube · seed 3072 · {COMPLETE_CELL_COUNT} audited base cells + optional 8-cell C2/C3 endpoints"
    )
    v2_report._v1._set_run_font(run, size=12, color="4B5563")
    _add_body(
        document,
        "The preceding pages preserve the detailed Legacy and V0/V1 method reference "
        "from the previous document. This final section is backed by the complete "
        f"{COMPLETE_CELL_COUNT}-cell audit, while its single master matrix also reserves the strict C2-E10/C3-E12 endpoint cells. Every displayed result "
        f"uses the same 50 start-goal pairs; the companion ledgers retain all "
        f"{COMPLETE_OUTCOME_COUNT:,} Boolean outcomes. All training uses one seed, so "
        "no statistical-significance claim is made.",
        color="7A5A00",
        bold=True,
    )

    _add_heading(document, "Decision summary")
    _add_table(
        document,
        ("Question", "Answer", "Evidence", "Decision"),
        (
            (
                "Best on prespecified F+G column",
                _compact_join(analysis.mode_winners["f_plus_g"]),
                f"{analysis.mode_best_counts['f_plus_g']}/50 "
                f"({analysis.mode_best_counts['f_plus_g'] * 2}%)",
                "Keep as primary baseline",
            ),
            (
                "Best fixed result overall",
                _compact_join(analysis.overall_winners),
                f"{analysis.overall_best_count}/50 ({analysis.overall_best_count * 2}%)",
                "Primary fixed-E10 baseline",
            ),
            (
                "Best equal-weight training average",
                best_training_mean_names,
                f"Five-score mean {analysis.best_training_mean_percent:.1f}%",
                "Treat as candidates, not a seed-robust winner",
            ),
            (
                "Best V2-EMA training variant",
                ema_best_variants,
                f"E10 five-score mean {ema_best_mean:.1f}%",
                "Retain this variant if joint training continues",
            ),
            (
                "Descriptive default test score",
                best_score_names,
                f"24-configuration mean {analysis.best_mean_percent:.1f}%",
                "Default readout candidate; confirm on dev pairs",
            ),
            (
                "Leading failure hypothesis",
                "Joint world-model drift",
                f"V1→V2 F-only {v1_f:.1f}%→{v2_f:.1f}%",
                f"Restore F toward the V1 mean ({v1_f:.1f}%)",
            ),
        ),
        (3300, 3400, 3600, 4100),
    )

    _add_heading(document, "Complete companion ledgers")
    _add_table(
        document,
        ("Artifact", "Coverage", "Repository-relative path", "SHA-256"),
        (
            (
                "CSV scalar ledger",
                f"{ledger_evidence.cell_count} O50 cells",
                ledger_evidence.csv_path,
                ledger_evidence.csv_sha256,
            ),
            (
                "JSON reconciliation ledger",
                f"{ledger_evidence.cell_count} x {EPISODES} = "
                f"{ledger_evidence.outcome_count:,} outcomes",
                ledger_evidence.json_path,
                ledger_evidence.json_sha256,
            ),
        ),
        (2600, 2900, 4500, 4400),
    )
    _add_body(
        document,
        f"The original analysis is fixed E10 only. The two fingerprinted files above retain "
        f"all {ledger_evidence.cell_count} scalar cells and all "
        f"{ledger_evidence.outcome_count:,} per-pair Boolean outcomes.",
        bold=True,
    )

    _add_heading(document, "How the seven evaluation methods are run", page_break=True)
    _add_body(
        document,
        "This section defines the seven columns in the 26-method matrix. z0 is the latent "
        "of the current real observation, z_g is the latent of the supplied goal image, "
        "and z_k^F is the imagined latent after k LeWM steps. The goal readout is "
        "w(g)=sqrt(192) z_g/||z_g||_2 and Q_G(z,A,g)=G(z,e,w(g))^T w(g).",
    )
    _add_table(
        document,
        ("Common protocol field", "Fixed setting"),
        COMMON_EVALUATION_PROTOCOL_ROWS,
        (3100, 11300),
    )
    _add_body(
        document,
        "Action input differs only by training version: V0 passes the normalized 25D "
        "action block directly to G; V1, V1-C2, V2 and V2-EMA use e=E_A(A) from the shared "
        "Action Encoder. V1-C3 State-V has no action input after F produces z5. Every A_k "
        "is one 25D block of five consecutive 5D primitive actions.",
        bold=True,
    )
    _add_table(
        document,
        ("Score", "Actual F and G path", "Cost minimized by CEM", "What is used"),
        EVALUATION_METHOD_ROWS,
        (1800, 4300, 4200, 4100),
    )
    _add_body(
        document,
        "V2-EMA still deploys the online F, online Action Encoder and online G. The EMA "
        "modules construct the stopped training target and are checked when loading the "
        "checkpoint, but they do not score CEM candidates. The Legacy G/C-only column is "
        "different: its older adapter first uses F to construct an H-1 tail state; only "
        "Direct Goal Critic Hybrid replaces G with a scalar C.",
        color="7A5A00",
        bold=True,
    )

    _add_heading(
        document,
        "26-method, seven-score master comparison and color legend",
    )
    _add_body(
        document,
        "Read across a row to compare all available scores for one training method; "
        "read down a column only inside one version block. V1 includes C2/C3 beside C. "
        "Yellow fill marks the best value in that row. Blue fill marks the best value "
        "in that score column within the same version. Teal fill marks a value that "
        "satisfies both conditions; all ties are marked. A neutral dash is missing and "
        "is ignored by both comparisons.",
        bold=True,
    )
    _add_fixed_master_legend(document)
    _add_body(
        document,
        "The Training loss column shows the method-specific G objective: l_i is "
        "the per-sample successor TD residual; qY = Y^T m; and w_i(.) is the normalized "
        "sample weight formed from a stopped signal. V0/V1 use the real branch. "
        "V2/V2-EMA optimize L_pred + 0.09 L_SIGReg + rho(L_method^real + "
        "L_method^pred). Exact C-G3 signals and weighting rules appear in "
        "V2 / V2-EMA network, target and loss.",
    )
    _add_body(
        document,
        "V1-C2 initializes every parameter from V1-C E10, freezes LeWM and its Action "
        "Encoder, and fine-tunes only G with L_C2=L_C+CE(p_F,p_Q): p_F is the softmax "
        "of negative candidate-normalized frozen-F terminal cost, while p_Q is the "
        "softmax of candidate-normalized Q_G(z0,A1,g). V1-C3 freezes the complete V1-C "
        "parent, including both G copies, and trains only an EMA-targeted MRN State-V "
        "critic with expectile-Huber TD (tau=0.03, gamma=0.98, n<=50 primitive steps).",
        bold=True,
    )
    _add_heading(
        document,
        "One 26 by 7 master matrix with training losses",
        level=2,
        page_break=True,
    )
    _add_fixed_master_table(document, cells, endpoint_cells)

    _add_heading(
        document,
        "Best training method and evaluation score",
        page_break=True,
    )
    _add_body(
        document,
        "There is no evaluation-independent training winner. The prespecified F+G "
        f"leader is {f_plus_g_winner}; under an equal-weight five-score average, "
        f"{best_training_mean_names} tie at {analysis.best_training_mean_percent:.1f}%; "
        f"the highest fixed cell is {overall_winner}. "
        f"Across all 24 configurations, {best_score_names} has the largest descriptive "
        f"mean ({analysis.best_mean_percent:.1f}%). These are single-seed comparisons.",
        color="7A5A00",
        bold=True,
    )
    _add_heading(document, "Within-version column winners", level=2)
    _add_table(
        document,
        (
            "Version",
            "F-only",
            "G-only",
            "F+G tail",
            "First-Q",
            "Mean-Q",
            "First-Q2",
            "State-V",
        ),
        _fixed_version_winner_rows(cells, endpoint_cells),
        (1450, 1850, 1850, 1850, 1850, 1850, 1850, 1850),
    )
    _add_body(
        document,
        "This table is the textual counterpart of the blue fills in the master "
        "matrix. It deliberately resets the column comparison at every version boundary.",
    )
    _add_heading(document, "V1-C2/C3 endpoint evidence", level=2)
    if endpoint_analysis.available:
        endpoint = _endpoint_index(endpoint_cells)
        endpoint_rows = []
        for mode in (*ALL_CONTROLLED_MODES, FIRST_Q2_MODE):
            parent_count = (
                endpoint_analysis.first_q2_parent_count
                if mode == FIRST_Q2_MODE
                else int(
                    _find(
                        cells,
                        version="v1",
                        variant="c",
                        epoch=10,
                        mode=mode,
                    )["success_count"]
                )
            )
            assert parent_count is not None
            c2_count = int(
                _endpoint_field(endpoint[("v1_c2", mode)], "success_count")
            )
            endpoint_rows.append(
                (
                    MODE_LABELS[mode],
                    f"{parent_count}/50",
                    f"{c2_count}/50",
                    f"{(c2_count - parent_count) * 2:+d} pp",
                )
            )
        _add_table(
            document,
            ("Shared score", "V1-C parent", "V1-C2", "C2-parent"),
            tuple(endpoint_rows),
            (3600, 3600, 3600, 3600),
        )
        _add_body(
            document,
            f"Across the six comparable scores, C2 is {endpoint_analysis.c2_improved} "
            f"up / {endpoint_analysis.c2_tied} tied / {endpoint_analysis.c2_harmed} down; "
            f"mean delta {endpoint_analysis.c2_mean_delta_percent:+.1f} pp. C3 State-V "
            f"is {endpoint_analysis.state_v_c3_count}/50 "
            f"({endpoint_analysis.state_v_c3_count * 2}%).",
            bold=True,
        )
        if endpoint_analysis.state_v_c3_epoch3_count is not None:
            contingency = (
                endpoint_analysis.state_v_c3_epoch3_vs_epoch12_contingency
            )
            assert contingency is not None
            _add_body(
                document,
                f"On the identical 50 pairs, the C3 early diagnostic is E3 "
                f"{endpoint_analysis.state_v_c3_epoch3_count}/50 "
                f"({endpoint_analysis.state_v_c3_epoch3_count * 2}%) versus the "
                f"final E12 endpoint {endpoint_analysis.state_v_c3_count}/50 "
                f"({endpoint_analysis.state_v_c3_count * 2}%). Paired outcomes: "
                f"both success {contingency['both_success']}, E3 only "
                f"{contingency['epoch3_only']}, E12 only "
                f"{contingency['epoch12_only']}, both failure "
                f"{contingency['both_failure']}; exact two-sided McNemar "
                f"p={endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p:.6g}. "
                "E3 is diagnostic only and is excluded from the eight-cell master table.",
            )
        parent_f_only = int(
            _find(cells, version="v1", variant="c", epoch=10, mode="f_only")[
                "success_count"
            ]
        )
        parent_first_q = int(
            _find(
                cells,
                version="v1",
                variant="c",
                epoch=10,
                mode="f_plus_g_first",
            )["success_count"]
        )
        assert endpoint_analysis.state_v_c3_count is not None
        assert endpoint_analysis.first_q2_parent_count is not None
        assert endpoint_analysis.first_q2_c2_count is not None
        first_q_delta_pp = (
            2 * endpoint_analysis.c2_shared_deltas["f_plus_g_first"]
        )
        first_q2_delta_pp = 2 * (
            endpoint_analysis.first_q2_c2_count
            - endpoint_analysis.first_q2_parent_count
        )
        _add_heading(document, "Endpoint interpretation and next steps", level=3)
        _add_body(
            document,
            f"C2 is mixed across the six shared scores: {endpoint_analysis.c2_improved} "
            f"up, {endpoint_analysis.c2_tied} tied and {endpoint_analysis.c2_harmed} "
            f"down, with mean delta {endpoint_analysis.c2_mean_delta_percent:+.1f} pp. "
            f"First-Q changes {first_q_delta_pp:+d} pp and First-Q2 changes "
            f"{first_q2_delta_pp:+d} pp. This single-seed evidence does not establish "
            "a robust cross-readout ranking gain.",
            bold=True,
        )
        _add_body(
            document,
            f"Final C3 State-V changes "
            f"{2 * (endpoint_analysis.state_v_c3_count - parent_f_only):+d} pp versus "
            f"V1-C F-only and {2 * (endpoint_analysis.state_v_c3_count - parent_first_q):+d} "
            "pp versus V1-C First-Q. Train State-V next on a controlled mixture of real "
            "latents and stopped-gradient F-imagined terminal latents to reduce the "
            "train/inference OOD gap.",
        )
        if (
            endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p
            is not None
        ):
            p_value = (
                endpoint_analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p
            )
            _add_body(
                document,
                f"The paired E3/E12 exact McNemar p-value is {p_value:.6g}. "
                + (
                    "It is not significant at 0.05, so E3 must not replace the final "
                    "endpoint post hoc. Select epochs on held-out development pairs or "
                    "pre-register the early-stop rule."
                    if p_value >= 0.05
                    else "Epoch selection must still use held-out development pairs or a "
                    "pre-registered early-stop rule rather than the formal O50 set."
                ),
            )
        _add_body(
            document,
            "All endpoint comparisons are descriptive: one training seed and the same 50 "
            "formal pairs do not establish multi-seed superiority or causality.",
            color="7A5A00",
            bold=True,
        )
    else:
        _add_body(
            document,
            "Pending strict endpoint evidence: V1-C First-Q2, six V1-C2 scores and "
            "V1-C3 State-V must arrive as one validated eight-cell ledger. Until then, "
            "the master matrix shows neutral dashes and no endpoint conclusion is generated.",
            color="7A5A00",
            bold=True,
        )
    _add_heading(document, "Fixed E10 means by training version", level=2)
    _add_table(
        document,
        (
            "Training version",
            "F-only",
            "G-only",
            "F+G",
            "First-Q",
            "Mean-Q",
            "Five-score mean",
        ),
        _fixed_version_mean_rows(cells),
        (3000, 1850, 1850, 1850, 1850, 1850, 2150),
    )
    _add_heading(document, "Fixed E10 means by evaluation score", level=2)
    _add_table(
        document,
        ("Score", "Versions", "Cells", "Four-version mean", "Best fixed cell"),
        _fixed_score_mean_rows(cells),
        (2500, 3000, 1500, 2400, 5000),
    )

    _add_heading(document, "Causes and next objectives", page_break=True)
    _add_table(
        document,
        ("Finding", "Evidence", "Interpretation", "Next objective"),
        (
            (
                "Action embedding helps G",
                f"V0→V1 G-only {analysis.version_mode_means[('v0', 'g_only')]:.1f}→"
                f"{analysis.version_mode_means[('v1', 'g_only')]:.1f}%; First-Q "
                f"{analysis.version_mode_means[('v0', 'f_plus_g_first')]:.1f}→"
                f"{analysis.version_mode_means[('v1', 'f_plus_g_first')]:.1f}%",
                "Association is consistent with a better critic readout",
                "Keep shared V1 Action Encoder",
            ),
            (
                "Joint tuning damages F",
                f"V1→V2 F-only {v1_f:.1f}→{v2_f:.1f}%",
                "Consistent with TD-gradient representation drift",
                f"Restore F-only mean to >={v1_f:.1f}%",
            ),
            (
                "EMA is insufficient",
                "Five-score mean delta "
                f"{_signed_pp(sum(analysis.version_mode_means[('v2_ema_sg', mode)] - analysis.version_mode_means[('v2', mode)] for mode in ALL_CONTROLLED_MODES) / len(ALL_CONTROLLED_MODES))}",
                "Target stabilization does not undo online drift",
                "Freeze/low-LR then staged unfreeze",
            ),
            (
                f"{best_score_names} leads descriptively",
                f"24-configuration mean {analysis.best_mean_percent:.1f}%",
                "Readout exposure to rollout OOD remains a testable hypothesis",
                "Use as default candidate; tune score parameters on dev",
            ),
            (
                "Tail effect is heterogeneous",
                f"F+G vs F-only: {tail_gain_count} up / {tail_tie_count} tied / "
                f"{tail_harm_count} down; mean {_signed_pp(tail_mean_delta)}",
                "Final imagined-state G can help or add OOD/scale error",
                "Gate/calibrate the tail on independent dev pairs",
            ),
            (
                "Mean-Q complete coverage",
                f"Mean {analysis.mode_means['g_only_f_rollout_mean']:.1f}%; {mean_q_winner}",
                "All four versions now contribute equally to the comparison",
                "Use the complete result to decide primary versus ablation status",
            ),
            (
                "Single-seed boundary",
                "All training seed 3072",
                "No population-level claim",
                ">=3 seeds + paired statistics",
            ),
        ),
        (2800, 4100, 4200, 3300),
    )
    _add_heading(document, "Reasoning from the fixed matrix", level=2)
    _add_body(
        document,
        f"First, V0 and V1 share the same F-only result for all six methods, while "
        f"the version mean falls from {analysis.version_mode_means[('v1', 'f_only')]:.1f}% "
        f"in V1 to {analysis.version_mode_means[('v2', 'f_only')]:.1f}% in V2. Because "
        "F-only does not call G at inference, the drop cannot be explained by a bad G "
        "readout alone. It is consistent with joint TD fine-tuning changing the deployed "
        "world-model representation or prediction quality.",
    )
    _add_body(
        document,
        f"Second, replacing raw actions with the shared Action Encoder raises the version "
        f"mean from {analysis.version_mode_means[('v0', 'g_only')]:.1f}% to "
        f"{analysis.version_mode_means[('v1', 'g_only')]:.1f}% for G-only and from "
        f"{analysis.version_mode_means[('v0', 'f_plus_g_first')]:.1f}% to "
        f"{analysis.version_mode_means[('v1', 'f_plus_g_first')]:.1f}% for First-Q. "
        "This supports keeping the semantic action embedding, although one training seed "
        "does not establish causality.",
    )
    _add_body(
        document,
        f"Third, F+G tail improves {tail_gain_count} of 24 fixed configurations, ties "
        f"{tail_tie_count}, and hurts {tail_harm_count}; its mean change from F-only is "
        f"{_signed_pp(tail_mean_delta)}. The tail reads G at the deepest imagined state, "
        "so rollout distribution shift and Q scale can either help or interfere. First-Q "
        "keeps Q on the real z0 and therefore avoids that specific source of error.",
    )
    _add_body(
        document,
        f"Finally, Mean-Q averages five Q values but four of them use imagined predecessor "
        f"states. Its 24-configuration mean is {analysis.mode_means['g_only_f_rollout_mean']:.1f}%. "
        "Averaging reduces dependence on one tail state, but it also repeats exposure to "
        "imagined-state error. The per-version and per-method results therefore matter more "
        "than the name of the aggregation rule.",
    )
    _add_body(
        document,
        "Recommended next experiment: freeze the V1 encoder/world model, compare "
        f"{_compact_join(analysis.overall_winners)} against "
        f"{_compact_join(analysis.mode_winners['f_plus_g'])} + F+G tail on independent development pairs, "
        "pre-register epoch/alpha, then run at least three training seeds. If joint "
        "training is retained, use a much smaller TD learning rate, latent/prediction "
        "anchors to V1 and conservative calibration on CEM candidate actions.",
        bold=True,
    )

    _add_heading(document, "Coverage map", page_break=True)
    _add_table(
        document,
        ("Family/version", "Methods", "Epochs", "Scores", "Cells"),
        (
            ("Legacy structures", "7", "E10", "F / G(C) / combined", "21"),
            ("V0 raw action", "6", "E10", "All five", "30"),
            ("V1 action encoder", "6", "E10", "All five", "30"),
            ("V2 joint", "6", "E3-E10", "3 original + E10 First/Mean", "156"),
            ("V2-EMA-SG", "6", "E3-E10", "All five", "240"),
            (
                "V1-C2/C3 endpoints",
                "2",
                "C2 E10 / C3 E12",
                "First-Q2 + State-V integrated in master",
                "8" if endpoint_analysis.available else "Pending",
            ),
            ("TOTAL", "-", "-", "Same O50 selection", str(COMPLETE_CELL_COUNT)),
        ),
        (3600, 1800, 2200, 4900, 1900),
    )
    _add_body(
        document,
        "The original 24-row comparison deliberately uses fixed E10; C2 uses final E10 "
        "and C3 uses final E12. The full V2 "
        "and V2-EMA E3-E10 trajectories are reproduced later as diagnostics and remain "
        "available in the audited companion ledgers.",
    )

    _add_heading(document, "V2 / V2-EMA network, target and loss", page_break=True)
    _add_body(
        document,
        "V2 jointly updates online LeWM, the shared 25D-to-192D Action Encoder and "
        "G. V2-EMA-SG uses an EMA world model, EMA Action Encoder and EMA G to form "
        "the fully stopped next-frame target. Neither version trains an Actor or reward model.",
    )
    v2_report._add_formula(
        document,
        "TD target",
        "Y_t = sg[zbar_(t+1) + gamma(1-d_t) Gbar(zbar_(t+1),ebar_(t+1),m_t)]",
        "zbar comes from the real next frame; ebar comes from the known dataset next action. gamma=0.95.",
    )
    v2_report._add_formula(
        document,
        "Total",
        "L_total = L_pred + 0.09 L_SIGReg + rho(u)(L_method^real + L_method^pred)",
        "The predicted branch is differentiable to online LeWM/Action Encoder; EMA modules only form targets.",
    )
    _add_body(
        document,
        "For branch b in {real,pred}, the per-sample base residual is "
        "l_i^b = ||G_phi(s_i^b,e_i,m_i)-Y_i||^2. Here m_i is the goal/random "
        "task vector and rho(u) is the TD warm-up. C adds a scalar goal residual; "
        "D/F/G1/G2/G3 use stopped advantage signals to reweight l_i^b.",
    )
    _add_table(
        document,
        ("Method", "Special signal", "Per-branch loss", "Role"),
        tuple(
            tuple(str(value) for value in row) for row in v2_report.METHOD_FORMULA_ROWS
        )
        + (
            (
                "C2 (V1 only)",
                "Frozen-F candidate ranking teacher",
                "L_C2=L_C+CE(softmax(-z(J_F)),softmax(z(Q_first)))",
                "Fine-tune G only; align real-state First-Q with the five-block F planner",
            ),
            (
                "C3 (V1 only)",
                "Same-episode primitive temporal cost",
                "E[omega_tau(r)Huber_1(r)], r=V_psi-sg(y), tau=0.03",
                "Train only MRN State-V; frozen V1-C parent; EMA bootstrap",
            ),
        ),
        (1200, 4300, 4300, 4600),
    )

    _add_heading(document, "Seven inference scores", page_break=True)
    _add_table(
        document,
        ("Score", "Planner signal", "G location", "Main risk"),
        (
            (
                "F-only",
                "Five-step terminal latent distance",
                "None",
                "World-model drift",
            ),
            (
                "G-only",
                "-Q_G(z0,A1,g), H=1",
                "Real z0 / first action",
                "Weak without F",
            ),
            (
                "F+G tail",
                "F prefix + final G tail",
                "Final imagined state",
                "Tail OOD / interference",
            ),
            (
                "First-Q",
                "F terminal cost - 0.25 Q_G(z0,A1,g)",
                "Real z0 / first action",
                "Q scale / alpha",
            ),
            (
                "Mean-Q",
                "-(1/5) sum Q over F rollout",
                "Five real/imagined predecessors",
                "Accumulated drift / OOD",
            ),
            (
                "First-Q2",
                "z_cand(F terminal cost) - 0.25 z_cand(Q_G(z0,A1,g))",
                "Real z0 / first action",
                "Candidate-set normalization / alpha",
            ),
            (
                "State-V",
                "EMA V_bar(F^5(z0,A1:5),z_g)",
                "G is not used",
                "Temporal-value calibration / imagined terminal state",
            ),
        ),
        (2500, 5000, 3700, 3200),
    )

    _add_heading(document, "Training / validation trajectories", page_break=True)
    _add_body(
        document,
        "Training totals are method-specific and cannot be ranked by absolute height. "
        "Use them only as within-method convergence diagnostics. The following charts "
        "show V0, V2 and V2-EMA; the preserved opening pages contain the Legacy and V1 "
        "curves. When hash-bound endpoint training metrics are present, the final compact "
        "figure adds C2/C3 convergence and diagnostics.",
    )
    v2_report._add_picture(
        document,
        v0_training_chart,
        title="Figure 1. V0 training and validation total loss (E1-E10)",
        description="Raw-action C, D, F, G1, G2 and G3. Log scale; compare convergence within a method, not absolute objective height across methods.",
    )
    v2_report._add_picture(
        document,
        v2_training_chart,
        title="Figure 2. V2 joint training and validation total loss (E1-E10)",
        description="Joint LeWM/Action Encoder/G fine-tuning trajectories for all six methods.",
    )
    v2_report._add_picture(
        document,
        training_chart,
        title="Figure 3. V2-EMA-SG training and validation loss (E1-E10)",
        description="Training and validation total-loss trajectories for C, D, F, G1, G2 and G3.",
    )
    if score_chart is not None:
        v2_report._add_picture(
            document,
            score_chart,
            title="Figure 4. V2-EMA First-Q and Mean-Q trajectories (E3-E10)",
            description="Formal O50 checkpoint trajectories for the two added evaluation scores across C, D, F, G1, G2 and G3.",
        )
    if endpoint_training_evidence is not None:
        _add_heading(document, "V1-C2/C3 endpoint training evidence", level=2)
        _add_table(
            document,
            ("Method", "Metric", "First epoch", "Final epoch", "Change"),
            _endpoint_training_summary_rows(endpoint_training_evidence),
            (1900, 4100, 2600, 2600, 3200),
        )
        c2_scale_ratio = 100.0 * (
            endpoint_training_evidence.c2_alignment_loss[-1]
            / endpoint_training_evidence.c2_train_loss[-1]
        )
        _add_body(
            document,
            f"C2 ranking CE ends at {endpoint_training_evidence.c2_alignment_loss[-1]:.4f} "
            f"versus train total {endpoint_training_evidence.c2_train_loss[-1]:.2f} "
            f"({c2_scale_ratio:.4f}% by raw value), while top-1 agreement changes only "
            f"from {100 * endpoint_training_evidence.c2_alignment_top1[0]:.2f}% to "
            f"{100 * endpoint_training_evidence.c2_alignment_top1[-1]:.2f}%. This does "
            "not prove a gradient-scale cause, but supports testing normalized losses and "
            "adaptive or gradient-balanced weighting instead of another fixed weight.",
            bold=True,
        )
        if endpoint_training_chart is not None:
            figure_number = 5 if score_chart is not None else 4
            v2_report._add_picture(
                document,
                endpoint_training_chart,
                title=(
                    f"Figure {figure_number}. V1-C2 and V1-C3 training, validation and "
                    "diagnostic trajectories"
                ),
                description=(
                    "C2 covers E1-E10 and C3 covers E1-E12. Loss magnitudes are only "
                    "comparable within a method; C2 top-1 random reference is 6.25%."
                ),
            )
    elif endpoint_analysis.available:
        _add_body(
            document,
            "The strict endpoint scores are present, but the archive does not declare both "
            "C2 and C3 training metrics. No training chart or inferred values are added.",
            color="7A5A00",
            bold=True,
        )

    for version, modes, label in (
        ("v2", ORIGINAL_MODES, "V2"),
        ("v2_ema_sg", ALL_CONTROLLED_MODES, "V2-EMA"),
    ):
        _add_heading(document, f"{label} exact checkpoint trajectories", page_break=True)
        _add_body(
            document,
            "These E3-E10 tables are diagnostic only. The main matrix remains fixed at "
            "E10, so no value below is substituted after looking at the formal O50 pairs.",
            color="7A5A00",
            bold=True,
        )
        for index, mode in enumerate(modes):
            if label == "V2-EMA" and mode == "f_plus_g":
                v2_report._v1._configure_append_section(document)
                _set_section_running_matter(
                    document.sections[-1],
                    header_text=(
                        "Results TD complete ledger · Cube O50 · "
                        f"{COMPLETE_CELL_COUNT} verified cells"
                    ),
                    footer_prefix="Validated result archive · Page ",
                )
            trajectory_heading = {
                "f_plus_g": "F and G tail",
                "f_plus_g_first": "F with first-Q",
            }.get(mode, MODE_LABELS[mode])
            _add_heading(
                document,
                trajectory_heading,
                level=2,
                page_break=index > 0
                and not (label == "V2-EMA" and mode == "f_plus_g"),
            )
            _add_table(
                document,
                ("Epoch", "C", "D", "F", "G1", "G2", "G3"),
                _trajectory_rows(cells, version, mode),
                (1600, 2130, 2130, 2130, 2130, 2130, 2150),
            )

    _add_heading(document, "Post-hoc checkpoint diagnosis", page_break=True)
    _add_body(
        document,
        "The maxima in this section were selected after observing E3-E10 on the same "
        "O50 pairs. They describe instability and selection risk; they are not formal "
        "replacements for the fixed E10 results.",
        color="9C2F17",
        bold=True,
    )
    _add_table(
        document,
        ("Version", "Score", "Post-hoc best", "Best", "Best at E10"),
        _global_posthoc_rows(cells),
        (1900, 3000, 3700, 2900, 2900),
    )
    _add_heading(document, "V2-EMA new-score stability", level=2)
    _add_table(
        document,
        ("Method", "Score", "Mean", "Sigma pp", "Best", "E10"),
        _ema_new_stability_rows(cells),
        (1500, 3600, 2100, 2200, 2900, 2100),
    )
    v2_report._v1._configure_append_section(document)
    _set_section_running_matter(
        document.sections[-1],
        header_text=(
            f"Results TD complete ledger · Cube O50 · {COMPLETE_CELL_COUNT} verified cells"
        ),
        footer_prefix="Validated result archive · Page ",
    )
    _add_heading(document, "V2-EMA training variants", level=2)
    _add_table(
        document,
        ("Training variant", "Five-score mean", "Best readout", "Best rate"),
        _ema_variant_mean_rows(cells),
        (2600, 3100, 5600, 3100),
    )

    _add_heading(document, "Audit boundary", page_break=True)
    _add_table(
        document,
        ("Audit field", "Locked value"),
        (
            (
                "Verified O50 cells",
                f"{COMPLETE_CELL_COUNT} / {COMPLETE_CELL_COUNT}",
            ),
            ("Scalar CSV path", ledger_evidence.csv_path),
            ("Scalar CSV SHA-256", ledger_evidence.csv_sha256),
            ("Reconciliation JSON path", ledger_evidence.json_path),
            ("Reconciliation JSON SHA-256", ledger_evidence.json_sha256),
            (
                "Per-pair Boolean outcomes",
                f"{ledger_evidence.outcome_count:,} / {COMPLETE_OUTCOME_COUNT:,}",
            ),
            ("Fixed E10 Mean-Q cells", "24 / 24"),
            (
                "C2/C3 endpoint extension",
                "8 / 8 strict cells" if endpoint_analysis.available else "Pending; neutral placeholders only",
            ),
            ("Episode selection", SELECTION_SHA256),
            ("Fixed valid-row-ranks hash", FIXED_SELECTION_RANKS_SHA256),
            ("Action normalization", ACTION_NORMALIZATION_SHA256),
            ("Training / planning seed", "3072 / 42"),
            ("EMA E3 retry", "G1/F+G and G2/F-only use isolated attempt_02"),
            (
                "Claim boundary",
                "One training seed; base rows fixed E10, C2 E10 and C3 E12",
            ),
        ),
        (3500, 10900),
    )

    if base_document is None:
        _add_heading(
            document,
            "Historical source document",
            page_break=True,
        )
        _add_body(
            document,
            "The detailed Legacy and V0/V1 reference is stored at "
            f"{_display_path(Path(DEFAULT_BASE_DOCUMENT))}. Use the current fixed-E10 "
            "matrix and its fingerprinted ledgers for cross-version comparisons.",
            color="7A5A00",
            bold=True,
        )

    if base_document is not None:
        _use_minimal_appendix_running_matter(document)
    stream = io.BytesIO()
    document.save(stream)
    payload = stream.getvalue()
    if not payload.startswith(b"PK"):
        raise RuntimeError("python-docx did not produce OOXML.")
    return payload


def _atomic_write(path: Path, payload: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary: str | None = None
    try:
        with tempfile.NamedTemporaryFile(
            mode="wb",
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as stream:
            temporary = stream.name
            stream.write(payload)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, path)
        temporary = None
    finally:
        if temporary is not None:
            Path(temporary).unlink(missing_ok=True)


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--ledger", default=str(DEFAULT_LEDGER))
    parser.add_argument(
        "--reconciliation-ledger", default=str(DEFAULT_RECONCILIATION_LEDGER)
    )
    parser.add_argument(
        "--endpoint-extension-ledger",
        default=None,
        help=(
            "Optional strict eight-cell V1-C/C2/C3 endpoint ledger. "
            f"Expected archive location: {DEFAULT_ENDPOINT_EXTENSION_LEDGER}"
        ),
    )
    parser.add_argument("--markdown-output", default=str(DEFAULT_MARKDOWN_OUTPUT))
    parser.add_argument("--docx-output", default=str(DEFAULT_DOCX_OUTPUT))
    parser.add_argument("--root-docx-output", default=str(DEFAULT_ROOT_DOCX_OUTPUT))
    parser.add_argument("--chart-dir", default=str(DEFAULT_CHART_DIR))
    parser.add_argument("--v0-training-csv", default=str(DEFAULT_V0_TRAINING_CSV))
    parser.add_argument("--v2-training-csv", default=str(DEFAULT_V2_TRAINING_CSV))
    parser.add_argument("--validate-only", action="store_true")
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    cells = load_complete_ledger(args.ledger)
    ledger_evidence = load_ledger_evidence(args.ledger, args.reconciliation_ledger)
    endpoint_cells: Sequence[Any] | None = None
    c3_epoch3_diagnostic: Any | None = None
    endpoint_training_evidence: EndpointTrainingEvidence | None = None
    if args.endpoint_extension_ledger is not None:
        from tdwm.results.actor_free_td_lewm_v1_c2_c3 import (
            load_endpoint_extension,
        )

        endpoint_extension = load_endpoint_extension(args.endpoint_extension_ledger)
        endpoint_cells = endpoint_extension.cells
        c3_epoch3_diagnostic = endpoint_extension.c3_epoch3_diagnostic
        endpoint_training_evidence = load_endpoint_training_evidence(
            args.endpoint_extension_ledger
        )
    if args.validate_only:
        endpoint_message = (
            " and 8 strict endpoint cells"
            if endpoint_cells is not None
            else " (endpoint extension not supplied)"
        )
        print(
            "PASS: complete Results TD ledger contains "
            f"{COMPLETE_CELL_COUNT} verified O50 cells and "
            f"{COMPLETE_OUTCOME_COUNT} reconciled outcomes{endpoint_message}."
        )
        return 0
    base_inputs = v2_report.load_validated_report_inputs(
        original_summary_path=v2_report.DEFAULT_ORIGINAL_SUMMARY,
        training_loss_csv_path=v2_report.DEFAULT_TRAINING_LOSS_CSV,
        new_summary_path=v2_report.DEFAULT_NEW_SUMMARY,
        new_results_csv_path=v2_report.DEFAULT_NEW_RESULTS_CSV,
        fixed_results_csv_path=v2_report.DEFAULT_FIXED_RESULTS_CSV,
        base_document_path=DEFAULT_BASE_DOCUMENT,
    )
    v0_training_chart = build_training_chart_from_archive(
        args.v0_training_csv, title="V0 raw-action training and validation loss"
    )
    v2_training_chart = build_training_chart_from_archive(
        args.v2_training_csv, title="V2 joint training and validation loss"
    )
    training_chart = v2_report.build_training_loss_chart(base_inputs)
    score_chart = v2_report.build_new_score_chart(base_inputs)
    endpoint_training_chart = (
        None
        if endpoint_training_evidence is None
        else build_endpoint_training_chart(endpoint_training_evidence)
    )
    chart_dir = Path(args.chart_dir)
    endpoint_chart_path = chart_dir / ENDPOINT_TRAINING_CHART_FILENAME
    markdown = build_markdown(
        cells,
        ledger_evidence,
        endpoint_cells,
        c3_epoch3_diagnostic,
        endpoint_training_evidence=endpoint_training_evidence,
        endpoint_training_chart_path=(
            None
            if endpoint_training_chart is None
            else _display_path(endpoint_chart_path)
        ),
    )
    document = build_docx(
        cells,
        ledger_evidence,
        endpoint_cells,
        c3_epoch3_diagnostic=c3_epoch3_diagnostic,
        endpoint_training_evidence=endpoint_training_evidence,
        endpoint_training_chart=endpoint_training_chart,
        v0_training_chart=v0_training_chart,
        v2_training_chart=v2_training_chart,
        training_chart=training_chart,
        score_chart=score_chart,
        base_document=base_inputs.base_document,
    )
    markdown_path = Path(args.markdown_output)
    docx_path = Path(args.docx_output)
    root_docx_path = Path(args.root_docx_output)
    _atomic_write(markdown_path, markdown.encode("utf-8"))
    _atomic_write(docx_path, document)
    _atomic_write(root_docx_path, document)
    _atomic_write(chart_dir / "v2_ema_training_validation_loss.png", training_chart)
    _atomic_write(chart_dir / "v2_ema_new_score_trajectories.png", score_chart)
    _atomic_write(chart_dir / "v0_training_validation_loss.png", v0_training_chart)
    _atomic_write(chart_dir / "v2_training_validation_loss.png", v2_training_chart)
    if endpoint_training_chart is not None:
        _atomic_write(endpoint_chart_path, endpoint_training_chart)
    print(f"Wrote {markdown_path}")
    print(f"Wrote {docx_path}")
    print(f"Updated {root_docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
