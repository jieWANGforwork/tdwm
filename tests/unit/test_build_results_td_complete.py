from __future__ import annotations

import base64
import csv
import hashlib
import importlib.util
import io
import json
import sys
from pathlib import Path
from typing import Any

import pytest

REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
MODULE_PATH = REPOSITORY_ROOT / "scripts/build_results_td_complete.py"
SPEC = importlib.util.spec_from_file_location(
    "build_results_td_complete_test", MODULE_PATH
)
assert SPEC is not None and SPEC.loader is not None
report = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = report
SPEC.loader.exec_module(report)


def _sha(label: str) -> str:
    return hashlib.sha256(label.encode()).hexdigest()


def _success_count(version: str, variant: str, mode: str) -> int:
    if (version, variant, mode) == ("v0", "d", "g_only_f_rollout_mean"):
        return 30
    if (version, variant, mode) == ("v1", "c", "f_plus_g_first"):
        return 21
    if (version, variant, mode) == ("v1", "g3", "f_plus_g"):
        return 20
    return 10


def _cells() -> tuple[Any, ...]:
    cells = []
    for version, variant, epoch, mode in sorted(report._expected_identities()):
        count = _success_count(version, variant, mode)
        cells.append(
            report.ResultCell(
                {
                    "version": version,
                    "variant": variant,
                    "method_id": variant,
                    "epoch": epoch,
                    "score_mode": mode,
                    "success_count": count,
                    "success_rate_percent": count * 2.0,
                }
            )
        )
    return tuple(cells)


def _ledger_evidence() -> Any:
    return report.LedgerEvidence(
        csv_path=(
            "reports/artifacts/actor_free_td_lewm_complete_cube_seed3072/"
            "all_o50_results.csv"
        ),
        csv_sha256="a" * 64,
        json_path=(
            "reports/artifacts/actor_free_td_lewm_complete_cube_seed3072/"
            "reconciliation_ledger.json"
        ),
        json_sha256="b" * 64,
        cell_count=477,
        outcome_count=23_850,
    )


def _endpoint_cells() -> tuple[dict[str, Any], ...]:
    counts = {
        ("v1_c", "first_q2"): 26,
        ("v1_c2", "f_only"): 27,
        ("v1_c2", "g_only"): 11,
        ("v1_c2", "f_plus_g"): 18,
        ("v1_c2", "f_plus_g_first"): 28,
        ("v1_c2", "g_only_f_rollout_mean"): 14,
        ("v1_c2", "first_q2"): 29,
        ("v1_c3", "state_v"): 31,
    }
    return tuple(
        {
            "method_key": method,
            "score_mode": mode,
            "success_count": count,
            "success_rate_percent": count * 2.0,
        }
        for (method, mode), count in counts.items()
    )


def _endpoint_cells_with_c3_outcomes(
    *, c3_count: int = 26
) -> tuple[dict[str, Any], ...]:
    cells = [dict(cell) for cell in _endpoint_cells()]
    c3 = next(
        cell
        for cell in cells
        if (cell["method_key"], cell["score_mode"]) == ("v1_c3", "state_v")
    )
    c3["success_count"] = c3_count
    c3["success_rate_percent"] = c3_count * 2.0
    c3["outcomes"] = [True] * c3_count + [False] * (50 - c3_count)
    return tuple(cells)


def _c3_epoch3_diagnostic(*, count: int = 28) -> dict[str, Any]:
    return {
        "method_key": "v1_c3",
        "score_mode": "state_v",
        "checkpoint_epoch": 3,
        "checkpoint_global_step": 3_000,
        "success_count": count,
        "success_rate_percent": count * 2.0,
        "outcomes": [True] * count + [False] * (50 - count),
    }


def _tiny_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
        "x8AAusB9WlT3r8AAAAASUVORK5CYII="
    )


def _write_endpoint_training_fixture(root: Path) -> Path:
    specifications = {
        "v1_c2": (
            10,
            {
                "train/loss_epoch": lambda epoch: 1000.0 - 10.0 * epoch,
                "validation/loss": lambda epoch: 100.0 - 2.0 * epoch,
                "train/first_q_alignment_loss_epoch": lambda epoch: 3.2
                - 0.005 * epoch,
                "train/first_q_alignment_top1_agreement_epoch": lambda epoch: 0.08
                + 0.001 * epoch,
            },
        ),
        "v1_c3": (
            12,
            {
                "train/loss_epoch": lambda epoch: 0.4 - 0.005 * epoch,
                "validation/loss": lambda epoch: 0.32 - 0.002 * epoch,
                "validation/mc_mae": lambda epoch: 12.8 - 0.05 * epoch,
                "validation/td_residual_mae": lambda epoch: 9.9 - 0.04 * epoch,
                "validation/spearman": lambda epoch: 0.64 + 0.0005 * epoch,
            },
        ),
    }
    training_sources = {}
    for method, (epoch_count, metrics) in specifications.items():
        path = root / "training" / method / "metrics.csv"
        path.parent.mkdir(parents=True, exist_ok=True)
        with path.open("w", newline="", encoding="utf-8") as stream:
            writer = csv.DictWriter(
                stream, fieldnames=("epoch", "step", *metrics), lineterminator="\n"
            )
            writer.writeheader()
            for epoch in range(epoch_count):
                writer.writerow(
                    {
                        "epoch": epoch,
                        "step": (epoch + 1) * 100,
                        **{name: function(epoch) for name, function in metrics.items()},
                    }
                )
        training_sources[method] = {
            "archive_directory": f"training/{method}",
            "files_sha256": {"metrics.csv": hashlib.sha256(path.read_bytes()).hexdigest()},
        }
    ledger = root / "reconciliation_ledger.json"
    ledger.write_text(json.dumps({"training_sources": training_sources}))
    return ledger


def _ledger_row(identity: tuple[str, str, int, str]) -> dict[str, str]:
    version, variant, epoch, mode = identity
    count = _success_count(version, variant, mode)
    family = (
        "actor_free_td_lewm" if version == "legacy" else f"actor_free_td_lewm_{version}"
    )
    method_id = "actor_free_td_lewm" if version == "legacy" else f"{family}_{variant}"
    return {
        "cell_id": f"fixture__{version}__{variant}__e{epoch:02d}__{mode}",
        "family": family,
        "version": version,
        "method_id": method_id,
        "method_label": variant,
        "variant": variant,
        "epoch": str(epoch),
        "global_step": str(epoch * report.STEPS_PER_EPOCH),
        "checkpoint_sha256": _sha(f"checkpoint:{version}:{variant}:{epoch}"),
        "score_mode": mode,
        "score_label": report.LEDGER_SCORE_LABELS[mode],
        "success_count": str(count),
        "episode_count": str(report.EPISODES),
        "success_rate": str(count / report.EPISODES),
        "success_rate_percent": str(count * 2.0),
        "training_seed": str(report.TRAINING_SEED),
        "planning_seed": str(report.PLANNING_SEED),
        "selection_sha256": report.SELECTION_SHA256,
        "action_normalization_sha256": report.ACTION_NORMALIZATION_SHA256,
        "training_commit": "abcdef0",
        "evaluation_commit": "abcdef0",
        "planning_horizon": ("1" if version != "legacy" and mode == "g_only" else "5"),
        "g_first_weight": "0.25" if mode == "f_plus_g_first" else "",
        "status": "VERIFIED",
        "comparison_role": "fixture",
        "source_kind": "fixture",
        "source_path": "/fixture",
        "source_results_sha256": _sha(f"results:{version}:{variant}:{epoch}:{mode}"),
        "source_protocol_sha256": _sha(f"protocol:{version}:{variant}:{epoch}:{mode}"),
        "outcomes_sha256": _sha(f"outcomes:{version}:{variant}:{epoch}:{mode}"),
    }


def _write_ledger(path: Path, identities: list[tuple[str, str, int, str]]) -> None:
    with path.open("w", newline="", encoding="utf-8") as stream:
        writer = csv.DictWriter(stream, fieldnames=report.FIELDS, lineterminator="\n")
        writer.writeheader()
        writer.writerows(_ledger_row(identity) for identity in identities)


def test_complete_grid_requires_477_cells_and_24_fixed_mean_q(tmp_path: Path) -> None:
    identities = sorted(report._expected_identities())
    assert len(identities) == report.COMPLETE_CELL_COUNT == 477
    ledger = tmp_path / "ledger.csv"
    _write_ledger(ledger, identities)

    cells = report.load_complete_ledger(ledger)

    fixed_mean_q = [
        cell
        for cell in cells
        if cell["version"] in report.VERSIONS
        and cell["epoch"] == 10
        and cell["score_mode"] == "g_only_f_rollout_mean"
    ]
    assert len(fixed_mean_q) == 24

    old_ledger = tmp_path / "old-ledger.csv"
    _write_ledger(old_ledger, identities[:465])
    with pytest.raises(report.CompleteResultsError, match="477 cells"):
        report.load_complete_ledger(old_ledger)


def test_base_analysis_stays_24_by_5_while_master_expands_to_26_by_7() -> None:
    cells = _cells()

    matrix = report._fixed_master_rows(cells, _endpoint_cells())
    analysis = report._fixed_analysis(cells)
    markdown = report.build_markdown(cells, _ledger_evidence(), _endpoint_cells())

    assert len(matrix) == 26
    assert all(len(row) == 9 for row in matrix)
    assert sum("—" in row for row in matrix) > 0
    assert next(row for row in matrix if row[:2] == ("V1", "C2"))[2:8] == (
        "27/50 (54%)",
        "11/50 (22%)",
        "18/50 (36%)",
        "28/50 (56%)",
        "14/50 (28%)",
        "29/50 (58%)",
    )
    assert next(row for row in matrix if row[:2] == ("V1", "C3"))[-1] == (
        "31/50 (62%)"
    )
    assert analysis.best_mean_modes == ("g_only_f_rollout_mean",)
    assert analysis.mode_winners["g_only_f_rollout_mean"] == ("V0-D",)
    assert analysis.overall_winners == ("V0-D + Mean-Q rollout",)
    assert "V0-D + Mean-Q rollout: 30/50 (60%)" in markdown
    assert "描述性领先测试评分为 Mean-Q rollout" in markdown
    assert "Mean-Q rollout = V0-D: 30/50 (60%)" in markdown
    assert "24 个固定训练配置上的均值" in markdown
    assert "all_o50_results.csv" in markdown
    assert "reconciliation_ledger.json" in markdown
    assert "23,850" in markdown
    assert "26 个方法 × 7 种评分" in markdown
    assert "C2 在 6 个可比评分上" in markdown


def test_missing_endpoint_ledger_produces_neutral_placeholders_only() -> None:
    cells = _cells()

    counts = report._fixed_master_counts(cells)
    matrix = report._fixed_master_rows(cells)
    markdown = report.build_markdown(cells, _ledger_evidence())

    c2_index = report._master_row_metadata().index(("v1", "c2"))
    c3_index = report._master_row_metadata().index(("v1", "c3"))
    assert counts[c2_index] == (None,) * 7
    assert counts[c3_index] == (None,) * 7
    assert matrix[c2_index][2:] == ("—",) * 7
    assert "尚未提供严格 endpoint ledger" in markdown


def test_endpoint_training_metrics_are_hash_bound_complete_and_reported(
    tmp_path: Path,
) -> None:
    ledger = _write_endpoint_training_fixture(tmp_path)
    evidence = report.load_endpoint_training_evidence(ledger)

    assert evidence is not None
    assert evidence.c2_epochs == tuple(range(1, 11))
    assert evidence.c3_epochs == tuple(range(1, 13))
    assert evidence.c2_validation_loss[-1] == pytest.approx(82.0)
    assert evidence.c3_spearman[-1] == pytest.approx(0.6455)
    chart = report.build_endpoint_training_chart(evidence)
    assert chart.startswith(b"\x89PNG\r\n\x1a\n")

    markdown = report.build_markdown(
        _cells(),
        _ledger_evidence(),
        _endpoint_cells_with_c3_outcomes(),
        _c3_epoch3_diagnostic(),
        endpoint_training_evidence=evidence,
        endpoint_training_chart_path="figures/endpoint.png",
    )
    assert "First-Q ranking CE" in markdown
    assert "figures/endpoint.png" in markdown
    assert "自适应权重或梯度平衡" in markdown
    assert "First-Q 为" in markdown and "First-Q2 为" in markdown
    assert "stop-gradient 的 F-imagined states" in markdown
    assert "McNemar" in markdown and "独立 dev pairs" in markdown


def test_endpoint_training_metrics_fail_closed_on_malformed_declared_curve(
    tmp_path: Path,
) -> None:
    ledger = _write_endpoint_training_fixture(tmp_path)
    path = tmp_path / "training/v1_c2/metrics.csv"
    rows = list(csv.DictReader(path.open(newline="", encoding="utf-8")))
    rows[-1]["validation/loss"] = ""
    with path.open("w", newline="", encoding="utf-8") as stream:
        writer = csv.DictWriter(stream, fieldnames=rows[0], lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)
    ledger_data = json.loads(ledger.read_text())
    ledger_data["training_sources"]["v1_c2"]["files_sha256"]["metrics.csv"] = (
        hashlib.sha256(path.read_bytes()).hexdigest()
    )
    ledger.write_text(json.dumps(ledger_data))

    with pytest.raises(report.CompleteResultsError, match="missing logical epochs"):
        report.load_endpoint_training_evidence(ledger)


def test_endpoint_training_evidence_is_optional_when_not_declared(
    tmp_path: Path,
) -> None:
    ledger = tmp_path / "reconciliation_ledger.json"
    ledger.write_text(json.dumps({"training_sources": {}}))
    assert report.load_endpoint_training_evidence(ledger) is None


def test_partial_endpoint_extension_is_rejected() -> None:
    with pytest.raises(report.CompleteResultsError, match="exactly 8 cells"):
        report._fixed_master_rows(_cells(), _endpoint_cells()[:-1])


def test_c3_epoch3_diagnostic_is_analyzed_but_never_added_to_master_table() -> None:
    endpoint_cells = _endpoint_cells_with_c3_outcomes()
    diagnostic = _c3_epoch3_diagnostic()
    analysis = report._endpoint_analysis(_cells(), endpoint_cells, diagnostic)

    assert analysis.cell_count == 8
    assert analysis.state_v_c3_count == 26
    assert analysis.state_v_c3_epoch3_count == 28
    assert analysis.state_v_c3_epoch3_vs_epoch12_contingency == {
        "both_success": 26,
        "epoch3_only": 2,
        "epoch12_only": 0,
        "both_failure": 22,
    }
    assert analysis.state_v_c3_epoch3_vs_epoch12_exact_mcnemar_p == 0.5

    matrix = report._fixed_master_rows(_cells(), endpoint_cells)
    assert len(matrix) == 26
    assert next(row for row in matrix if row[:2] == ("V1", "C3"))[-1] == (
        "26/50 (52%)"
    )
    markdown = report.build_markdown(
        _cells(),
        _ledger_evidence(),
        endpoint_cells,
        diagnostic,
    )
    assert "E3 28/50 (56%)" in markdown
    assert "E12 26/50 (52%)" in markdown
    assert "exact McNemar 双侧 p=0.5" in markdown
    assert "不进入 8-cell 主表" in markdown


def test_c3_epoch3_diagnostic_requires_strict_endpoint_and_exact_identity() -> None:
    with pytest.raises(report.CompleteResultsError, match="without the strict"):
        report._endpoint_analysis(_cells(), None, _c3_epoch3_diagnostic())

    diagnostic = _c3_epoch3_diagnostic()
    diagnostic["checkpoint_global_step"] = 2_999
    with pytest.raises(report.CompleteResultsError, match="global-step 3000"):
        report._endpoint_analysis(
            _cells(), _endpoint_cells_with_c3_outcomes(), diagnostic
        )


def test_four_version_mean_rows_include_mean_q_for_v0_and_v1() -> None:
    rows = report._fixed_version_mean_rows(_cells())

    assert len(rows) == 4
    assert rows[0][0] == "V0 raw action"
    assert rows[0][5] == "26.7%"
    assert rows[1][0] == "V1 action encoder"
    assert rows[1][5] == "20.0%"
    score_rows = report._fixed_score_mean_rows(_cells())
    mean_q = next(row for row in score_rows if row[0] == "Mean-Q rollout")
    assert mean_q[1] == "V0/V1/V2/V2-EMA"
    assert mean_q[2] == "24"


def test_companion_ledgers_are_counted_and_fingerprinted(tmp_path: Path) -> None:
    identities = sorted(report._expected_identities())
    csv_path = tmp_path / "all_o50_results.csv"
    json_path = tmp_path / "reconciliation_ledger.json"
    _write_ledger(csv_path, identities)
    json_path.write_text(
        json.dumps(
            {
                "cell_count": report.COMPLETE_CELL_COUNT,
                "episode_count_per_cell": report.EPISODES,
                "outcome_count": report.COMPLETE_OUTCOME_COUNT,
                "cells": [
                    {"outcomes": [False] * report.EPISODES}
                    for _ in range(report.COMPLETE_CELL_COUNT)
                ],
            }
        ),
        encoding="utf-8",
    )

    evidence = report.load_ledger_evidence(csv_path, json_path)

    assert evidence.cell_count == 477
    assert evidence.outcome_count == 23_850
    assert evidence.csv_sha256 == hashlib.sha256(csv_path.read_bytes()).hexdigest()
    assert evidence.json_sha256 == hashlib.sha256(json_path.read_bytes()).hexdigest()

    payload = json.loads(json_path.read_text(encoding="utf-8"))
    payload["outcome_count"] -= 1
    json_path.write_text(json.dumps(payload), encoding="utf-8")
    with pytest.raises(report.CompleteResultsError, match="outcome_count"):
        report.load_ledger_evidence(csv_path, json_path)


def test_tied_winner_summaries_are_bounded() -> None:
    winners = tuple(f"V0-M{index}" for index in range(12))

    summary = report._compact_join(winners)

    assert summary == "12 tied: V0-M0, V0-M1, V0-M2, …"
    assert "V0-M3" not in summary


def test_master_column_maxima_reset_at_every_version_boundary() -> None:
    rows = []
    for version_index, (_version, variants) in enumerate(report.MASTER_ROW_GROUPS):
        for method_index in range(len(variants)):
            rows.append(
                tuple(
                    100 * version_index + 10 * column_index + method_index
                    for column_index in range(len(report.MASTER_SCORE_MODES))
                )
            )

    maxima = report._fixed_master_version_column_maxima(tuple(rows))

    assert len(maxima) == len(report.VERSIONS)
    assert maxima[0] == (5, 15, 25, 35, 45, 55, 65)
    assert maxima[1] == (107, 117, 127, 137, 147, 157, 167)
    assert maxima[2] == (205, 215, 225, 235, 245, 255, 265)
    assert maxima[3] == (305, 315, 325, 335, 345, 355, 365)


def test_dynamic_maxima_ignore_missing_cells_and_empty_columns() -> None:
    rows = [list(row) for row in report._fixed_master_counts(_cells())]
    maxima = report._fixed_master_version_column_maxima(rows)

    assert maxima[0][-2:] == (None, None)
    assert maxima[1][-2:] == (None, None)
    assert maxima[2][-2:] == (None, None)
    assert maxima[3][-2:] == (None, None)


def test_docx_starts_with_complete_decision_view_and_has_full_master_matrix() -> None:
    from docx import Document
    from docx.oxml.ns import qn

    png = _tiny_png()
    evidence = _ledger_evidence()
    payload = report.build_docx(
        _cells(),
        evidence,
        _endpoint_cells(),
        v0_training_chart=png,
        v2_training_chart=png,
        training_chart=png,
        score_chart=png,
    )
    document = Document(io.BytesIO(payload))
    expected_header = [
        "Version",
        "Method",
        "Training loss",
        "F-only",
        "G-only",
        "F+G tail",
        "First-Q",
        "Mean-Q",
        "First-Q2",
        "State-V",
    ]
    table_headers = [
        [cell.text for cell in table.rows[0].cells] for table in document.tables
    ]

    assert table_headers[:6] == [
        ["Question", "Answer", "Evidence", "Decision"],
        ["Artifact", "Coverage", "Repository-relative path", "SHA-256"],
        ["Common protocol field", "Fixed setting"],
        ["Score", "Actual F and G path", "Cost minimized by CEM", "What is used"],
        ["Marker", "Meaning", "Comparison scope"],
        expected_header,
    ]
    master = document.tables[table_headers.index(expected_header)]
    assert len(master.rows) == 27
    assert all(len(row.cells) == 10 for row in master.rows)
    assert all(cell.text for row in master.rows[1:] for cell in row.cells[2:])
    expected_losses = tuple(
        report.METHOD_LOSS_LABELS[variant]
        for _version, variant in report._master_row_metadata()
    )
    assert tuple(row.cells[2].text for row in master.rows[1:]) == expected_losses
    master_text = "\n".join(cell.text for row in master.rows for cell in row.cells)
    assert "—" in master_text
    assert "–" not in master_text

    count_rows = report._fixed_master_counts(_cells(), _endpoint_cells())
    version_maxima = report._fixed_master_version_column_maxima(count_rows)
    group_by_row = tuple(
        group_index
        for group_index, (_version, variants) in enumerate(report.MASTER_ROW_GROUPS)
        for _variant in variants
    )
    highlight_fills = {"FFF2CC", "DDEBF7", "B7DEE8"}
    for row_index, (row, counts) in enumerate(zip(master.rows[1:], count_rows)):
        version_index = group_by_row[row_index]
        row_maximum = max(count for count in counts if count is not None)
        for score_index, (cell, count) in enumerate(zip(row.cells[3:], counts)):
            shading = cell._tc.tcPr.find(qn("w:shd"))
            fill = "" if shading is None else (shading.get(qn("w:fill")) or "").upper()
            row_best = count is not None and count == row_maximum
            column_best = (
                count is not None
                and version_maxima[version_index][score_index] is not None
                and count == version_maxima[version_index][score_index]
            )
            expected = (
                "B7DEE8"
                if row_best and column_best
                else "DDEBF7"
                if column_best
                else "FFF2CC"
                if row_best
                else None
            )
            if expected is None:
                assert fill not in highlight_fills
            else:
                assert fill == expected

            borders = cell._tc.tcPr.find(qn("w:tcBorders"))
            assert borders is None or all(
                (edge.get(qn("w:color")) or "").upper() != "2F75B5"
                for edge in borders
            )

    legend = document.tables[4]
    legend_fills = []
    for row in legend.rows[1:4]:
        shading = row.cells[0]._tc.tcPr.find(qn("w:shd"))
        legend_fills.append((shading.get(qn("w:fill")) or "").upper())
    assert legend_fills == ["FFF2CC", "DDEBF7", "B7DEE8"]

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    ordered_front_sections = [
        "Decision summary",
        "Complete companion ledgers",
        "How the seven evaluation methods are run",
        "26-method, seven-score master comparison and color legend",
        "One 26 by 7 master matrix with training losses",
        "Best training method and evaluation score",
        "Causes and next objectives",
        "Coverage map",
    ]
    indices = [paragraphs.index(text) for text in ordered_front_sections]
    assert indices == sorted(indices)
    assert "six methods × three score modes" not in "\n".join(paragraphs)

    evidence_text = "\n".join(
        cell.text for row in document.tables[1].rows for cell in row.cells
    )
    assert evidence.csv_path in evidence_text
    assert evidence.json_path in evidence_text
    assert evidence.csv_sha256 in evidence_text
    assert evidence.json_sha256 in evidence_text
    assert "477" in evidence_text
    assert "23,850" in evidence_text
    assert "477" in document.core_properties.subject
    assert "26-by-7" in document.core_properties.subject

    audit_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert report.FIXED_SELECTION_RANKS_SHA256 in audit_text
    assert report.FIXED_SELECTION_RANKS_SHA256.startswith("88c204770")


def test_docx_can_preserve_previous_reference_pages(tmp_path: Path) -> None:
    from docx import Document

    base = Document()
    base.add_paragraph("PRESERVED PREVIOUS RESULTS TD CONTENT")
    base_path = tmp_path / "base.docx"
    base.save(base_path)
    png = _tiny_png()

    payload = report.build_docx(
        _cells(),
        _ledger_evidence(),
        v0_training_chart=png,
        v2_training_chart=png,
        training_chart=png,
        score_chart=png,
        base_document=base_path,
    )
    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "PRESERVED PREVIOUS RESULTS TD CONTENT" in text
    assert "Complete 26-method, seven-score results and analysis" in text
    assert "How the seven evaluation methods are run" in text


def test_docx_includes_endpoint_training_diagnostics_without_changing_master(
    tmp_path: Path,
) -> None:
    from docx import Document

    evidence = report.load_endpoint_training_evidence(
        _write_endpoint_training_fixture(tmp_path)
    )
    assert evidence is not None
    png = _tiny_png()
    payload = report.build_docx(
        _cells(),
        _ledger_evidence(),
        _endpoint_cells_with_c3_outcomes(),
        c3_epoch3_diagnostic=_c3_epoch3_diagnostic(),
        endpoint_training_evidence=evidence,
        endpoint_training_chart=png,
        v0_training_chart=png,
        v2_training_chart=png,
        training_chart=png,
        score_chart=png,
    )
    document = Document(io.BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    headers = [[cell.text for cell in table.rows[0].cells] for table in document.tables]
    master_header = [
        "Version",
        "Method",
        "Training loss",
        "F-only",
        "G-only",
        "F+G tail",
        "First-Q",
        "Mean-Q",
        "First-Q2",
        "State-V",
    ]
    master = document.tables[headers.index(master_header)]

    assert "Endpoint interpretation and next steps" in text
    assert "robust cross-readout ranking gain" in text
    assert "held-out development pairs" in text
    assert "V1-C2/C3 endpoint training evidence" in text
    assert "adaptive or gradient-balanced weighting" in text
    assert len(master.rows) == 27
    assert all(len(row.cells) == 10 for row in master.rows)
